#!/usr/bin/env python3
"""Biblioteka wspolna: markdown korpusu -> model dokumentu -> DOCX.

Korpus pochodzi z konwersji .docx/.pdf/.xlsx do markdown, wiec naglowki
wystepuja w dwoch konwencjach naraz:
  * markdownowej   (#, ##, ###)
  * "pogrubieniowej" (**1\\. Tytul**, **1.1 Tytul**, **CZESC I — ...**)
Parser rozpoznaje obie i sprowadza je do jednej hierarchii poziomow.
"""
import re

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

GRANAT = RGBColor(0x14, 0x2B, 0x45)
STAL = RGBColor(0x4A, 0x5A, 0x6A)
AKCENT = RGBColor(0x8C, 0x1D, 0x2C)

# ---------------------------------------------------------------- parser ---
RE_MD_H = re.compile(r"^(#{1,6})\s+(.+?)\s*#*$")
RE_CZESC = re.compile(r"^\*\*\s*(CZ[EĘ][SŚ][CĆ][^*]{0,120}?)\s*\*\*\s*$")
RE_MODUL = re.compile(r"^\*\*\s*(Modu[lł]\s+[A-ZSCDX]?\d+\s*:[^*]{0,120}?)\s*\*\*\s*$", re.I)
RE_NUM = re.compile(r"^\*\*\s*(\d+(?:\.\d+)*)\s*\\?\.?\s+([^*]{2,140}?)\s*\*\*\s*$")
RE_BOLD_ONLY = re.compile(r"^\*\*\s*([^*]{3,140}?)\s*\*\*\s*[.:]?$")
RE_TABLE = re.compile(r"^\s*\|.*\|\s*$")
RE_SEP = re.compile(r"^\s*\|[\s:|-]+\|\s*$")
RE_LI = re.compile(r"^\s*([-*+]|\d+[.)])\s+(.*)$")
RE_HR = re.compile(r"^\s*(-{3,}|\*{3,}|_{3,})\s*$")


def czysc(t: str) -> str:
    """Zdejmuje artefakty konwersji, zostawia tresc."""
    t = t.replace("\\.", ".").replace("\\-", "-").replace("\\_", "_")
    t = t.replace("�", "")
    return t.strip()


def parsuj(md: str):
    """markdown -> lista elementow: ('h',poziom,tekst) ('p',tekst) ('tab',wiersze) ('li',poziom,tekst)"""
    linie = md.split("\n")
    el, i = [], 0
    while i < len(linie):
        ln = linie[i].rstrip()
        if not ln.strip() or RE_HR.match(ln):
            i += 1
            continue

        # tabela
        if RE_TABLE.match(ln):
            wiersze = []
            while i < len(linie) and RE_TABLE.match(linie[i].rstrip()):
                row = linie[i].strip()
                if not RE_SEP.match(row):
                    kom = [czysc(c) for c in row.strip("|").split("|")]
                    wiersze.append(kom)
                i += 1
            if wiersze:
                szer = max(len(w) for w in wiersze)
                wiersze = [w + [""] * (szer - len(w)) for w in wiersze]
                el.append(("tab", wiersze))
            continue

        m = RE_MD_H.match(ln)
        if m:
            el.append(("h", min(len(m.group(1)), 5), czysc(m.group(2))))
            i += 1
            continue
        m = RE_CZESC.match(ln)
        if m:
            el.append(("h", 1, czysc(m.group(1))))
            i += 1
            continue
        m = RE_MODUL.match(ln)
        if m:
            el.append(("h", 1, czysc(m.group(1))))
            i += 1
            continue
        m = RE_NUM.match(ln)
        if m:
            poziom = min(1 + m.group(1).count(".") + 1, 5)
            el.append(("h", poziom, czysc(f"{m.group(1)}. {m.group(2)}")))
            i += 1
            continue
        m = RE_LI.match(ln)
        if m:
            wciecie = len(linie[i]) - len(linie[i].lstrip())
            el.append(("li", 1 if wciecie < 3 else 2, czysc(m.group(2))))
            i += 1
            continue
        m = RE_BOLD_ONLY.match(ln)
        if m and len(m.group(1)) < 110 and not m.group(1).endswith((".", "?")):
            el.append(("h", 4, czysc(m.group(1))))
            i += 1
            continue

        # akapit — sklejamy do pustej linii
        buf = [ln]
        i += 1
        while i < len(linie) and linie[i].strip() and not RE_TABLE.match(linie[i]) \
                and not RE_MD_H.match(linie[i]) and not RE_LI.match(linie[i]) \
                and not RE_HR.match(linie[i]):
            buf.append(linie[i].rstrip())
            i += 1
        el.append(("p", czysc(" ".join(buf))))
    return el


# ------------------------------------------------------------- skladanie ---
RE_RUN = re.compile(r"(\*\*[^*]+\*\*|\*[^*]+\*|`[^`]+`|\[[^\]]+\]\([^)]*\))")


def wpisz_runy(p, tekst: str):
    """Zapisuje tekst z zachowaniem pogrubien, kursywy i kodu."""
    for cz in RE_RUN.split(tekst):
        if not cz:
            continue
        if cz.startswith("**") and cz.endswith("**"):
            p.add_run(cz[2:-2]).bold = True
        elif cz.startswith("*") and cz.endswith("*") and len(cz) > 2:
            p.add_run(cz[1:-1]).italic = True
        elif cz.startswith("`") and cz.endswith("`"):
            r = p.add_run(cz[1:-1])
            r.font.name = "Consolas"
            r.font.size = Pt(9)
        elif cz.startswith("["):
            m = re.match(r"\[([^\]]+)\]\(([^)]*)\)", cz)
            p.add_run(m.group(1) if m else cz)
        else:
            p.add_run(cz)


_ZAKLADKI = {}


def zakladka(paragraf, nazwa: str):
    """Zakladka Worda — cel odsylaczy i pol PAGEREF.

    Identyfikatory musza byc unikalne w obrebie pliku, a nazwy skracane do
    38 znakow moga sie zejsc — przy kolizji doklejamy licznik, inaczej
    PAGEREF wskazywalby nie to miejsce.
    """
    baza = re.sub(r"[^A-Za-z0-9_]", "_", nazwa)[:34].strip("_") or "Z"
    nazwa = baza
    i = 1
    while nazwa in _ZAKLADKI:
        i += 1
        nazwa = f"{baza}_{i}"
    ident = len(_ZAKLADKI) + 1000
    _ZAKLADKI[nazwa] = ident

    start = OxmlElement("w:bookmarkStart")
    start.set(qn("w:id"), str(ident))
    start.set(qn("w:name"), nazwa)
    end = OxmlElement("w:bookmarkEnd")
    end.set(qn("w:id"), str(ident))
    paragraf._p.insert(0, start)
    paragraf._p.append(end)
    return nazwa


def pole(paragraf, instrukcja: str, tekst_zastepczy: str = "—"):
    """Pole Worda (TOC, PAGEREF, PAGE) — Word wylicza wartosc przy aktualizacji."""
    r1 = paragraf.add_run()
    fld = OxmlElement("w:fldChar")
    fld.set(qn("w:fldCharType"), "begin")
    r1._r.append(fld)
    r2 = paragraf.add_run()
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = instrukcja
    r2._r.append(instr)
    r3 = paragraf.add_run()
    sep = OxmlElement("w:fldChar")
    sep.set(qn("w:fldCharType"), "separate")
    r3._r.append(sep)
    r4 = paragraf.add_run(tekst_zastepczy)
    r5 = paragraf.add_run()
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    r5._r.append(end)
    return r4


def hiperlacze(paragraf, tekst: str, cel: str, wewnetrzne: bool = False):
    """Odsylacz: wewnetrzny (do zakladki) albo zewnetrzny (do innego pliku).

    Cel zewnetrzny moze wskazywac zakladke w innym pliku: "plik.docx#ZAKLADKA".
    Word wymaga wtedy relacji do samego pliku ORAZ atrybutu w:anchor —
    fragment zostawiony w adresie relacji nie zadziala.
    """
    h = OxmlElement("w:hyperlink")
    if wewnetrzne:
        h.set(qn("w:anchor"), cel)
    else:
        plik, _, kotwica = cel.partition("#")
        rid = paragraf.part.relate_to(
            plik, "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
            is_external=True)
        h.set(qn("r:id"), rid)
        if kotwica:
            h.set(qn("w:anchor"), kotwica)
    r = OxmlElement("w:r")
    rPr = OxmlElement("w:rPr")
    col = OxmlElement("w:color")
    col.set(qn("w:val"), "1F4E79")
    u = OxmlElement("w:u")
    u.set(qn("w:val"), "single")
    rPr.append(col)
    rPr.append(u)
    r.append(rPr)
    t = OxmlElement("w:t")
    t.text = tekst
    r.append(t)
    h.append(r)
    paragraf._p.append(h)


def nowy_dokument(tytul: str, autor: str = "Eternal Labs Sp. z o.o."):
    doc = Document()
    doc.core_properties.title = tytul
    doc.core_properties.author = autor
    doc.core_properties.language = "pl-PL"

    s = doc.sections[0]
    s.page_height, s.page_width = Cm(29.7), Cm(21.0)
    s.top_margin = s.bottom_margin = Cm(2.2)
    s.left_margin, s.right_margin = Cm(2.5), Cm(2.0)

    st = doc.styles["Normal"]
    st.font.name = "Calibri"
    st.font.size = Pt(10.5)
    st.paragraph_format.space_after = Pt(6)
    st.paragraph_format.line_spacing = 1.13

    for lvl, rozm, kolor, przed in ((1, 17, GRANAT, 20), (2, 13.5, GRANAT, 15),
                                    (3, 11.5, GRANAT, 11), (4, 10.5, STAL, 9),
                                    (5, 10, STAL, 7)):
        h = doc.styles[f"Heading {lvl}"]
        h.font.name = "Calibri"
        h.font.size = Pt(rozm)
        h.font.color.rgb = kolor
        h.font.bold = True
        h.paragraph_format.space_before = Pt(przed)
        h.paragraph_format.space_after = Pt(4)
        h.paragraph_format.keep_with_next = True
    return doc


def stopka_z_numeracja(doc, tekst_lewy: str):
    for sek in doc.sections:
        p = sek.footer.paragraphs[0]
        p.text = ""
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(tekst_lewy + "   ·   s. ")
        r.font.size = Pt(8)
        r.font.color.rgb = STAL
        pole(p, " PAGE ", "1").font.size = Pt(8)
        r2 = p.add_run(" z ")
        r2.font.size = Pt(8)
        r2.font.color.rgb = STAL
        pole(p, " NUMPAGES ", "1").font.size = Pt(8)


def strona_tytulowa(doc, nadtytul, tytul, podtytul, metryka: list):
    for _ in range(3):
        doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(nadtytul)
    r.font.size = Pt(12)
    r.font.color.rgb = AKCENT
    r.bold = True

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(tytul)
    r.font.size = Pt(30)
    r.font.color.rgb = GRANAT
    r.bold = True

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(podtytul)
    r.font.size = Pt(12.5)
    r.font.color.rgb = STAL

    doc.add_paragraph()
    t = doc.add_table(rows=0, cols=2)
    t.style = "Light List Accent 1"
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    for k, v in metryka:
        c = t.add_row().cells
        c[0].text = k
        c[1].text = str(v)
        for par in c[0].paragraphs:
            for run in par.runs:
                run.bold = True
                run.font.size = Pt(9.5)
        for par in c[1].paragraphs:
            for run in par.runs:
                run.font.size = Pt(9.5)
    doc.add_page_break()


def spis_tresci(doc, poziomy="1-3"):
    doc.add_heading("Spis treści", level=1)
    p = doc.add_paragraph()
    pole(p, f' TOC \\o "{poziomy}" \\h \\z \\u ',
         "Spis treści zostanie zbudowany po otwarciu w Wordzie "
         "(zaznacz wszystko i naciśnij F9).")
    doc.add_page_break()


def dodaj_tabele(doc, wiersze, szer_pierwszej=None):
    if not wiersze:
        return
    t = doc.add_table(rows=0, cols=len(wiersze[0]))
    t.style = "Light Grid Accent 1"
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    for idx, w in enumerate(wiersze):
        cells = t.add_row().cells
        for c, tekst in zip(cells, w):
            c.text = ""
            par = c.paragraphs[0]
            par.paragraph_format.space_after = Pt(2)
            wpisz_runy(par, tekst)
            for run in par.runs:
                run.font.size = Pt(8.5)
                if idx == 0:
                    run.bold = True
    return t


def zrodlo(doc, teksty):
    """Linijka zrodla pod elementem — potwierdzenie slusznosci tezy."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(1)
    p.paragraph_format.space_after = Pt(10)
    p.paragraph_format.left_indent = Cm(0.3)
    r = p.add_run("Źródło: ")
    r.font.size = Pt(8)
    r.bold = True
    r.font.color.rgb = AKCENT
    r2 = p.add_run(teksty)
    r2.font.size = Pt(8)
    r2.italic = True
    r2.font.color.rgb = STAL
    return p


def wstaw_elementy(doc, elementy, przesuniecie=0, max_poziom=5,
                   kotwica_prefix=None, zebrane_kotwice=None):
    """Wstawia sparsowane elementy do dokumentu, nadajac zakladki naglowkom."""
    for e in elementy:
        if e[0] == "h":
            poziom = min(max(e[1] + przesuniecie, 1), max_poziom)
            h = doc.add_heading("", level=poziom)
            wpisz_runy(h, e[2])
            if kotwica_prefix is not None and poziom <= 3:
                nazwa = zakladka(h, f"{kotwica_prefix}{re.sub(r'[^A-Za-z0-9]', '_', e[2])[:30]}")
                if zebrane_kotwice is not None:
                    zebrane_kotwice.append((poziom, e[2], nazwa))
        elif e[0] == "p":
            wpisz_runy(doc.add_paragraph(), e[1])
        elif e[0] == "li":
            p = doc.add_paragraph(style="List Bullet" if e[1] == 1 else "List Bullet 2")
            wpisz_runy(p, e[2])
        elif e[0] == "tab":
            dodaj_tabele(doc, e[1])
            doc.add_paragraph().paragraph_format.space_after = Pt(2)
