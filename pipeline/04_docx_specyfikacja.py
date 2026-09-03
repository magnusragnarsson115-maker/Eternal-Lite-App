#!/usr/bin/env python3
"""Krok 4 — SPECYFIKACJA TECHNICZNA (.docx) z calej tresci korpusu.

Trzon      : dokument o najwyzszym autorytecie w grupie specyfikacyjnej
             (Master 5.4 FINAL) — w calosci, bo po deduplikacji zawiera
             tresc wszystkich wczesniejszych wersji.
Aneks A/B  : tresc z pozostalych dokumentow G1 i G2, ktora NIE zostala
             pokryta przez trzon — czyli dokladnie to, co kazdy z tych
             plikow wnosi ponad specyfikacje glowna.
Aneks C/D  : indeks dokumentow zrodlowych i rejestr zrodel zewnetrznych.

Kazdy naglowek trzonu dostaje zakladke SPEC_* — biznesplan linkuje do niej.
"""
import json
import os
from collections import defaultdict

from docx.shared import Pt

import docxlib as L

KANON = ".work/kanon.jsonl"
IDX = "out/indeks_zrodel.json"
AUTH = ".work/autorytet.json"
RAPORT = ".work/dedup_raport.json"
KOTWICE = ".work/kotwice_spec.json"
WYJSCIE = "out/ETERNAL_Specyfikacja_Techniczna.docx"

GRUPY_SPEC = ("G1", "G2")


def main():
    kanon = [json.loads(l) for l in open(KANON, encoding="utf-8")]
    idx = json.load(open(IDX, encoding="utf-8"))
    auth = json.load(open(AUTH, encoding="utf-8"))
    raport = json.load(open(RAPORT, encoding="utf-8"))
    grupy = idx["grupy"]
    meta = {d["id"]: d for d in idx["dokumenty"]}

    wg_doc = defaultdict(list)
    for b in kanon:
        wg_doc[b["doc"]].append(b)
    for v in wg_doc.values():
        v.sort(key=lambda b: b["n"])

    spec_docs = [d for d in meta if meta[d]["grupa"] in GRUPY_SPEC and wg_doc.get(d)]
    trzon = max(spec_docs, key=lambda d: auth.get(d, {}).get("punkty", 0.0))
    pozostale = sorted((d for d in spec_docs if d != trzon),
                       key=lambda d: (meta[d]["grupa"],
                                      -sum(len(b["txt"]) for b in wg_doc[d])))

    def tekst_doc(doc_id):
        out = []
        for b in wg_doc[doc_id]:
            out.append(b["txt"])
            if b["uzupelnienia"]:
                out.append("")
                out.append("**Treść unikalna odzyskana z wariantów równoległych "
                           "tego samego ustalenia:**")
                out.extend(f"- {u['txt']}  *[{u['doc']}]*" for u in b["uzupelnienia"])
            out.append("")
        return "\n".join(out)

    znakow_trzon = sum(len(b["txt"]) for b in wg_doc[trzon])
    znakow_aneks = sum(len(b["txt"]) for d in pozostale for b in wg_doc[d])

    doc = L.nowy_dokument("ETERNAL — Specyfikacja funkcjonalno-techniczna")
    L.strona_tytulowa(
        doc, "ETERNAL ECOSYSTEM",
        "SPECYFIKACJA\nFUNKCJONALNO-TECHNICZNA",
        "Dokument scalony z całego korpusu źródłowego\npo deduplikacji zwykłej i znaczeniowej",
        [("Trzon dokumentu", f"{meta[trzon]['nazwa']} [{trzon}]"),
         ("Dokumentów źródłowych", f"{len(idx['dokumenty'])}"),
         ("W tym wykorzystanych tu", f"{len(spec_docs)} (grupy G1 i G2)"),
         ("Objętość trzonu", f"{znakow_trzon:,} znaków".replace(",", " ")),
         ("Objętość aneksów", f"{znakow_aneks:,} znaków".replace(",", " ")),
         ("Usunięto duplikatów 1:1", f"{raport['usunieto_1_1']:,}".replace(",", " ")),
         ("Scalono wariantów znaczeniowych", f"{raport['scalonych_wariantow']:,}".replace(",", " ")),
         ("Status", "POUFNE — do użytku wewnętrznego")])

    # --- nota metodyczna ---
    doc.add_heading("Nota metodyczna — jak powstał ten dokument", level=1)
    doc.add_paragraph(
        f"Dokument nie jest nową redakcją specyfikacji. Jest scaleniem "
        f"{len(idx['dokumenty'])} plików źródłowych — konwersacji roboczych, opracowań, "
        f"arkuszy i wcześniejszych wersji specyfikacji — przeprowadzonym w czterech krokach.")
    for i, (tyt, opis) in enumerate([
        ("Scalenie", f"Wszystkie pliki źródłowe połączono w jeden korpus "
                     f"({raport['blokow_zrodlowych']:,} bloków treści)."),
        ("Deduplikacja zwykła (1:1)", f"Usunięto {raport['usunieto_1_1']:,} bloków identycznych "
                                      "co do znaku, występujących w kilku plikach naraz. "
                                      "Powtórzenia wewnątrz jednego pliku zachowano — to jego "
                                      "struktura (nagłówki tabel przy każdej karcie funkcji), "
                                      "a nie redundancja."),
        ("Deduplikacja zaawansowana (znaczeniowa)",
         f"Zgrupowano {raport['scalonych_wariantow']:,} bloków o tym samym znaczeniu, "
         "wyrażonych innymi słowami. Z każdej grupy zachowano wariant najbardziej "
         "aktualny i najpełniejszy, mierzony numerem wersji w nazwie pliku, datą "
         "dokumentu, znacznikami typu FINAL/SCALONA/KOMPLETNA oraz kompletnością. "
         f"Z wariantów odrzuconych odzyskano {raport['odzyskanych_zdan']:,} zdań niosących "
         "treść, której w wersji zachowanej nie było — są w dokumencie oznaczone."),
        ("Złożenie", "Trzonem jest dokument o najwyższym autorytecie w grupie "
                     "specyfikacyjnej. Aneksy zawierają wyłącznie to, co pozostałe pliki "
                     "wnoszą ponad trzon."),
    ], 1):
        p = doc.add_paragraph(style="List Number")
        p.add_run(tyt + ". ").bold = True
        p.add_run(opis)

    doc.add_heading("Hierarchia źródeł — co obowiązuje przy rozbieżności", level=2)
    doc.add_paragraph(
        "Przy sprzeczności między plikami obowiązuje pozycja wyżej w tabeli. "
        "Punktacja jest wyliczona, nie uznaniowa: numer wersji w nazwie × 2,5, "
        "data dokumentu, znaczniki wersji obowiązującej, kompletność treści.")
    wiersze = [["#", "Dokument", "Grupa", "Wersja", "Data", "Punkty", "Rola"]]
    for i, d in enumerate(sorted(spec_docs, key=lambda d: -auth[d]["punkty"])[:18], 1):
        a = auth[d]
        wiersze.append([str(i), f"{meta[d]['nazwa']} [{d}]", meta[d]["grupa"],
                        f"{a['wersja']:g}" if a["wersja"] else "—", a["data"] or "—",
                        f"{a['punkty']:.1f}", "TRZON" if d == trzon else "aneks"])
    L.dodaj_tabele(doc, wiersze)
    doc.add_page_break()

    L.spis_tresci(doc, "1-3")

    # --- trzon ---
    kotwice = []
    doc.add_heading("CZĘŚĆ GŁÓWNA — SPECYFIKACJA", level=1)
    L.zrodlo(doc, f"{meta[trzon]['nazwa']} [{trzon}], grupa {meta[trzon]['grupa']} — "
                  f"{grupy[meta[trzon]['grupa']]}; pakiet {meta[trzon]['pakiet']}. "
                  f"Dokument o najwyższym autorytecie w korpusie "
                  f"({auth[trzon]['punkty']:.1f} pkt), po deduplikacji zawiera treść "
                  f"wcześniejszych wersji specyfikacji.")
    L.wstaw_elementy(doc, L.parsuj(tekst_doc(trzon)), przesuniecie=0,
                     kotwica_prefix="SPEC_", zebrane_kotwice=kotwice)

    # --- aneksy A i B ---
    for grupa_id, litera, tytul in (("G1", "A", "SPECYFIKACJE I ARCHITEKTURA"),
                                    ("G2", "B", "REJESTRY FUNKCJI I MACIERZE")):
        czlonkowie = [d for d in pozostale if meta[d]["grupa"] == grupa_id]
        if not czlonkowie:
            continue
        doc.add_page_break()
        doc.add_heading(f"ANEKS {litera} — {tytul}", level=1)
        doc.add_paragraph(
            f"Poniżej wyłącznie treść, której nie ma w części głównej. "
            f"Dokumentów: {len(czlonkowie)}. Pod nagłówkiem każdego z nich podano plik "
            f"źródłowy, jego grupę i objętość wniesioną ponad trzon.")
        for d in czlonkowie:
            znakow = sum(len(b["txt"]) for b in wg_doc[d])
            if znakow < 200:
                continue
            doc.add_heading(f"{litera}.{czlonkowie.index(d) + 1} {meta[d]['nazwa']}", level=2)
            a = auth.get(d, {})
            L.zrodlo(doc, f"{meta[d]['nazwa']} [{d}] · grupa {grupa_id} — {grupy[grupa_id]} · "
                          f"pakiet {meta[d]['pakiet']} · autorytet {a.get('punkty', 0):.1f} pkt"
                          + (f" · data {a['data']}" if a.get("data") else "")
                          + f" · wnosi {znakow:,} znaków ponad część główną".replace(",", " "))
            L.wstaw_elementy(doc, L.parsuj(tekst_doc(d)), przesuniecie=2)

    # --- aneks C: indeks dokumentow zrodlowych ---
    doc.add_page_break()
    doc.add_heading("ANEKS C — INDEKS DOKUMENTÓW ŹRÓDŁOWYCH SPECYFIKACJI", level=1)
    doc.add_paragraph(
        "Pełny wykaz plików, z których zbudowano ten dokument. Kolumna „wniesiono” "
        "podaje objętość treści, która przeszła do specyfikacji po deduplikacji — "
        "zero oznacza, że cała treść pliku występuje już w innym, nowszym dokumencie.")
    wiersze = [["ID", "Plik źródłowy", "Grupa", "Pakiet", "Znaków źródłowo", "Wniesiono"]]
    for d in sorted(spec_docs, key=lambda d: (meta[d]["grupa"], meta[d]["nazwa"])):
        wiersze.append([d, meta[d]["nazwa"], meta[d]["grupa"], meta[d]["pakiet"],
                        f"{meta[d]['znaki']:,}".replace(",", " "),
                        f"{sum(len(b['txt']) for b in wg_doc[d]):,}".replace(",", " ")])
    L.dodaj_tabele(doc, wiersze)

    # --- aneks D: zrodla zewnetrzne ---
    doc.add_page_break()
    doc.add_heading("ANEKS D — ŹRÓDŁA ZEWNĘTRZNE", level=1)
    doc.add_paragraph(
        "Podstawy prawne i zasoby infrastruktury publicznej, na których opierają się "
        "ustalenia regulacyjne specyfikacji. Wykaz pochodzi z korpusu źródłowego "
        "(rozdział „Indeks źródeł” specyfikacji głównej) i wymaga weryfikacji stanu "
        "prawnego przed użyciem zewnętrznym.")
    zrodla_zewn = [
        ["Podstawa", "Gdzie", "Czego dotyczy w specyfikacji"],
        ["MDR — rozporządzenie (UE) 2017/745", "eur-lex.europa.eu, zał. I, II, III, VIII, XIV",
         "Klasyfikacja wyrobu, granica regulacyjna, obowiązki wytwórcy"],
        ["IVDR — rozporządzenie (UE) 2017/746", "eur-lex.europa.eu",
         "Diagnostyka in vitro, moduł biochemii Station"],
        ["EHDS — rozporządzenie (UE) 2025/327", "eur-lex.europa.eu",
         "Wtórne wykorzystanie danych zdrowotnych, interoperacyjność"],
        ["AI Act — rozporządzenie (UE) 2024/1689", "eur-lex.europa.eu, zał. I i III, art. 50",
         "Nadzór nad AI, moduł A18, obowiązki informacyjne"],
        ["Dyrektywa (UE) 2024/2853 o odpowiedzialności za produkt", "eur-lex.europa.eu",
         "Odpowiedzialność za oprogramowanie i AI"],
        ["Ustawa o wyrobach medycznych z 7.04.2022 (Dz.U. 2022 poz. 974)", "isap.sejm.gov.pl",
         "Wprowadzenie do obrotu w Polsce, obowiązki wobec URPL"],
        ["Ustawa o systemie informacji w ochronie zdrowia", "isap.sejm.gov.pl",
         "Integracja z P1, EDM"],
        ["Ustawa o krajowym systemie cyberbezpieczeństwa (NIS2)", "isap.sejm.gov.pl",
         "Wymogi bezpieczeństwa infrastruktury"],
        ["Wytyczne MDCG", "health.ec.europa.eu — sekcja wyrobów medycznych",
         "Kwalifikacja oprogramowania jako wyrobu"],
        ["URPL — organ kompetentny", "gov.pl/web/urpl", "Rejestracja, nadzór rynku"],
        ["Wykaz jednostek notyfikowanych", "Baza NANDO Komisji Europejskiej",
         "Wybór jednostki do oceny zgodności"],
        ["RPWDL 2.0", "rpwdl.ezdrowie.gov.pl", "Rejestr podmiotów leczniczych, model kliniczny"],
        ["P1, EPP, IKP — Centrum e-Zdrowia", "ezdrowie.gov.pl",
         "Integracja z systemem publicznym, pozycja wobec państwa"],
        ["EUDAMED", "ec.europa.eu/tools/eudamed", "Rejestracja wyrobu i UDI"],
        ["Profile IHE (XDS.b, ATNA), HL7 CDA", "ihe.net, materiały CeZ",
         "Interoperacyjność dokumentacji medycznej"],
        ["Profile Bluetooth SIG (GATT)", "bluetooth.com/specifications",
         "Warstwa abstrakcji urządzeń, agregacja z opasek"],
        ["Normy PN, IEC 62304, IEC 62366, ISO 13485, ISO 14971", "pkn.pl, iec.ch, iso.org",
         "Cykl życia oprogramowania, użyteczność, jakość, ryzyko"],
    ]
    L.dodaj_tabele(doc, zrodla_zewn)

    L.stopka_z_numeracja(doc, "ETERNAL — Specyfikacja funkcjonalno-techniczna · POUFNE")
    doc.save(WYJSCIE)

    json.dump({"plik": os.path.basename(WYJSCIE), "trzon": trzon,
               "kotwice": [{"poziom": p, "tytul": t, "zakladka": z} for p, t, z in kotwice]},
              open(KOTWICE, "w", encoding="utf-8"), ensure_ascii=False, indent=1)

    print(f"{WYJSCIE}  ({os.path.getsize(WYJSCIE):,} B)")
    print(f"  trzon: {meta[trzon]['nazwa']} [{trzon}] — {znakow_trzon:,} znakow")
    print(f"  aneksy: {len(pozostale)} dokumentow — {znakow_aneks:,} znakow")
    print(f"  zakladek do linkowania z biznesplanu: {len(kotwice)}")


if __name__ == "__main__":
    main()
