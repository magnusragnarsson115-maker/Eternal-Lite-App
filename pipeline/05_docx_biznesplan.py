#!/usr/bin/env python3
"""Krok 5 — BIZNESPLAN (.docx) z calej tresci korpusu.

Roznice wobec specyfikacji, wynikajace z wymagan:
  * pod KAZDYM elementem linia zrodla — plik zrodlowy, jego grupa oraz
    podstawa zewnetrzna (akt prawny, zasob publiczny), jesli sekcja sie na
    nia powoluje;
  * pod kazdym elementem odsylacz do specyfikacji technicznej — hiperlacze
    do zakladki w drugim pliku .docx;
  * pelny indeks WSZYSTKICH dokumentow zrodlowych: nazwa, grupa, strona.

Trzon: Biznesplan 4.0. Materialem dowodowym sa dokumenty grup biznesowych,
przypisane do sekcji biznesplanu wedlug tresci, kazdy z wlasna linia zrodla.
"""
import json
import os
import re
import unicodedata
from collections import defaultdict

from docx.shared import Cm, Pt

import docxlib as L

KANON = ".work/kanon.jsonl"
IDX = "out/indeks_zrodel.json"
AUTH = ".work/autorytet.json"
RAPORT = ".work/dedup_raport.json"
KOTWICE = ".work/kotwice_spec.json"
WYJSCIE = "out/ETAP5_ETERNAL_Biznesplan.docx"
PLIK_SPEC = "ETAP4_ETERNAL_Specyfikacja_Techniczna.docx"

GRUPY_BIZNES = ("G3", "G5", "G6", "G8", "G4")

# podstawy zewnetrzne rozpoznawane w tresci sekcji -> co przywolac pod sekcja
PODSTAWY = [
    (r"\bMDR\b|2017/745|wyrob(?:u|em|y)? medyczn",
     "MDR — rozporządzenie (UE) 2017/745 (eur-lex.europa.eu)"),
    (r"\bIVDR\b|2017/746|in vitro", "IVDR — rozporządzenie (UE) 2017/746 (eur-lex.europa.eu)"),
    (r"\bEHDS\b|2025/327|przestrzen danych zdrowotnych",
     "EHDS — rozporządzenie (UE) 2025/327 (eur-lex.europa.eu)"),
    (r"AI Act|2024/1689|sztuczn(?:ej|a) inteligencj",
     "AI Act — rozporządzenie (UE) 2024/1689 (eur-lex.europa.eu)"),
    (r"\bRODO\b|\bGDPR\b|dane osobow", "RODO — rozporządzenie (UE) 2016/679"),
    (r"\bNIS2\b|cyberbezpieczen", "Ustawa o krajowym systemie cyberbezpieczeństwa / NIS2 (isap.sejm.gov.pl)"),
    (r"\bURPL\b", "URPL — organ kompetentny (gov.pl/web/urpl)"),
    (r"\bMDCG\b", "Wytyczne MDCG (health.ec.europa.eu)"),
    (r"\bP1\b|\bIKP\b|mojeIKP|e-?[Zz]drowie|Centrum e-Zdrowia|\bCeZ\b",
     "P1 / IKP / EPP — Centrum e-Zdrowia (ezdrowie.gov.pl)"),
    (r"\bEDM\b|HL7|\bCDA\b|\bIHE\b", "HL7 CDA, profile IHE — interoperacyjność EDM (ihe.net, CeZ)"),
    (r"\bRPWDL\b|podmiot leczniczy|dzialalnosc lecznicz",
     "RPWDL 2.0 (rpwdl.ezdrowie.gov.pl); ustawa o działalności leczniczej"),
    (r"\bEUDAMED\b|\bUDI\b", "EUDAMED (ec.europa.eu/tools/eudamed)"),
    (r"ISO 13485|ISO 14971|IEC 62304|IEC 62366|\bnorm[ay]\b",
     "Normy ISO 13485, ISO 14971, IEC 62304, IEC 62366 (pkn.pl, iso.org, iec.ch)"),
    (r"jednostk(?:a|i) notyfikowan", "Wykaz jednostek notyfikowanych — baza NANDO"),
    (r"Dz\.\s?U\.|ustaw(?:a|y|ie) o wyrobach",
     "Ustawa o wyrobach medycznych z 7.04.2022 (Dz.U. 2022 poz. 974, isap.sejm.gov.pl)"),
]

# slowa kluczowe sekcji biznesplanu -> dobor materialu dowodowego i sekcji specyfikacji
KLUCZE_SEKCJI = {
    "streszczenie": ["streszczenie", "zarzadcze", "podsumowanie", "wykonawcze"],
    "problem": ["problem", "fragmentacja", "rozproszenie", "bol"],
    "rozwiazanie": ["rozwiazanie", "warstwa decyzyjna", "agregacja", "jeden system"],
    "teraz": ["dlaczego teraz", "okno", "moment", "timing"],
    "rynek": ["rynek", "tam", "sam", "som", "pole gry", "konkurencja", "ikp", "publiczn"],
    "produkt": ["produkt", "portfel", "modul", "funkcj", "capsule", "station", "pet", "matrix"],
    "model": ["monetyzac", "przychod", "model biznesowy", "freemium", "kto placi", "cennik", "pakiet"],
    "regulacje": ["mdr", "certyfikac", "regulac", "zgodnosc", "klasa", "notyfikowan", "norm"],
    "dowod": ["dowod", "dane", "walidac", "hipotez", "weryfikac", "badani"],
    "wejscie": ["wejscie na rynek", "go to market", "kanal", "sprzedaz", "marketing", "klinik"],
    "konkurencja": ["konkurencj", "gracz", "apple", "google", "whoop", "oura", "porownanie"],
    "fosa": ["fosa", "przewag", "wyroznia", "bariera", "kontrola", "ip", "wlasnosc"],
    "technologia": ["technolog", "architektur", "stos", "komponent", "orkiestrac", "dostawc", "api"],
    "zespol": ["zespol", "struktura", "ludzie", "fundacja", "spolka", "statut", "korporacyjn"],
    "finanse": ["koszt", "finans", "budzet", "naklad", "rentownosc", "cash", "wycena"],
    "finansowanie": ["finansowanie", "grant", "inwestor", "runda", "kapital", "dotacj", "bez funduszy"],
    "ryzyko": ["ryzyk", "zagrozeni", "braki", "luki", "audyt", "ocena", "krytyczn", "sprzecznosc"],
    "kamienie": ["kamien", "milowe", "roadmap", "harmonogram", "sekwencj", "etap", "plan", "90dni", "kolejnosc"],
    "ograniczenia": ["nie obiecuje", "niezrobione", "ograniczeni", "zastrzezeni", "wylaczon"],
    "zrodla": ["zrodl", "odnosnik", "indeks", "rejestr", "bibliograf"],
}


def deacc(s: str) -> str:
    return "".join(c for c in unicodedata.normalize("NFKD", s) if not unicodedata.combining(c))


def dopasuj_temat(tytul: str) -> str:
    t = deacc(tytul).lower()
    najlepszy, wynik = None, 0
    for temat, klucze in KLUCZE_SEKCJI.items():
        s = sum(3 if k in t else 0 for k in klucze)
        if s > wynik:
            najlepszy, wynik = temat, s
    return najlepszy or "produkt"


def main():
    kanon = [json.loads(l) for l in open(KANON, encoding="utf-8")]
    idx = json.load(open(IDX, encoding="utf-8"))
    auth = json.load(open(AUTH, encoding="utf-8"))
    raport = json.load(open(RAPORT, encoding="utf-8"))
    spec = json.load(open(KOTWICE, encoding="utf-8"))
    grupy = idx["grupy"]
    meta = {d["id"]: d for d in idx["dokumenty"]}

    wg_doc = defaultdict(list)
    for b in kanon:
        wg_doc[b["doc"]].append(b)
    for v in wg_doc.values():
        v.sort(key=lambda b: b["n"])

    def tekst_doc(doc_id):
        out = []
        for b in wg_doc[doc_id]:
            out.append(b["txt"])
            if b["uzupelnienia"]:
                out.append("")
                out.append("**Treść unikalna odzyskana z wariantów równoległych:**")
                out.extend(f"- {u['txt']}  *[{u['doc']}]*" for u in b["uzupelnienia"])
            out.append("")
        return "\n".join(out)

    biznes_docs = [d for d in meta if meta[d]["grupa"] in GRUPY_BIZNES and wg_doc.get(d)]
    trzon = max((d for d in biznes_docs if meta[d]["grupa"] == "G3"),
                key=lambda d: auth.get(d, {}).get("punkty", 0.0))

    # --- podzial trzonu na sekcje ---
    # Trzon bierzemy z pliku ZRODLOWEGO, nie z blokow po deduplikacji: biznesplan
    # ma byc dokumentem samodzielnym i kompletnym. Deduplikacja przenosi czesc
    # jego akapitow do dokumentu o wyzszym autorytecie (np. sekcje "Fosa" do
    # specyfikacji), co zostawiloby tu puste sekcje. Duplikacji to nie tworzy,
    # bo Czesc II sklada sie wylacznie z blokow kanonicznych — a te, ktore
    # trzon juz zawiera, sa z niej usuniete.
    surowy = open(os.path.join("out/zrodla", meta[trzon]["plik"]),
                  encoding="utf-8").read()
    surowy = re.sub(r"^<!--.*?-->\s*", "", surowy, flags=re.S)
    surowy = re.sub(r"^#\s+.*\n", "", surowy, count=1)
    elementy = L.parsuj(surowy)
    sekcje, biezaca = [], None
    for e in elementy:
        if e[0] == "h" and e[1] <= 2:
            biezaca = {"tytul": e[2], "elementy": []}
            sekcje.append(biezaca)
        elif biezaca is None:
            biezaca = {"tytul": "Wprowadzenie", "elementy": [e]}
            sekcje.append(biezaca)
        else:
            biezaca["elementy"].append(e)
    sekcje = [s for s in sekcje if s["elementy"]]
    for s in sekcje:
        s["temat"] = dopasuj_temat(s["tytul"])

    # --- przypisanie materialu dowodowego do sekcji ---
    tematy_sekcji = {}
    for s in sekcje:
        tematy_sekcji.setdefault(s["temat"], s["tytul"])

    przypisanie = defaultdict(list)
    for d in biznes_docs:
        if d == trzon or sum(len(b["txt"]) for b in wg_doc[d]) < 400:
            continue
        nazwa = deacc(meta[d]["nazwa"]).lower()
        tresc = deacc(" ".join(b["txt"] for b in wg_doc[d])[:60000]).lower()
        skala = max(len(tresc) / 10000.0, 1.0)
        najlepszy, wynik = None, 0.0
        for temat, klucze in KLUCZE_SEKCJI.items():
            if temat not in tematy_sekcji:
                continue
            # nazwa pliku wazy najmocniej; trafienia w tresci normalizujemy
            # dlugoscia, inaczej najdluzszy dokument wygrywalby kazdy temat
            s = 5.0 * sum(1 for k in klucze if k in nazwa)
            s += sum(tresc.count(k) for k in klucze) / skala
            if s > wynik:
                najlepszy, wynik = temat, s
        przypisanie[najlepszy if wynik >= 2.0 else "_pozostale"].append(d)
    for lista in przypisanie.values():
        lista.sort(key=lambda d: -auth.get(d, {}).get("punkty", 0.0))

    # --- dobor sekcji specyfikacji do linkowania ---
    def link_do_spec(temat, tytul):
        cel = deacc(tytul).lower() + " " + " ".join(KLUCZE_SEKCJI.get(temat, []))
        najlepsza, wynik = None, 0
        for k in spec["kotwice"]:
            t = deacc(k["tytul"]).lower()
            s = sum(2 for w in KLUCZE_SEKCJI.get(temat, []) if w in t)
            s += sum(1 for w in set(deacc(tytul).lower().split()) if len(w) > 5 and w in t)
            if s > wynik:
                najlepsza, wynik = k, s
        return najlepsza if wynik >= 2 else None

    # --- budowa dokumentu ---
    doc = L.nowy_dokument("ETERNAL — Biznesplan")
    L.strona_tytulowa(
        doc, "ETERNAL ECOSYSTEM", "BIZNESPLAN",
        "Dokument scalony z całego korpusu źródłowego\n"
        "z indeksem plików, źródłem pod każdym elementem\n"
        "i odsyłaczami do specyfikacji technicznej",
        [("Trzon dokumentu", f"{meta[trzon]['nazwa']} [{trzon}]"),
         ("Dokumentów w indeksie", f"{len(idx['dokumenty'])}"),
         ("Materiał dowodowy", f"{sum(len(v) for v in przypisanie.values())} dokumentów"),
         ("Dokument powiązany", PLIK_SPEC),
         ("Usunięto duplikatów 1:1", f"{raport['usunieto_1_1']:,}".replace(",", " ")),
         ("Scalono wariantów znaczeniowych", f"{raport['scalonych_wariantow']:,}".replace(",", " ")),
         ("Status", "POUFNE — do użytku wewnętrznego")])

    doc.add_heading("Nota metodyczna i instrukcja czytania", level=1)
    doc.add_paragraph(
        f"Biznesplan powstał ze scalenia {len(idx['dokumenty'])} plików źródłowych. "
        "Trzonem jest dokument o najwyższym autorytecie w grupie biznesowej; "
        "pozostałe pliki wchodzą jako materiał dowodowy, przypisany do tej sekcji "
        "biznesplanu, której dotyczy ich treść.")
    for tyt, opis in [
        ("Źródło pod każdym elementem",
         "Pod każdą sekcją znajduje się linia „Źródło” — plik źródłowy z identyfikatorem, "
         "jego grupa oraz podstawa zewnętrzna (akt prawny, zasób publiczny), jeżeli sekcja "
         "się na nią powołuje. Teza bez wskazanego źródła nie występuje."),
        ("Odsyłacz do specyfikacji",
         f"Sekcje mające odpowiednik techniczny kończy odsyłacz „→ Specyfikacja techniczna”. "
         f"Jest to działające hiperłącze do zakładki w pliku {PLIK_SPEC}; oba pliki muszą leżeć "
         "w tym samym katalogu."),
        ("Indeks plików",
         f"Aneks A zawiera wszystkie {len(idx['dokumenty'])} dokumentów źródłowych: nazwę, grupę, "
         "objętość i stronę, na której materiał z pliku występuje w tym biznesplanie."),
        ("Liczby i luki",
         "Dokument nie uzupełnia luk korpusu. Tam, gdzie źródła podają rozbieżne liczby, "
         "obowiązuje wersja z dokumentu o wyższym autorytecie, a wariant odrzucony jest "
         "zachowany jako treść odzyskana."),
    ]:
        p = doc.add_paragraph(style="List Bullet")
        p.add_run(tyt + ". ").bold = True
        p.add_run(opis)

    doc.add_page_break()
    L.spis_tresci(doc, "1-3")

    # --- CZESC I: trzon z linia zrodla i linkiem do specyfikacji ---
    doc.add_heading("CZĘŚĆ I — BIZNESPLAN", level=1)
    uzyte_na_stronie, zakladki_sekcji = {}, {}
    for nr, s in enumerate(sekcje, 1):
        h = doc.add_heading("", level=2)
        L.wpisz_runy(h, s["tytul"])
        zakladki_sekcji[nr] = L.zakladka(h, f"BP_{nr}")
        L.wstaw_elementy(doc, s["elementy"], przesuniecie=1)

        tekst_sekcji = " ".join(e[-1] if e[0] != "tab" else " ".join(
            " ".join(w) for w in e[1]) for e in s["elementy"])
        podstawy = [op for pat, op in PODSTAWY if re.search(pat, deacc(tekst_sekcji), re.I)]
        wsparcie = przypisanie.get(s["temat"], [])[:4]
        czesci = [f"{meta[trzon]['nazwa']} [{trzon}], grupa {meta[trzon]['grupa']} — "
                  f"{grupy[meta[trzon]['grupa']]}"]
        if wsparcie:
            czesci.append("potwierdzenie w: " + "; ".join(
                f"{meta[d]['nazwa']} [{d}]" for d in wsparcie))
        if podstawy:
            czesci.append("podstawa zewnętrzna: " + "; ".join(podstawy[:4]))
        L.zrodlo(doc, ". ".join(czesci) + ".")

        k = link_do_spec(s["temat"], s["tytul"])
        if k:
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Cm(0.3)
            p.paragraph_format.space_after = Pt(14)
            r = p.add_run("→ Specyfikacja techniczna: ")
            r.font.size = Pt(8.5)
            r.bold = True
            L.hiperlacze(p, k["tytul"], f"{PLIK_SPEC}#{k['zakladka']}")

    # --- CZESC II: material dowodowy ---
    doc.add_page_break()
    doc.add_heading("CZĘŚĆ II — MATERIAŁ DOWODOWY", level=1)
    doc.add_paragraph(
        "Treść pozostałych dokumentów korpusu, po deduplikacji, uporządkowana według sekcji "
        "biznesplanu, której dotyczy. Powtórzenia względem Części I zostały usunięte — poniżej "
        "jest wyłącznie to, co każdy plik wnosi ponad trzon.")
    nr_roz = 0
    rozdzialy = list(tematy_sekcji.items())
    if przypisanie.get("_pozostale"):
        rozdzialy.append(("_pozostale", "materiał przekrojowy, bez jednej sekcji wiodącej"))
    for temat, tytul_sekcji in rozdzialy:
        docs_t = przypisanie.get(temat, [])
        if not docs_t:
            continue
        nr_roz += 1
        doc.add_heading(f"II.{nr_roz} Do sekcji: {tytul_sekcji}", level=2)
        for d in docs_t:
            znakow = sum(len(b["txt"]) for b in wg_doc[d])
            doc.add_heading(meta[d]["nazwa"], level=3)
            uzyte_na_stronie[d] = L.zakladka(doc.paragraphs[-1], f"BPDOC_{d}")
            a = auth.get(d, {})
            L.zrodlo(doc, f"{meta[d]['nazwa']} [{d}] · grupa {meta[d]['grupa']} — "
                          f"{grupy[meta[d]['grupa']]} · pakiet {meta[d]['pakiet']} · "
                          f"autorytet {a.get('punkty', 0):.1f} pkt"
                          + (f" · data {a['data']}" if a.get("data") else "")
                          + f" · wnosi {znakow:,} znaków ponad trzon".replace(",", " "))
            L.wstaw_elementy(doc, L.parsuj(tekst_doc(d)), przesuniecie=3)

    # --- ANEKS A: pelny indeks dokumentow ---
    doc.add_page_break()
    doc.add_heading("ANEKS A — INDEKS WSZYSTKICH DOKUMENTÓW ŹRÓDŁOWYCH", level=1)
    doc.add_paragraph(
        f"Wszystkie {len(idx['dokumenty'])} plików korpusu. Kolumna „Strona” wskazuje miejsce "
        "w tym biznesplanie, gdzie materiał z pliku został użyty; „spec.” oznacza plik "
        f"wykorzystany w dokumencie {PLIK_SPEC}, a „—” plik, którego treść po deduplikacji "
        "w całości pokryły dokumenty nowsze. Numery stron Word wylicza przy aktualizacji pól "
        "(Ctrl+A, następnie F9).")

    t = doc.add_table(rows=1, cols=6)
    t.style = "Light Grid Accent 1"
    for c, tekst in zip(t.rows[0].cells,
                        ["ID", "Nazwa pliku", "Grupa", "Nazwa grupy", "Znaków", "Strona"]):
        c.text = ""
        r = c.paragraphs[0].add_run(tekst)
        r.bold = True
        r.font.size = Pt(8.5)
    for d in sorted(meta, key=lambda d: (meta[d]["grupa"], meta[d]["nazwa"].lower())):
        m = meta[d]
        cells = t.add_row().cells
        for c, tekst in zip(cells[:5], [d, m["nazwa"], m["grupa"], m["grupa_nazwa"],
                                        f"{m['znaki']:,}".replace(",", " ")]):
            c.text = ""
            par = c.paragraphs[0]
            par.paragraph_format.space_after = Pt(1)
            par.add_run(tekst).font.size = Pt(8)
        c = cells[5]
        c.text = ""
        par = c.paragraphs[0]
        par.paragraph_format.space_after = Pt(1)
        if d in uzyte_na_stronie:
            L.pole(par, f" PAGEREF {uzyte_na_stronie[d]} \\h ", "•").font.size = Pt(8)
        elif d == trzon:
            L.pole(par, f" PAGEREF {zakladki_sekcji[1]} \\h ", "•").font.size = Pt(8)
        elif m["grupa"] in ("G1", "G2"):
            par.add_run("spec.").font.size = Pt(8)
        else:
            par.add_run("—").font.size = Pt(8)

    # --- ANEKS B: mapa biznesplan <-> specyfikacja ---
    doc.add_page_break()
    doc.add_heading("ANEKS B — MAPA: BIZNESPLAN ↔ SPECYFIKACJA TECHNICZNA", level=1)
    doc.add_paragraph(
        "Odpowiedniki sekcji w obu dokumentach. Każdy wiersz jest odsyłaczem do "
        f"pliku {PLIK_SPEC}.")
    t = doc.add_table(rows=1, cols=3)
    t.style = "Light Grid Accent 1"
    for c, tekst in zip(t.rows[0].cells, ["Sekcja biznesplanu", "Temat",
                                          "Odpowiednik w specyfikacji"]):
        c.text = ""
        r = c.paragraphs[0].add_run(tekst)
        r.bold = True
        r.font.size = Pt(8.5)
    for nr, s in enumerate(sekcje, 1):
        k = link_do_spec(s["temat"], s["tytul"])
        cells = t.add_row().cells
        for c, tekst in zip(cells[:2], [s["tytul"], s["temat"]]):
            c.text = ""
            c.paragraphs[0].add_run(tekst).font.size = Pt(8)
        c = cells[2]
        c.text = ""
        if k:
            L.hiperlacze(c.paragraphs[0], k["tytul"], f"{PLIK_SPEC}#{k['zakladka']}")
        else:
            c.paragraphs[0].add_run("— brak bezpośredniego odpowiednika").font.size = Pt(8)

    L.stopka_z_numeracja(doc, "ETERNAL — Biznesplan · POUFNE")
    doc.save(WYJSCIE)

    print(f"{WYJSCIE}  ({os.path.getsize(WYJSCIE):,} B)")
    print(f"  trzon: {meta[trzon]['nazwa']} [{trzon}], sekcji: {len(sekcje)}")
    print(f"  material dowodowy: {sum(len(v) for v in przypisanie.values())} dokumentow "
          f"w {nr_roz} rozdzialach")
    print(f"  indeks: {len(meta)} dokumentow")
    linki = sum(1 for s in sekcje if link_do_spec(s["temat"], s["tytul"]))
    print(f"  odsylaczy do specyfikacji: {linki}/{len(sekcje)} sekcji")


if __name__ == "__main__":
    main()
