# -*- coding: utf-8 -*-
"""Silnik skladu DOCX: markdown -> Word (naglowki, tabele, pola TOC/PAGEREF, zakladki, linki)."""
import re
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_SECTION
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

GRANAT = RGBColor(0x14, 0x2A, 0x4C)
SZARY  = RGBColor(0x5A, 0x5A, 0x5A)
AKCENT = RGBColor(0x8A, 0x6D, 0x3B)

# ---------------------------------------------------------------- pola Word
def _pole(par, instr):
    r = par.add_run()
    fc = OxmlElement("w:fldChar"); fc.set(qn("w:fldCharType"), "begin"); r._r.append(fc)
    r = par.add_run()
    it = OxmlElement("w:instrText"); it.set(qn("xml:space"), "preserve"); it.text = instr
    r._r.append(it)
    r = par.add_run()
    fc = OxmlElement("w:fldChar"); fc.set(qn("w:fldCharType"), "separate"); r._r.append(fc)
    r = par.add_run("…")
    r = par.add_run()
    fc = OxmlElement("w:fldChar"); fc.set(qn("w:fldCharType"), "end"); r._r.append(fc)

def spis_tresci(doc, poziomy="1-3"):
    p = doc.add_paragraph()
    _pole(p, f' TOC \\o "{poziomy}" \\h \\z \\u ')
    return p

_BK = [0]
def zakladka(par, nazwa):
    _BK[0] += 1
    s = OxmlElement("w:bookmarkStart"); s.set(qn("w:id"), str(_BK[0])); s.set(qn("w:name"), nazwa)
    e = OxmlElement("w:bookmarkEnd");   e.set(qn("w:id"), str(_BK[0]))
    par._p.insert(0, s); par._p.append(e)

def pageref(par, nazwa):
    _pole(par, f" PAGEREF {nazwa} \\h ")

def numery_stron(sekcja, tekst_lewy=""):
    st = sekcja.footer.paragraphs[0]
    st.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if tekst_lewy:
        r = st.add_run(tekst_lewy + "   ·   "); r.font.size = Pt(8); r.font.color.rgb = SZARY
    r = st.add_run("s. "); r.font.size = Pt(8); r.font.color.rgb = SZARY
    _pole(st, " PAGE ")
    r = st.add_run(" z "); r.font.size = Pt(8); r.font.color.rgb = SZARY
    _pole(st, " NUMPAGES ")
    for r in st.runs:
        r.font.size = Pt(8); r.font.color.rgb = SZARY

def hiperlink(par, tekst, url):
    rid = par.part.relate_to(url,
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
        is_external=True)
    h = OxmlElement("w:hyperlink"); h.set(qn("r:id"), rid)
    r = OxmlElement("w:r"); rpr = OxmlElement("w:rPr")
    c = OxmlElement("w:color"); c.set(qn("w:val"), "1F4E79"); rpr.append(c)
    u = OxmlElement("w:u"); u.set(qn("w:val"), "single"); rpr.append(u)
    r.append(rpr)
    t = OxmlElement("w:t"); t.text = tekst; r.append(t)
    h.append(r); par._p.append(h)

# ---------------------------------------------------------------- style
def przygotuj(doc, tytul_footer=""):
    st = doc.styles["Normal"]
    st.font.name = "Calibri"; st.font.size = Pt(10)
    st.element.rPr.rFonts.set(qn("w:eastAsia"), "Calibri")
    pf = st.paragraph_format
    pf.space_after = Pt(6); pf.line_spacing = 1.10
    for lvl, (rozm, kolor, przed) in {
        1: (18, GRANAT, 20), 2: (14, GRANAT, 16), 3: (11.5, GRANAT, 12),
        4: (10.5, AKCENT, 10), 5: (10, SZARY, 8),
    }.items():
        s = doc.styles[f"Heading {lvl}"]
        s.font.name = "Calibri"; s.font.size = Pt(rozm); s.font.bold = True
        s.font.color.rgb = kolor
        s.paragraph_format.space_before = Pt(przed); s.paragraph_format.space_after = Pt(5)
        s.paragraph_format.keep_with_next = True
    for sec in doc.sections:
        sec.top_margin = sec.bottom_margin = Cm(2.0)
        sec.left_margin = sec.right_margin = Cm(2.0)
        numery_stron(sec, tytul_footer)
    return doc

def strona_tytulowa(doc, nadtytul, tytul, podtytul, meta_wiersze):
    for _ in range(4): doc.add_paragraph()
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(nadtytul); r.font.size = Pt(11); r.font.bold = True
    r.font.color.rgb = AKCENT; r.font.name = "Calibri"
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(tytul); r.font.size = Pt(30); r.font.bold = True; r.font.color.rgb = GRANAT
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(podtytul); r.font.size = Pt(12.5); r.font.color.rgb = SZARY
    doc.add_paragraph()
    t = doc.add_table(rows=0, cols=2); t.alignment = WD_TABLE_ALIGNMENT.CENTER
    t.style = "Table Grid"
    for k, v in meta_wiersze:
        c = t.add_row().cells
        rr = c[0].paragraphs[0].add_run(k); rr.font.bold = True; rr.font.size = Pt(9)
        rr = c[1].paragraphs[0].add_run(v); rr.font.size = Pt(9)
    _obramuj(t, "D8D8D8")
    doc.add_page_break()

def _obramuj(tabela, kolor="BFBFBF"):
    tp = tabela._tbl.tblPr
    b = OxmlElement("w:tblBorders")
    for e in ("top", "left", "bottom", "right", "insideH", "insideV"):
        x = OxmlElement(f"w:{e}")
        x.set(qn("w:val"), "single"); x.set(qn("w:sz"), "4"); x.set(qn("w:color"), kolor)
        b.append(x)
    tp.append(b)

def _cien(kom, kolor):
    sh = OxmlElement("w:shd"); sh.set(qn("w:val"), "clear"); sh.set(qn("w:fill"), kolor)
    kom._tc.get_or_add_tcPr().append(sh)

# ---------------------------------------------------------------- inline md
_INL = re.compile(r"(\*\*.+?\*\*|__.+?__|\*[^*]+?\*|_[^_]+?_|`[^`]+?`)", re.S)
def wstaw_tekst(par, tekst, rozm=10, kolor=None, bold=False, italic=False):
    tekst = re.sub(r"\\([\\`*_{}\[\]()#+\-.!])", r"\1", tekst)
    for cz in _INL.split(tekst):
        if not cz: continue
        b, i, mono = bold, italic, False
        s = cz
        if (s.startswith("**") and s.endswith("**")) or (s.startswith("__") and s.endswith("__")):
            b, s = True, s[2:-2]
        elif (s.startswith("*") and s.endswith("*") and len(s) > 2) or \
             (s.startswith("_") and s.endswith("_") and len(s) > 2):
            i, s = True, s[1:-1]
        elif s.startswith("`") and s.endswith("`") and len(s) > 2:
            mono, s = True, s[1:-1]
        if not s: continue
        r = par.add_run(s)
        r.font.size = Pt(rozm); r.font.bold = b; r.font.italic = i
        if mono: r.font.name = "Consolas"
        if kolor is not None: r.font.color.rgb = kolor
    return par

# ---------------------------------------------------------------- tabele md
def _wiersz(ln):
    s = ln.strip()
    if s.startswith("|"): s = s[1:]
    if s.endswith("|"):   s = s[:-1]
    return [c.strip() for c in s.split("|")]

def wstaw_tabele(doc, linie, maks_kol=9):
    wiersze = [_wiersz(l) for l in linie if l.strip()]
    wiersze = [w for w in wiersze if not all(re.fullmatch(r":?-{2,}:?", c or "-") for c in w)]
    if not wiersze: return None
    ncol = min(max(len(w) for w in wiersze), maks_kol)
    wiersze = [(w + [""] * ncol)[:ncol] for w in wiersze]
    # pomin tabele calkowicie puste
    if not any(any(c for c in w) for w in wiersze[1:] if len(wiersze) > 1):
        if len(wiersze) < 2: return None
    t = doc.add_table(rows=0, cols=ncol)
    t.style = "Table Grid"; t.autofit = True
    for idx, w in enumerate(wiersze):
        kom = t.add_row().cells
        for j, v in enumerate(w):
            p = kom[j].paragraphs[0]; p.paragraph_format.space_after = Pt(1)
            wstaw_tekst(p, v, rozm=8.5, bold=(idx == 0),
                        kolor=GRANAT if idx == 0 else None)
            if idx == 0: _cien(kom[j], "EDF0F5")
    _obramuj(t)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)
    return t
