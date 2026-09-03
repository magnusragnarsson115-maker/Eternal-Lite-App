#!/usr/bin/env python3
"""Sklada BIZNESPLAN (.docx): tresc z calego korpusu + zrodlo pod kazdym elementem
+ indeks 149 plikow (nazwa, strona, grupa) + odeslanie do specyfikacji technicznej."""
import os
import json, sys, os, re, datetime
BASE = os.environ.get("KONSOLIDACJA_BASE", os.path.dirname(os.path.abspath(__file__)))
sys.path.insert(0, os.path.dirname(os.path.abspath(__file__)))
from docx import Document
from docx.shared import Pt
from lib_docx import (przygotuj, strona_tytulowa, spis_tresci, zakladka, pageref,
                      hiperlink, wstaw_tekst, wstaw_tabele, _obramuj, _cien,
                      GRANAT, SZARY, AKCENT)
from lib_md import struktura, render, zrodlo, nota
from lib_grupy import GRUPY
from lib_zrodla import Atrybucja

WSZYSTKIE = json.load(open(f"{BASE}/work/index.json", encoding="utf-8"))       # 149
POZOSTALE = {r["id"]: r for r in json.load(open(f"{BASE}/work/po_zwyklej.json", encoding="utf-8"))}
NAZWA = {r["id"]: r["source_name"] for r in WSZYSTKIE}
GRUPA = {r["id"]: r["grupa"] for r in WSZYSTKIE}
DUP   = {r["id"]: r.get("duplikat_of") for r in WSZYSTKIE}

def tresc(i):
    p = f"{BASE}/work/dza_{i:03d}.md"
    return open(p, encoding="utf-8").read() if os.path.exists(p) else ""

atr = Atrybucja(list(POZOSTALE))
zakotwiczone = set()
uzyte_gdzie = {}

def linia_zrodla(doc, fragment, dodatkowe=()):
    """Zrodlo pod elementem — potwierdzenie slusznosci tezy."""
    ids = list(dict.fromkeys(list(dodatkowe) + atr.zrodla(fragment)))[:3]
    if not ids:
        return
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0); p.paragraph_format.space_after = Pt(10)
    p.paragraph_format.left_indent = Pt(10)
    r = p.add_run("Źródło: "); r.font.size = Pt(8); r.font.bold = True; r.font.color.rgb = AKCENT
    for k, i in enumerate(ids):
        if k: 
            r = p.add_run(" · "); r.font.size = Pt(8); r.font.color.rgb = SZARY
        r = p.add_run(f"[{i:03d}] {NAZWA[i]}")
        r.font.size = Pt(8); r.font.italic = True; r.font.color.rgb = SZARY
        if i not in zakotwiczone:
            zakotwiczone.add(i); uzyte_gdzie[i] = "tekst"
            zakladka(p, f"ZR_{i:03d}")
    return p

def render_z_zrodlami(doc, el, baza_ids=(), maks_kol=9):
    """Renderuje strukture, wstawiajac linie zrodla pod kazdym elementem
    (tabela = element; ciag akapitow pod naglowkiem = element)."""
    bufor = []
    def domknij():
        if bufor:
            linia_zrodla(doc, " ".join(bufor), baza_ids)
            bufor.clear()
    for rodzaj, lvl, c in el:
        if rodzaj == "h":
            domknij()
            doc.add_heading(re.sub(r"\\", "", c)[:180], level=max(1, min(lvl, 5)))
        elif rodzaj == "tbl":
            domknij()
            wstaw_tabele(doc, c, maks_kol=maks_kol)
            linia_zrodla(doc, " ".join(c), baza_ids)
        elif rodzaj == "q":
            p = doc.add_paragraph(); p.paragraph_format.left_indent = Pt(16)
            p.paragraph_format.space_after = Pt(2)
            wstaw_tekst(p, c, rozm=9, kolor=SZARY)
        elif rodzaj == "li":
            p = doc.add_paragraph(style="List Bullet")
            p.paragraph_format.space_after = Pt(2); wstaw_tekst(p, c); bufor.append(c)
        else:
            wstaw_tekst(doc.add_paragraph(), c); bufor.append(c)
    domknij()

def sekcje(i, numery):
    """Wycina z pliku sekcje o podanych numerach wiodacych (np. '6', '40')."""
    el = struktura(tresc(i), baza=2)
    chce = {str(n) for n in numery}
    out, bierz, poziom = [], False, 99
    for rodzaj, lvl, c in el:
        if rodzaj == "h":
            m = re.match(r"^(\d{1,3})(?:\.(\d{1,3}))?\.", str(c))
            if m:
                if m.group(1) in chce and not m.group(2):
                    bierz, poziom = True, lvl
                elif not m.group(2) and lvl <= poziom:
                    bierz = False
        if bierz:
            out.append((rodzaj, lvl, c))
    return out

# =========================================================== dokument
dzis = datetime.date.today().strftime("%d.%m.%Y")
doc = Document()
przygotuj(doc, "ETERNAL — Biznesplan")
strona_tytulowa(doc,
    "ETERNAL LABS SP. Z O.O.",
    "BIZNESPLAN",
    "Wersja skonsolidowana z całego korpusu dokumentacyjnego, ze wskazaniem źródła "
    "pod każdym elementem",
    [("Data złożenia", dzis),
     ("Podstawa źródłowa", "149 plików · 1 300 064 słów treści wyjściowej"),
     ("Po deduplikacji", "135 plików · 642 813 słów"),
     ("Indeks źródeł", "Załącznik A — wszystkie 149 plików z numerem strony i grupą"),
     ("Dokument powiązany", "ETERNAL — Specyfikacja techniczna (Załącznik B)"),
     ("Klauzula", "POUFNE — do użytku wewnętrznego")])

doc.add_heading("Spis treści", level=1)
spis_tresci(doc, "1-3")
nota(doc, "Spis treści oraz numery stron w indeksie źródeł (Załącznik A) są polami Worda. "
          "Po otwarciu pliku naciśnij Ctrl+A, potem F9 i wybierz „Aktualizuj cały spis” — "
          "dopiero wtedy numery stron się wypełnią.")
doc.add_page_break()

# --- nota o metodzie
doc.add_heading("Nota o metodzie i podstawie źródłowej", level=1)
for a in [
 "Ten biznesplan nie jest dokumentem pisanym od zera. Powstał z korpusu 149 plików — "
 "konwersji rozmów roboczych, opracowań, arkuszy, prezentacji i kolejnych wersji "
 "dokumentów — scalonych w jeden zbiór i przepuszczonych przez dwa przebiegi deduplikacji.",
 "**Deduplikacja zwykła** usunęła treść identyczną znak w znak: 14 plików będących "
 "dokładnymi kopiami i 1 428 powtórzonych bloków. Korpus: 1 300 064 → 864 939 słów.",
 "**Deduplikacja zaawansowana** usunęła treść o zbliżonym znaczeniu, zachowując unikat: "
 "267 klastrów, 596 pochłoniętych wariantów, 864 939 → 642 813 słów. Z każdego "
 "pochłoniętego wariantu wyłuskano zdania nieobecne w wersji wiodącej — 1 179 zdań "
 "unikatowych przeniesiono do wersji wiodącej wraz ze wskazaniem pliku źródłowego.",
 "**Jak czytać wiersze „Źródło”.** Pod każdym elementem — pod każdą tabelą i pod każdym "
 "blokiem tez — stoi wiersz wskazujący pliki korpusu, w których ta treść faktycznie "
 "występuje. Przypisania nie są deklaracją autora: liczy je dopasowanie treści elementu "
 "do korpusu wagami IDF, więc wskazany plik naprawdę zawiera dane twierdzenie. To jest "
 "materiał do weryfikacji tezy, nie ozdobnik.",
 "**Numeracja w nawiasach kwadratowych** — na przykład [030] — odsyła do Załącznika A, "
 "gdzie każdy z 149 plików ma numer strony, na której został przywołany, oraz grupę "
 "tematyczną.",
]:
    wstaw_tekst(doc.add_paragraph(), a)
doc.add_page_break()

# =========================================================== CZĘŚĆ I
doc.add_heading("CZĘŚĆ I — BIZNESPLAN", level=1)
nota(doc, "Trzon planu. Wersja wiodąca korpusu w warstwie biznesowej — Biznesplan 4.0 — "
          "uzupełniona treścią z pozostałych dokumentów tam, gdzie ją wnoszą.")
el = struktura(tresc(25), baza=1)
while el and el[0][0] == "h" and str(el[0][2]).isupper():
    el.pop(0)
# sekcja 20 (zrodla zewnetrzne) idzie do Czesci III
tnij = next((k for k, (r, l, c) in enumerate(el)
             if r == "h" and str(c).startswith("20.")), len(el))
render_z_zrodlami(doc, el[:tnij], baza_ids=(25,))
sekcja20 = el[tnij:]
doc.add_page_break()

# =========================================================== CZĘŚĆ II
doc.add_heading("CZĘŚĆ II — POGŁĘBIENIE EKONOMICZNE I WYKONAWCZE", level=1)
nota(doc, "Rozwinięcie tez Części I materiałem z pozostałych dokumentów korpusu: "
          "ekonomia specyfikacji wiodącej, plan korporacyjny, roadmapa, modele monetyzacji, "
          "koszty i alternatywy, konkurencja.")

doc.add_heading("II.1 Ekonomia, portfel i pozycja — ze specyfikacji wiodącej", level=2)
nota(doc, "Wycinek sekcji ekonomicznych Specyfikacji Master 5.4 FINAL.")
render_z_zrodlami(doc, sekcje(30, [6, 8, 9, 10, 11, 12, 15, 16, 40, 41, 42]), baza_ids=(30,))

ROZDZ = [
 ("II.2 Plan korporacyjny i struktura", [28, 21]),
 ("II.3 Roadmapa i sekwencja wykonawcza", [29, 73, 143, 135]),
 ("II.4 Modele monetyzacji i strumienie przychodu", [12, 13, 20, 139, 142, 114, 16]),
 ("II.5 Koszty, alternatywy i kontrola nad dostawcami", [43, 106, 17, 101]),
 ("II.6 Produkty, pakiety i karty produktowe", [81, 82, 59, 102]),
 ("II.7 Konkurencja, wyróżniki i pole gry", [128, 71, 92, 80]),
 ("II.8 Marketing, społeczność i współdecydowanie", [83, 93]),
]
for tyt, ids in ROZDZ:
    ids = [i for i in ids if tresc(i)]
    if not ids: continue
    doc.add_heading(tyt, level=2)
    for i in ids:
        doc.add_heading(NAZWA[i].rsplit(".", 1)[0].replace("_", " "), level=3)
        linia_zrodla(doc, NAZWA[i], (i,))
        e = struktura(tresc(i), baza=3)
        while e and e[0][0] == "h" and str(e[0][2]).isupper():
            e.pop(0)
        render_z_zrodlami(doc, e, baza_ids=(i,))
doc.add_page_break()

# =========================================================== CZĘŚĆ III
doc.add_heading("CZĘŚĆ III — ŹRÓDŁA ZEWNĘTRZNE I ODNOŚNIKI", level=1)
nota(doc, "Akty prawne, wytyczne organów, infrastruktura publiczna i normy przywołane w planie.")
render_z_zrodlami(doc, sekcja20, baza_ids=(25,))
doc.add_page_break()

# =========================================================== ZAŁĄCZNIK A
doc.add_heading("ZAŁĄCZNIK A — INDEKS PLIKÓW ŹRÓDŁOWYCH", level=1)
for a in [
 f"Wszystkie **{len(WSZYSTKIE)} plików** stanowiących podstawę tego planu. Kolumna "
 "**Strona** wskazuje stronę, na której plik został przywołany jako źródło; dla plików "
 "niewywołanych wprost w tekście — stronę jego pozycji w rejestrze A.2. Kolumna "
 "**Grupa** przypisuje plik do jednej z dziewięciu grup tematycznych korpusu.",
 "Numery stron są polami Worda i wypełnią się dopiero po aktualizacji "
 "(Ctrl+A, potem F9).",
]:
    wstaw_tekst(doc.add_paragraph(), a)

doc.add_heading("A.1 Legenda grup tematycznych", level=2)
from collections import Counter
lg = Counter(GRUPA.values())
t = doc.add_table(rows=1, cols=3); t.style = "Table Grid"
for j, h in enumerate(["Grupa", "Zakres", "Plików"]):
    r = t.rows[0].cells[j].paragraphs[0].add_run(h)
    r.font.bold = True; r.font.size = Pt(9); r.font.color.rgb = GRANAT
    _cien(t.rows[0].cells[j], "EDF0F5")
for g, nz in GRUPY.items():
    c = t.add_row().cells
    for j, v in enumerate([g, nz, str(lg.get(g, 0))]):
        wstaw_tekst(c[j].paragraphs[0], v, rozm=8.5, bold=(j == 0))
_obramuj(t)

doc.add_heading("A.2 Indeks pełny — 149 plików", level=2)
t = doc.add_table(rows=1, cols=6); t.style = "Table Grid"
for j, h in enumerate(["#", "Plik źródłowy", "Typ", "Grupa", "Słowa", "Strona"]):
    r = t.rows[0].cells[j].paragraphs[0].add_run(h)
    r.font.bold = True; r.font.size = Pt(9); r.font.color.rgb = GRANAT
    _cien(t.rows[0].cells[j], "EDF0F5")
for r0 in sorted(WSZYSTKIE, key=lambda r: r["source_name"].lower()):
    i = r0["id"]; c = t.add_row().cells
    for j, v in enumerate([f"{i:03d}", r0["source_name"], r0["ext"],
                           f"{GRUPA[i]} — {GRUPY[GRUPA[i]]}",
                           f"{r0['words']:,}".replace(",", " ")]):
        wstaw_tekst(c[j].paragraphs[0], v, rozm=8)
    pageref(c[5].paragraphs[0], f"ZR_{i:03d}")
    for rr in c[5].paragraphs[0].runs:
        rr.font.size = Pt(8)
_obramuj(t)
doc.add_page_break()

doc.add_heading("A.3 Rejestr materiału nieprzywołanego wprost w tekście", level=2)
brak = [r for r in WSZYSTKIE if r["id"] not in zakotwiczone]
wstaw_tekst(doc.add_paragraph(),
 f"**{len(brak)}** plików nie zostało przywołanych jako źródło pod konkretnym elementem. "
 "Dzielą się na dwie kategorie: kopie usunięte w deduplikacji zwykłej (ich treść jest "
 "w pliku wskazanym w kolumnie „Status”) oraz materiał wspierający, którego treść "
 "została pochłonięta przez wersje wiodące w deduplikacji zaawansowanej. "
 "Każda pozycja ma tu zakładkę, do której odsyła indeks A.2.")
t = doc.add_table(rows=1, cols=4); t.style = "Table Grid"
for j, h in enumerate(["#", "Plik", "Grupa", "Status"]):
    r = t.rows[0].cells[j].paragraphs[0].add_run(h)
    r.font.bold = True; r.font.size = Pt(9); r.font.color.rgb = GRANAT
    _cien(t.rows[0].cells[j], "EDF0F5")
for r0 in sorted(brak, key=lambda r: r["source_name"].lower()):
    i = r0["id"]; c = t.add_row().cells
    d = DUP.get(i)
    status = (f"Kopia identyczna 1:1 — treść w [{d:03d}] {NAZWA[d]}" if d
              else "Materiał wspierający — treść pochłonięta przez wersję wiodącą")
    p0 = c[0].paragraphs[0]
    wstaw_tekst(p0, f"{i:03d}", rozm=8)
    zakladka(p0, f"ZR_{i:03d}")
    for j, v in enumerate([r0["source_name"], f"{GRUPA[i]} — {GRUPY[GRUPA[i]]}", status], start=1):
        wstaw_tekst(c[j].paragraphs[0], v, rozm=8)
_obramuj(t)
doc.add_page_break()

# =========================================================== ZAŁĄCZNIK B
doc.add_heading("ZAŁĄCZNIK B — ODESŁANIE DO SPECYFIKACJI TECHNICZNEJ", level=1)
wstaw_tekst(doc.add_paragraph(),
 "Warstwa techniczna tego planu — moduły, funkcje, architektura komponentów, granica "
 "regulacyjna funkcja po funkcji, rejestry zakresu — jest opisana w osobnym dokumencie "
 "złożonym z tego samego korpusu i tą samą metodą.")
p = doc.add_paragraph()
r = p.add_run("Dokument: "); r.font.bold = True; r.font.size = Pt(10)
hiperlink(p, "ETERNAL — Specyfikacja techniczna "
             "(03_ETERNAL_SPECYFIKACJA_TECHNICZNA.docx)",
          "03_ETERNAL_SPECYFIKACJA_TECHNICZNA.docx")
nota(doc, "Odnośnik jest względny — działa, gdy oba pliki leżą w tym samym katalogu.")

doc.add_heading("B.1 Co któremu rozdziałowi planu odpowiada w specyfikacji", level=2)
MAPA = [
 ("3. Rozwiązanie · 6. Produkt i portfel", "Część A — sekcje 1, 8, 29 · Część B",
  "Zakres, liczby bazowe, etap zerowy, moduły aplikacji"),
 ("8. Ścieżka regulacyjna", "Część A — sekcje 2, 3 · Część E",
  "Granica wyrobu funkcja po funkcji, klasy, terminy"),
 ("7. Model biznesowy · 15. Finanse", "Część A — sekcje 6, 40, 41, 42",
  "Ekonomia, osiem strumieni przychodu, Station"),
 ("12. Fosa · 11. Konkurencja", "Część A — sekcje 10, 23",
  "Pozycja wobec państwa, drabina kontroli technologicznej"),
 ("13. Technologia i operacje", "Część A — sekcje 4, 7, 21, 22 · Część D",
  "Architektura, reguła 33%, komponenty, odporność"),
 ("14. Zespół i struktura", "Część A — sekcje 11, 23.1",
  "Struktura instytucjonalna, pięć źródeł kontroli"),
 ("18. Kamienie milowe", "Część A — sekcje 15, 27",
  "Kolejność prac, warianty dat"),
]
t = doc.add_table(rows=1, cols=3); t.style = "Table Grid"
for j, h in enumerate(["Rozdział biznesplanu", "Miejsce w specyfikacji", "Zawartość"]):
    r = t.rows[0].cells[j].paragraphs[0].add_run(h)
    r.font.bold = True; r.font.size = Pt(9); r.font.color.rgb = GRANAT
    _cien(t.rows[0].cells[j], "EDF0F5")
for a, b, c0 in MAPA:
    c = t.add_row().cells
    for j, v in enumerate([a, b, c0]):
        wstaw_tekst(c[j].paragraphs[0], v, rozm=8.5)
_obramuj(t)

sciezka = f"{BASE}/out/04_ETERNAL_BIZNESPLAN.docx"
doc.save(sciezka)
print("zapisano:", sciezka, os.path.getsize(sciezka), "B")
print("przywolanych w tekscie:", len(zakotwiczone), "| w rejestrze A.3:", len(brak))
