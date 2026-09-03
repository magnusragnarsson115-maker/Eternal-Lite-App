# -*- coding: utf-8 -*-
"""Odzyskanie struktury naglowkow z konwersji docx + render markdown -> Word."""
import os
import re, sys
BASE = os.environ.get("KONSOLIDACJA_BASE", os.path.dirname(os.path.abspath(__file__)))
sys.path.insert(0, os.path.dirname(os.path.abspath(__file__)))
from docx.shared import Pt, RGBColor
from lib_docx import wstaw_tekst, wstaw_tabele, SZARY, AKCENT, GRANAT

# konwersja docx zgubila poziomy — sekcje zostaly pogrubieniami z numeracja
_CZESC = re.compile(r"^\*\*\s*(CZ[ĘE]ŚĆ|CZESC|ZAŁĄCZNIK|ZALACZNIK|ANEKS)\b[^*]{0,120}\*\*\s*$")
_N1    = re.compile(r"^\*\*\s*(\d{1,3})\\?\.\s+([^*]{2,110}?)\s*\*\*\s*$")
_N2    = re.compile(r"^\*\*\s*(\d{1,3}\.\d{1,3})\\?\.?\s+([^*]{2,110}?)\s*\*\*\s*$")
_N3    = re.compile(r"^\*\*\s*(\d{1,3}\.\d{1,3}\.\d{1,3})\\?\.?\s+([^*]{2,110}?)\s*\*\*\s*$")
_CAPS  = re.compile(r"^\*\*\s*([A-ZĄĆĘŁŃÓŚŹŻ0-9][A-ZĄĆĘŁŃÓŚŹŻ0-9 ,\-–—/()\.]{4,80})\s*\*\*\s*$")
_MDH   = re.compile(r"^(#{1,6})\s+(.*\S)\s*$")
_LI    = re.compile(r"^\s*([-*+]|\d{1,3}[.)])\s+(.*)$")
_TBL   = re.compile(r"^\s*\|")

def struktura(tekst, baza=1):
    """Zwraca liste (rodzaj, poziom, tresc): h/p/li/tbl."""
    linie = tekst.split("\n")
    out, i = [], 0
    while i < len(linie):
        ln = linie[i]
        s = ln.strip()
        if not s:
            i += 1; continue
        if _TBL.match(s):
            blok = []
            while i < len(linie) and _TBL.match(linie[i].strip() or "x"):
                if not linie[i].strip(): break
                blok.append(linie[i]); i += 1
            out.append(("tbl", 0, blok)); continue
        if s.startswith(">"):
            out.append(("q", 0, s.lstrip("> ").strip())); i += 1; continue
        m = _MDH.match(s)
        if m:
            out.append(("h", min(baza + len(m.group(1)) - 1, 5), m.group(2))); i += 1; continue
        if _CZESC.match(s):
            out.append(("h", baza, s.strip("*").strip())); i += 1; continue
        for rx, lvl in ((_N3, baza + 3), (_N2, baza + 2), (_N1, baza + 1)):
            m = rx.match(s)
            if m:
                out.append(("h", min(lvl, 5), f"{m.group(1)}. {m.group(2)}")); break
        else:
            m = _CAPS.match(s)
            if m and len(m.group(1)) <= 80 and not m.group(1).endswith("."):
                out.append(("h", min(baza + 2, 5), m.group(1).strip())); i += 1; continue
            m = _LI.match(s)
            if m:
                out.append(("li", 0, m.group(2))); i += 1; continue
            if set(s) <= set("-–—_* "):
                i += 1; continue
            out.append(("p", 0, s)); i += 1; continue
        i += 1
    return out

def render(doc, elementy, maks_kol=9):
    """Wypisuje strukture do dokumentu. Zwraca liczbe wstawionych naglowkow."""
    n = 0
    for rodzaj, lvl, tresc in elementy:
        if rodzaj == "h":
            doc.add_heading(re.sub(r"\\", "", tresc)[:180], level=max(1, min(lvl, 5)))
            n += 1
        elif rodzaj == "tbl":
            wstaw_tabele(doc, tresc, maks_kol=maks_kol)
        elif rodzaj == "q":
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Pt(16)
            p.paragraph_format.space_after = Pt(2)
            wstaw_tekst(p, tresc, rozm=9, kolor=SZARY)
        elif rodzaj == "li":
            p = doc.add_paragraph(style="List Bullet")
            p.paragraph_format.space_after = Pt(2)
            wstaw_tekst(p, tresc)
        else:
            p = doc.add_paragraph()
            wstaw_tekst(p, tresc)
    return n

def zrodlo(doc, tekst):
    """Linia zrodla pod elementem — potwierdzenie slusznosci tezy."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(1); p.paragraph_format.space_after = Pt(9)
    p.paragraph_format.left_indent = Pt(10)
    r = p.add_run("Źródło: "); r.font.size = Pt(8); r.font.bold = True; r.font.color.rgb = AKCENT
    r = p.add_run(tekst); r.font.size = Pt(8); r.font.italic = True; r.font.color.rgb = SZARY
    return p

def nota(doc, tekst):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Pt(12); p.paragraph_format.space_after = Pt(8)
    wstaw_tekst(p, tekst, rozm=9, kolor=SZARY, italic=True)
    return p
