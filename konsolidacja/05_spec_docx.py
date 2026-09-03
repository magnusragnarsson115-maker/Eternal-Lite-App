#!/usr/bin/env python3
"""Skalda SPECYFIKACJE TECHNICZNA (.docx) z calej tresci po deduplikacji."""
import os
import json, sys, os, datetime
BASE = os.environ.get("KONSOLIDACJA_BASE", os.path.dirname(os.path.abspath(__file__)))
sys.path.insert(0, os.path.dirname(os.path.abspath(__file__)))
from docx import Document
from docx.shared import Pt
from lib_docx import (przygotuj, strona_tytulowa, spis_tresci, zakladka,
                      wstaw_tekst, wstaw_tabele, GRANAT, SZARY, AKCENT)
from lib_md import struktura, render, zrodlo, nota
from lib_grupy import GRUPY

rec = {r["id"]: r for r in json.load(open(f"{BASE}/work/po_zwyklej.json", encoding="utf-8"))}
def tresc(i):
    p = f"{BASE}/work/dza_{i:03d}.md"
    return open(p, encoding="utf-8").read() if os.path.exists(p) else ""
def slowa(i): return len(tresc(i).split())

# ---- kompozycja dokumentu -------------------------------------------------
CZESCI = [
 ("A", "SPECYFIKACJA MASTER — WERSJA WIODĄCA",
  "Rdzeń specyfikacji. Wersja 5.4 FINAL pochłonęła w deduplikacji zaawansowanej "
  "treść wersji 3.0 i 3.1; fragmenty unikatowe ze starszych wersji są oznaczone "
  "w tekście jako uzupełnienia scalone.", [30]),
 ("B", "SPECYFIKACJA APLIKACJI I DOKUMENTACJA POWIĄZANA",
  "Warstwa aplikacyjna — moduły A1–A24 i dokumentacja opisująca ich działanie.",
  [24, 48, 49, 107, 9, 5, 62, 63]),
 ("C", "REJESTRY FUNKCJI, MODUŁÓW I TAKSONOMIA",
  "Obowiązujące rejestry zakresu: liczby funkcji, moduły, taksonomia kodów.", "G2"),
 ("D", "ARCHITEKTURA I TECHNOLOGIA",
  "Komponenty, warstwy, agregacja danych, sprzęt, dostawcy i interfejsy.", "G5"),
 ("E", "ZGODNOŚĆ REGULACYJNA, CERTYFIKACJA I OTOCZENIE PUBLICZNE",
  "Granica wyrobu medycznego, ścieżki certyfikacji, relacja z infrastrukturą "
  "państwową (P1, IKP, EHDS).", "G4"),
]
def pliki_czesci(spec):
    if isinstance(spec, list):
        return [i for i in spec if i in rec and tresc(i)]
    ids = [r["id"] for r in rec.values() if r["grupa"] == spec and tresc(r["id"])]
    return sorted(ids, key=lambda i: -rec[i]["prio"])

uzyte, plan = set(), []
for kod, tyt, opis, spec in CZESCI:
    ids = [i for i in pliki_czesci(spec) if i not in uzyte]
    uzyte.update(ids)
    plan.append((kod, tyt, opis, ids))

dzis = datetime.date.today().strftime("%d.%m.%Y")
laczne = sum(slowa(i) for _, _, _, ids in plan for i in ids)

doc = Document()
przygotuj(doc, "ETERNAL — Specyfikacja techniczna")
strona_tytulowa(doc,
    "ETERNAL LABS SP. Z O.O.",
    "SPECYFIKACJA TECHNICZNA",
    "Konsolidacja korpusu dokumentacyjnego po deduplikacji zwykłej i zaawansowanej",
    [("Data złożenia", dzis),
     ("Podstawa", "149 plików źródłowych z 7 paczek konwersji"),
     ("Po deduplikacji", "135 plików · 642 813 słów"),
     ("Zakres dokumentu", f"{len(uzyte)} plików technicznych · ok. {laczne:,} słów".replace(",", " ")),
     ("Wersja wiodąca", "ETERNAL_Specyfikacja_Master_5_4_FINAL"),
     ("Klauzula", "POUFNE — do użytku wewnętrznego")])

doc.add_heading("Spis treści", level=1)
spis_tresci(doc, "1-3")
nota(doc, "Spis treści jest polem Worda. Po otwarciu pliku naciśnij Ctrl+A, potem F9 "
          "i wybierz „Aktualizuj cały spis”, aby wypełnić numery stron.")
doc.add_page_break()

# ---- nota o pochodzeniu ---------------------------------------------------
doc.add_heading("Nota o pochodzeniu treści i metodzie konsolidacji", level=1)
for a in [
 "Dokument powstał z 149 plików źródłowych — konwersji rozmów, opracowań, arkuszy, "
 "prezentacji i dokumentów roboczych — scalonych w jeden korpus, a następnie "
 "przepuszczonych przez dwa przebiegi deduplikacji.",
 "**Deduplikacja zwykła** usunęła treść identyczną znak w znak: 14 plików będących "
 "dokładnymi kopiami oraz 1 428 powtórzonych bloków. Korpus zszedł z 1 300 064 do "
 "864 939 słów, bez ingerencji w brzmienie czegokolwiek.",
 "**Deduplikacja zaawansowana** usunęła treść o zbliżonym znaczeniu i kontekście. "
 "267 klastrów bliskoznacznych, 596 wariantów pochłoniętych, korpus 864 939 → 642 813 słów. "
 "Kluczowa reguła: z każdego pochłoniętego wariantu wyłuskano zdania, których nie było "
 "w wersji wiodącej — **1 179 zdań unikatowych** doklejono do wersji wiodącej wraz ze "
 "wskazaniem pliku źródłowego. Nic nie zostało utracone; treść została przeniesiona.",
 "Wersję wiodącą wybierał priorytet liczony z numeru wersji w nazwie pliku, znaczników "
 "FINAL / KOMPLETNA / scalona oraz wagi dokumentu. Dlatego rdzeniem specyfikacji jest "
 "Master 5.4 FINAL, a nie wcześniejsze 3.0 i 3.1.",
]:
    wstaw_tekst(doc.add_paragraph(), a)

doc.add_heading("Skład dokumentu", level=2)
t = doc.add_table(rows=1, cols=4); t.style = "Table Grid"
for j, h in enumerate(["Część", "Zakres", "Plików", "Słów"]):
    r = t.rows[0].cells[j].paragraphs[0].add_run(h)
    r.font.bold = True; r.font.size = Pt(9); r.font.color.rgb = GRANAT
for kod, tyt, opis, ids in plan:
    c = t.add_row().cells
    wstaw_tekst(c[0].paragraphs[0], f"**{kod}**", rozm=8.5)
    wstaw_tekst(c[1].paragraphs[0], tyt, rozm=8.5)
    wstaw_tekst(c[2].paragraphs[0], str(len(ids)), rozm=8.5)
    wstaw_tekst(c[3].paragraphs[0], f"{sum(slowa(i) for i in ids):,}".replace(",", " "), rozm=8.5)
from lib_docx import _obramuj; _obramuj(t)
doc.add_page_break()

# ---- czesci ---------------------------------------------------------------
zrodla_uzyte = []
for kod, tyt, opis, ids in plan:
    doc.add_heading(f"CZĘŚĆ {kod} — {tyt}", level=1)
    nota(doc, opis)
    for i in ids:
        r = rec[i]
        h = doc.add_heading(r["source_name"].rsplit(".", 1)[0].replace("_", " "), level=2)
        zakladka(h, f"SPEC_{i:03d}")
        zrodlo(doc, f"[{i:03d}] {r['source_name']} · grupa {r['grupa']} — "
                    f"{GRUPY[r['grupa']]} · {slowa(i):,} słów po deduplikacji".replace(",", " "))
        zrodla_uzyte.append(i)
        el = struktura(tresc(i), baza=2 if i == 30 else 3)
        # usun powtorzony blok tytulowy pliku na starcie
        while el and el[0][0] == "h" and el[0][1] >= 3 and el[0][2].isupper():
            el.pop(0)
        render(doc, el)
    doc.add_page_break()

# ---- zalacznik: indeks zrodel --------------------------------------------
doc.add_heading("Załącznik — indeks plików źródłowych specyfikacji", level=1)
nota(doc, "Pełny indeks wszystkich 149 plików korpusu znajduje się w biznesplanie.")
t = doc.add_table(rows=1, cols=5); t.style = "Table Grid"
for j, h in enumerate(["#", "Plik źródłowy", "Typ", "Grupa", "Słowa"]):
    r = t.rows[0].cells[j].paragraphs[0].add_run(h)
    r.font.bold = True; r.font.size = Pt(9); r.font.color.rgb = GRANAT
for i in sorted(zrodla_uzyte):
    r = rec[i]; c = t.add_row().cells
    for j, v in enumerate([f"{i:03d}", r["source_name"], r["ext"],
                           f"{r['grupa']} {GRUPY[r['grupa']]}",
                           f"{slowa(i):,}".replace(",", " ")]):
        wstaw_tekst(c[j].paragraphs[0], v, rozm=8)
_obramuj(t)

os.makedirs(f"{BASE}/out", exist_ok=True)
sciezka = f"{BASE}/out/03_ETERNAL_SPECYFIKACJA_TECHNICZNA.docx"
doc.save(sciezka)
print("zapisano:", sciezka, os.path.getsize(sciezka), "B")
print("plikow w specyfikacji:", len(zrodla_uzyte), "| slow:", f"{laczne:,}")
json.dump(zrodla_uzyte, open(f"{BASE}/work/spec_zrodla.json", "w"))
