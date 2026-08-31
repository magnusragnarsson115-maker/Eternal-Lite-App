# -*- coding: utf-8 -*-
"""Architektura komponentow Eternal — specyfikacja bramy, dostawcow i certyfikacji."""
import json
import os
import sys
import collections
import datetime

sys.path.insert(0, os.path.dirname(os.path.abspath(__file__)))
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from mkdocx import setup, toc
from dane_komponenty import (K, WARSTWA, SZCZEBEL, WYZWALACZE, SKLADOWE,
                             EKONOMIA, MODULY, BEZPIECZENSTWO)

R = json.load(open('build/KOMPONENTY.json'))
TODAY = datetime.date.today().strftime('%d.%m.%Y')
OUT = '/home/user/Eternal-Lite-App/out/ETERNAL_ARCHITEKTURA_KOMPONENTOW.docx'
CW = collections.Counter(r['warstwa'] for r in R)
CK = collections.Counter(r['klasa'] for r in R)


def cien(cell, hexcol):
    sh = OxmlElement('w:shd')
    sh.set(qn('w:val'), 'clear')
    sh.set(qn('w:fill'), hexcol)
    cell._tc.get_or_add_tcPr().append(sh)


TLA = {'A': 'D7F0DD', 'B': 'F3E3C3', 'C': 'F8D7DA'}


def bogaty(p, txt, size=10):
    for i, cz in enumerate(str(txt).split('**')):
        if cz:
            r = p.add_run(cz)
            r.font.size = Pt(size)
            r.bold = bool(i % 2)


def tab(doc, rows, szer, warstwa_kol=None, naglowek='1B3A6B'):
    t = doc.add_table(rows=0, cols=len(rows[0]))
    t.style = 'Table Grid'
    for ri, row in enumerate(rows):
        cells = t.add_row().cells
        for ci, val in enumerate(row):
            cells[ci].text = ''
            p = cells[ci].paragraphs[0]
            bogaty(p, val, 8.5 if ri else 8.5)
            if ri == 0:
                for r_ in p.runs:
                    r_.bold = True
                    r_.font.color.rgb = RGBColor.from_string('FFFFFF')
                cien(cells[ci], naglowek)
            elif warstwa_kol is not None and ci == warstwa_kol and str(val) in TLA:
                cien(cells[ci], TLA[str(val)])
            cells[ci].width = Cm(szer[ci])
    return t


doc = Document()
setup(doc)
s = doc.sections[0]
s.orientation = WD_ORIENT.LANDSCAPE
s.page_width, s.page_height = s.page_height, s.page_width
s.left_margin = s.right_margin = Cm(1.8)

for txt, sz, bold, col in [
        ('ETERNAL LIFE', 26, True, 'B8431F'),
        ('Architektura komponentów i brama dostawców', 16, True, '1B3A6B'),
        ('Co stoi za każdą funkcją, kto to dostarcza, kiedy trzeba to zmienić', 12, False, None)]:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(txt)
    r.font.size = Pt(sz)
    r.bold = bold
    if col:
        r.font.color.rgb = RGBColor.from_string(col)
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('%d funkcji · %d klas komponentów · %d funkcji wymagających certyfikacji\n'
              'Stan na %s' % (len(R), len(K), CW['C'], TODAY))
r.font.size = Pt(10)
doc.add_page_break()
doc.add_heading('Spis treści', 1)
toc(doc)
doc.add_page_break()

# ---------------------------------------------------------------- 1
doc.add_heading('1. Po co ten dokument', 1)
for t in [
 'Specyfikacja Master 5.4 zawiera polecenie, którego nigdy nie wykonano: **„Do każdej karty '
 'funkcji dopisać: KLASA KOMPONENTU, WARSTWA ZGODNOŚCI (A/B/C), CZAS WYJŚCIA w dniach '
 'i PRÓG ZMIANY wyrażony liczbą. Bez tych czterech pól karta opisuje jedną piątą pracy.”** '
 'Ten dokument i towarzysząca mu macierz wykonują to polecenie dla wszystkich %d funkcji.'
 % len(R),
 'Master 5.4 wyróżnia pięć klas komponentów — Środowisko, Zgodność, Architektura, Dane '
 'i Funkcjonalna — i sam stwierdza, że **w specyfikacji obecna jest tylko piąta**. Cztery '
 'pozostałe są tutaj uzupełnione: każda klasa funkcjonalna ma wskazane, jakie składowe '
 'środowiska, zgodności, architektury i danych pociąga za sobą.',
 'Trzecia rzecz, której w korpusie nie było wcale: **brama nie tylko wywołuje dostawcę, '
 'ale przypisuje do niego użytkownika**. To jest różnica między routerem a instrumentem '
 'kontroli. Router wybiera trasę dla żądania; brama Eternal wybiera dostawcę dla osoby, '
 'na cały okres jej obecności w systemie, i mierzy, ile ten wybór kosztuje i ile przynosi.',
 'Dokument nie zastępuje kart funkcji ze specyfikacji. Dokłada do nich warstwę, której '
 'tam nie ma, i wskazuje, gdzie liczby są szacunkiem, a nie cennikiem.',
]:
    p = doc.add_paragraph()
    bogaty(p, t)

doc.add_heading('Rozkład wyniku', 2)
tab(doc, [['Miara', 'Wartość', 'Co z tego wynika']] + [
 ['Funkcji z przypisanym komponentem', str(len(R)), 'Komplet rejestru — nie ma funkcji bez klasy'],
 ['Klas komponentów', '%d (25 z Master 5.4 + 5 uzupełnionych)' % len(K),
  'K10 płatności, K19 tożsamość, K24 kolejki, K29 laboratoria, K30 wsparcie psychologiczne '
  'nie miały definicji w korpusie'],
 ['Warstwa A — wellness, poza MDR', str(CW['A']),
  '%d%% funkcji sprzedawalnych bez certyfikacji od dnia pierwszego' % round(100 * CW['A'] / len(R))],
 ['Warstwa B — klinika, poza MDR', str(CW['B']),
  'Dokumentacja i organizacja wizyty; %d z nich to granica — własna ocena przenosi je do C'
  % sum(1 for r in R if r['granica'] == 'TAK')],
 ['Warstwa C — wyrób klasy IIa i wyżej', str(CW['C']),
  'Tylko te wymagają dossier albo modelu proxy do cudzego CE'],
], [7.0, 5.5, 13.5])

# ---------------------------------------------------------------- 2
doc.add_page_break()
doc.add_heading('2. Brama Eternal — cztery funkcje, nie jedna', 1)
doc.add_paragraph(
 'Brama jest opisywana w korpusie jako „prawo wstępu — Eternal Kompatybilny plus klucz do '
 'Gateway”. To jedna czwarta tego, czym musi być, żeby reguła 33% w ogóle dała się '
 'wyegzekwować.')
tab(doc, [['Funkcja bramy', 'Na czym polega', 'Bez tego']] + [
 ['1. Przypisanie', 'Przy rejestracji użytkownik dostaje przypisanie do konkretnego dostawcy '
  'w każdej klasie komponentu. Przypisanie jest **trwałe** — historia zdrowotna jednej osoby '
  'pochodzi z jednego źródła, bo dane z dwóch urządzeń o różnej kalibracji nie są '
  'porównywalne w czasie.',
  'Reguła 33% nie ma czego mierzyć. „Udział dostawcy” bez przypisania to udział w ruchu, '
  'a nie w bazie użytkowników'],
 ['2. Wywołanie', 'Rdzeń nigdy nie woła API dostawcy. Woła adapter klasy. Wymiana dostawcy '
  'to wymiana jednego pliku.',
  'Każda zmiana dostawcy jest przepisaniem funkcji. Plan wyjścia zostaje deklaracją'],
 ['3. Pomiar', 'Brama liczy wywołania per użytkownik per dostawca. To jedyne miejsce, '
  'w którym koszt dostawcy da się podzielić na osoby, a nie na faktury.',
  'Nie da się policzyć marży na użytkowniku. LTV pozostaje liczbą z prezentacji, '
  'a nie z rachunku'],
 ['4. Odebranie', 'Prawo przeniesienia użytkownika do innego dostawcy bez zmiany w kodzie '
  'aplikacji i bez udziału dostawcy, którego się opuszcza.',
  'Pozostałe trzy prawa są deklaracją — to sformułowanie z Master 5.4 §7'],
], [4.5, 12.5, 9.0])

doc.add_heading('2.1 Jak brama wybiera dostawcę dla użytkownika', 2)
doc.add_paragraph(
 'Wybór nie jest losowy i nie jest wyłącznie równoważeniem obciążenia. Przechodzi przez '
 'cztery filtry w tej kolejności — pierwszy, który odrzuca dostawcę, kończy sprawę:')
tab(doc, [['#', 'Filtr', 'Reguła', 'Przykład']] + [
 ['1', 'Zgodność', 'Warstwa funkcji ogranicza zbiór dopuszczalnych dostawców. Funkcja '
  'warstwy C nie może trafić do dostawcy klasy wellness — niezależnie od ceny i jakości.',
  'A6.8 Predykcja ryzyka chorób nie pójdzie do Terra API. Może pójść do dostawcy '
  'z własnym CE (Vitalera) albo do naszego wyrobu z dossier'],
 ['2', 'Możliwość', 'Sprzęt i system użytkownika. Nie każdy dostawca obsługuje każde '
  'urządzenie.',
  'Użytkownik iPhone z Apple Watch → HealthKit. Użytkownik Garmina w programie '
  'pracowniczym → Terra albo Rook'],
 ['3', 'Udział', 'Reguła 33%: dostawca, który obsługuje już jedną trzecią aktywnych '
  'użytkowników w klasie, przestaje przyjmować nowych. Ostrzeżenie zapala się przy 25%.',
  'Gdy Terra przekroczy 33% klasy K01, nowi użytkownicy Garmina idą do Rook, '
  'nawet jeśli Terra jest wygodniejsza'],
 ['4', 'Koszt', 'Dopiero na końcu — spośród dostawców, którzy przeszli trzy poprzednie '
  'filtry, wybierany jest ten o niższym koszcie na użytkownika.',
  'HealthKit (0 zł) przed Junction (2,00 zł/user/mies), o ile obsługuje urządzenie'],
], [1.0, 3.0, 11.0, 11.0])

doc.add_paragraph()
p = doc.add_paragraph()
bogaty(p, 'Kolejność filtrów jest istotna. **Odwrócenie jej — najpierw koszt, potem '
 'zgodność — jest najczęstszym sposobem, w jaki architektura tego typu łamie prawo bez '
 'niczyjej złej woli.** Tańszy dostawca wygrywa przetarg wewnętrzny, a fakt, że nie ma CE, '
 'wychodzi przy pierwszym audycie.')

doc.add_heading('2.2 Kiedy brama przepina użytkownika', 2)
doc.add_paragraph(
 'Przypisanie jest trwałe, ale nie wieczne. Przepięcie następuje w oknie migracyjnym '
 '(proponowane: pierwszy tydzień miesiąca), z zapisem w dzienniku i z zachowaniem ciągłości '
 'danych — nowe źródło jest kalibrowane wobec starego przez okres nakładania się serii.')
tab(doc, [['Zdarzenie', 'Reakcja bramy']] + [
 ['Dostawca przekroczył 33% klasy', 'Stop dla nowych przypisań; migracja nadwyżki '
  'w kolejnych oknach, zaczynając od użytkowników o najniższym wolumenie danych'],
 ['Przekroczony próg kosztowy klasy', 'Przeniesienie ruchu na wariant A (open source) '
  'albo C (własne) — patrz wyzwalacz W1'],
 ['Funkcja przeszła z warstwy A/B do C', 'Natychmiastowe wyłączenie dostawców bez CE '
  'dla tej funkcji. To wyzwalacz W3 i nie ma okna migracyjnego'],
 ['Awaria albo złamanie SLA', 'Przełączenie na wariant zapasowy w czasie zapisanym '
  'w karcie funkcji (kolumna „czas wyjścia”)'],
 ['Żądanie użytkownika', 'Przepięcie na wniosek, z informacją o skutku dla porównywalności '
  'historii pomiarów'],
], [8.0, 18.0])

# ---------------------------------------------------------------- 3
doc.add_page_break()
doc.add_heading('3. Reguła 33% i ekonomia na użytkowniku', 1)
for t in [
 'Reguła 33% pochodzi z Master 5.4: żaden dostawca nie obsługuje więcej niż jednej trzeciej '
 'aktywnych użytkowników **w obrębie klasy komponentu**. Minimum to trzy warianty '
 'zaimplementowane, z czego przynajmniej jeden żywy, z 1–5% realnego ruchu. Wariant '
 'wypisany w dokumencie, ale nieuruchomiony, nie jest wariantem — jest planem.',
 'Reguła ma cenę i korpus ją nazywa: **utrata rabatu wolumenowego, potrojony nakład '
 'integracyjny, trzy warianty w dossier, trzy umowy powierzenia.** To nie jest darmowa '
 'ostrożność i nie należy jej stosować wszędzie.',
 'Reguła **nie obowiązuje** tam, gdzie nie ma czego dywersyfikować: P1 i Centrum e-Zdrowia, '
 'jednostka notyfikowana, reżim MDR, producenci pasków. Odpowiedzią na monopol jest stanie '
 'się niezbędnym, nie redundancja.',
 'Pułapka, którą korpus nazywa wprost: **trzej dostawcy na jednej chmurze to jedna zależność '
 'w trzech opakowaniach.** Reguła 33% mierzy dostawcę. Potrzebny jest drugi licznik — '
 'na poziomie technologii źródłowej. Terra, Rook i Junction sięgające po ten sam SDK Garmina '
 'to formalnie trzech dostawców i faktycznie jeden punkt awarii.',
]:
    p = doc.add_paragraph()
    bogaty(p, t)

doc.add_heading('3.1 Koszt wariantów w przeliczeniu na użytkownika', 2)
doc.add_paragraph(
 'To jest liczba, której w dokumentacji nie było: ile kosztuje jeden użytkownik miesięcznie '
 'w każdej klasie. Bez niej marża na użytkowniku jest nieobliczalna, a to właśnie jej brak '
 'jest przedmiotem ustalenia nr 10 z analizy poprawności (LTV liczone dla subskrypcji, '
 'której w modelu darmowym nie ma).')
tab(doc, [['Klasa', 'Dostawca', 'Model rozliczenia', 'PLN/user/mies', 'Udział %', 'Podstawa']]
    + [[k, d, m, ('%.2f' % c).replace('.', ','), str(u), p_] for k, d, m, c, u, p_ in EKONOMIA],
    [2.0, 6.5, 5.5, 2.4, 2.0, 8.0])

koszt = {}
for k, d, m, c, u, p_ in EKONOMIA:
    koszt.setdefault(k, []).append(c * u / 100.0)
suma = sum(sum(v) for v in koszt.values())
doc.add_paragraph()
p = doc.add_paragraph()
bogaty(p, 'Suma ważona udziałami dla klas wycenionych: **%.2f PLN na użytkownika '
 'miesięcznie**. To nie jest pełny koszt jednostkowy — obejmuje %d z %d klas, bez '
 'wynagrodzeń, bez K20 (certyfikat P1 i KS-BLOZ to koszt roczny, nie per user) i bez '
 'kosztów sprzętu. Traktować jako dolną granicę, nie jako wynik.'
 % (suma, len(koszt), len(K)))
p = doc.add_paragraph()
bogaty(p, 'Dla porządku: **wynagrodzenia są największą pozycją kosztową i nie występują '
 'ani w tym rachunku, ani w żadnym z siedemnastu kosztów stałych wymienionych w korpusie.** '
 'Master 5.4 nazywa to „kategorią nienazwaną”. Marża na użytkowniku policzona bez nich '
 'jest marżą brutto na infrastrukturze, niczym więcej.')

# ---------------------------------------------------------------- 4
doc.add_page_break()
doc.add_heading('4. Warstwy zgodności — med czy wellness', 1)
doc.add_paragraph(
 'Podział na trzy warstwy pochodzi z Master 5.4 i jest tym elementem strategii, który '
 'w niezależnej analizie wypadł najlepiej: pozwala wejść na rynek bez certyfikacji, '
 'nie łamiąc prawa. Warunkiem jest konsekwencja — jedna funkcja przesunięta z A do C '
 'bez zauważenia przenosi cały produkt w reżim MDR.')
tab(doc, [['Warstwa', 'Zakres', 'Charakter', 'Certyfikacja', 'Funkcji', 'Co obejmuje']]
    + [[w, v[0], v[2], v[3], str(CW.get(w, 0)), v[1]] for w, v in WARSTWA.items()],
    [2.0, 4.5, 2.6, 3.4, 1.8, 12.0], warstwa_kol=0)

doc.add_heading('4.1 Granica, o którą się potyka', 2)
gran = [r for r in R if r['granica'] == 'TAK']
p = doc.add_paragraph()
bogaty(p, '%d funkcji siedzi dokładnie na granicy B/C: mierzą albo pokazują parametr, '
 'ale go nie interpretują. **Dopóki aplikacja pokazuje liczbę, jest to warstwa B. '
 'W chwili, gdy pokazuje ją na czerwono z podpisem „poza normą”, jest to warstwa C.** '
 'Różnica nie leży w pomiarze, tylko w tym, kto wyciąga wniosek.' % len(gran))
doc.add_paragraph(
 'Korpus wskazuje A3.5 „Trójkolorowe alerty” jako warstwę C i ma rację — trzy kolory to '
 'ocena, nie prezentacja. Ta sama funkcja bez kolorów, z samą liczbą i zakresem '
 'referencyjnym podanym przez laboratorium, zostaje w warstwie A.')
tab(doc, [['Kod', 'Funkcja', 'Klasa', 'Co przenosi ją do C']]
    + [[r['kod'], r['nazwa'][:70], r['klasa'], 'Własna ocena, próg kliniczny albo '
        'oznaczenie wartości jako nieprawidłowej'] for r in gran],
    [2.0, 10.0, 2.2, 12.0])

# ---------------------------------------------------------------- 5
doc.add_page_break()
doc.add_heading('5. Pięć szczebli kontroli i najdroższy krok', 1)
tab(doc, [['Szczebel', 'Co znaczy', 'Kontrola', 'Rola wg MDR']]
    + [[str(k_), v[0], v[1], v[2]] for k_, v in sorted(SZCZEBEL.items())],
    [2.0, 7.0, 7.0, 10.0])
doc.add_paragraph()
p = doc.add_paragraph()
bogaty(p, '**Najdroższy krok w całej architekturze zachodzi przez umieszczenie logotypu '
 'na obudowie.** Wprowadzenie wyrobu do obrotu pod własną nazwą czyni z nas producenta '
 'ze wszystkimi obowiązkami: dossier, PRRC, EUDAMED. To decyzja regulacyjna o koszcie '
 'w setkach tysięcy, podejmowana zwykle jako decyzja marketingowa. Szczebel trzeci — '
 'partnerstwo — jest w większości przypadków optymalny i niedoceniany.')
doc.add_paragraph(
 'Trzy decyzje z Master 5.4 są ze sobą powiązane i nie da się ich rozstrzygnąć osobno: '
 'czy wydajemy własną ocenę, czy sprzęt nosi naszą markę i czy użytkownik widzi nazwy '
 'dostawców. Trzy „tak” dają jeden model, trzy „nie” dają inny, a mieszanka daje sprzeczność, '
 'którą zauważy pierwszy audytor.')

# ---------------------------------------------------------------- 6
doc.add_page_break()
doc.add_heading('6. Kiedy zmienić model — osiem wyzwalaczy', 1)
doc.add_paragraph(
 'Model dostawcy nie jest wyborem raz na zawsze. Poniżej osiem zdarzeń, po których '
 'dotychczasowy układ przestaje obowiązywać. Każde ma podaną miarę — próg wyrażony liczbą, '
 'nie wrażeniem.')
tab(doc, [['Kod', 'Wyzwalacz', 'Co jest mierzone', 'Co się zmienia', 'Źródło']]
    + [list(w) for w in WYZWALACZE], [1.6, 5.0, 8.5, 8.5, 4.4], naglowek='B8431F')

# ---------------------------------------------------------------- 7
doc.add_page_break()
doc.add_heading('7. Gotowe moduły — kupić, licencjonować, partnerować czy zbudować', 1)
for t in [
 'Pytanie „budować czy kupić” ma w tej architekturze piątą odpowiedź, której zwykle się '
 'nie stawia: **kupić dostęp, zbudować warstwę danych, i pozwolić, żeby to warstwa danych '
 'decydowała o tym, kto jest zastępowalny.** Partner dostarcza usługę; my dostarczamy '
 'kontekst, w którym ta usługa ma sens. Kontekstu nie da się przenieść do konkurenta.',
 'Jedno zastrzeżenie do postawionego pytania. Konkurowanie z platformą lepszym produktem, '
 'niższą ceną albo lepszą dystrybucją jest normalną strategią i tak są opisane poniższe '
 'pozycje. **Działania wymierzone w podstawy funkcjonowania konkurenta poza konkurencją '
 'na rynku nie są tu opisywane i nie są potrzebne** — w tym układzie wystarczy zbudować '
 'substytut wewnątrz ekosystemu i obniżyć koszt przejścia dla użytkownika. O reszcie '
 'zdecyduje dystrybucja.',
]:
    p = doc.add_paragraph()
    bogaty(p, t)
tab(doc, [['Podmiot', 'Co robi', 'Co daje nam', 'Czego nie daje', 'Nasza postawa',
           'Warunek zmiany postawy']] + [list(m) for m in MODULY],
    [4.2, 4.6, 4.6, 4.6, 4.4, 5.6])

# ---------------------------------------------------------------- 8
doc.add_page_break()
doc.add_heading('8. Bezpieczeństwo bramy', 1)
doc.add_paragraph(
 'Brama, która tylko routuje, jest pojedynczym punktem, przez który przechodzą wszystkie '
 'dane zdrowotne w systemie. Osiem zasad poniżej odróżnia ją od takiego punktu.')
tab(doc, [['Zasada', 'Na czym polega', 'Źródło']] + [list(b) for b in BEZPIECZENSTWO],
    [6.0, 14.0, 6.0])

# ---------------------------------------------------------------- 9
doc.add_page_break()
doc.add_heading('9. Składowe I–IV — czego nie było w specyfikacji', 1)
doc.add_paragraph(
 'Master 5.4 wyróżnia pięć klas komponentów i sam stwierdza, że obecna w specyfikacji jest '
 'wyłącznie piąta — funkcjonalna. Poniżej cztery pozostałe, z konkretnym wskazaniem, '
 'co je wypełnia. Każda klasa funkcjonalna w macierzy ma wskazane, które z nich pociąga.')
tab(doc, [['Klasa', 'Nazwa', 'Zakres wg Master 5.4', 'Co konkretnie w Eternal']]
    + [[k_, v[0], v[1], v[2]] for k_, v in SKLADOWE.items()], [2.0, 3.4, 8.0, 12.6])

doc.add_heading('9.1 Trzydzieści klas funkcjonalnych', 2)
doc.add_paragraph(
 'Dwadzieścia pięć klas pochodzi z Master 5.4. Pięć — K10 płatności, K19 tożsamość i zgody, '
 'K24 kolejki, K29 laboratoria zewnętrzne, K30 wsparcie psychologiczne — uzupełnia lukę: '
 'w korpusie nie miały definicji, mimo że funkcje z nich korzystające istnieją. '
 'K24 nie miał w korpusie nawet nazwy.')
tab(doc, [['Klasa', 'Nazwa', 'Funkcji', 'Wariant startowy', 'Próg wyjścia', 'Szczebel docelowy']]
    + [[k_, v[0], str(CK.get(k_, 0)), v[6], v[4], '%d — %s' % (v[7], SZCZEBEL[v[7]][0])]
       for k_, v in sorted(K.items())],
    [2.0, 5.4, 1.8, 8.0, 5.4, 3.4])

# ---------------------------------------------------------------- 10
doc.add_page_break()
doc.add_heading('10. Certyfikacja — co, kiedy i którą ścieżką', 1)
cert = [r for r in R if r['warstwa'] == 'C']
mod_c = collections.Counter('%s — %s' % (r['modul'], r['modul_nazwa']) for r in cert)
p = doc.add_paragraph()
bogaty(p, '**%d funkcji z %d wymaga certyfikacji** — to %d%% rejestru. Reszta jest '
 'sprzedawalna bez jednostki notyfikowanej od dnia pierwszego. Ta proporcja jest wynikiem '
 'strategii warstwowej, nie przypadkiem: warstwy zostały zaprojektowane tak, żeby '
 'certyfikacja dotyczyła możliwie wąskiego rdzenia.'
 % (len(cert), len(R), round(100 * len(cert) / len(R))))
tab(doc, [['Moduł', 'Funkcji w C', 'Etapy', 'Klasy komponentów', 'Ścieżka']]
    + [[m, str(n),
        ', '.join(sorted({r['etap'] for r in cert
                          if '%s — %s' % (r['modul'], r['modul_nazwa']) == m})),
        ', '.join(sorted({r['klasa'] for r in cert
                          if '%s — %s' % (r['modul'], r['modul_nazwa']) == m})),
        ('Proxy do cudzego CE (Labplus) albo własne dossier'
         if any(r['klasa'] == 'K28' for r in cert
                if '%s — %s' % (r['modul'], r['modul_nazwa']) == m)
         else 'Dossier klasy IIa: 80–150 tys. zł, 6–12 mies.')]
       for m, n in mod_c.most_common()],
    [7.5, 2.2, 4.0, 4.0, 8.3], naglowek='B8431F')

doc.add_heading('10.1 Kolejność, w jakiej to robić', 2)
tab(doc, [['Kiedy', 'Co', 'Dlaczego teraz']] + [
 ['Etap MVP (2026)', 'Zero certyfikacji. Warstwy A i B, %d funkcji. Proxy do Labplus '
  'dla wyników laboratoryjnych.' % (CW['A'] + CW['B']),
  'Certyfikacja przed przychodem to wydanie 80–150 tys. zł na hipotezę'],
 ['Pierwszy przychód B2B', 'Klasyfikacja wg MDCG 2019-11 dla funkcji granicznych, '
  'deklaracja przeznaczenia per funkcja.',
  'Klient B2B zapyta o to w pierwszej rozmowie. Odpowiedź „jeszcze nie sprawdzaliśmy” '
  'kończy rozmowę'],
 ['12–24 miesiące', 'Dossier klasy IIa dla rdzenia warstwy C: A3.5, A6.5, A6.8, D2.x. '
  'Równolegle system zarządzania jakością ISO 13485 i PRRC.',
  'To jest ścieżka wskazana w samym korpusie; wcześniej nie ma czego certyfikować'],
 ['Po dossier', 'Rejestracja w EUDAMED, UDI, nadzór po wprowadzeniu do obrotu.',
  'Obowiązki producenta nie kończą się na certyfikacie — zaczynają się na nim'],
 ['Osobno, niezależnie', 'AI Act: system wysokiego ryzyka wg załącznika III, obowiązek '
  'oznaczania treści generowanej od 2.08.2026.',
  'To osobny reżim obok MDR, nie jego część. W decku go nie ma — patrz ustalenie 7 '
  'analizy poprawności'],
], [4.5, 11.0, 10.5])

# ---------------------------------------------------------------- 11
doc.add_page_break()
doc.add_heading('11. Czego w tym dokumencie nie ma', 1)
for t in [
 '**Cenniki dostawców są z korpusu i z publicznych stron, nie z ofert.** Żadna z liczb '
 'nie pochodzi z negocjacji. Vitalera nie ma publicznego cennika — pozycja jest oznaczona '
 'jako [BRAK], a nie oszacowana.',
 '**Koszty na użytkownika oznaczone [SZACUNEK] to przeliczenie wolumenu, nie pomiar.** '
 'Przyjęte założenia (2 dokumenty OCR, 50 zapytań do modelu, 20 minut nagrań, 15 zdjęć '
 'posiłków na użytkownika miesięcznie) są podane wprost i wymagają weryfikacji na realnym '
 'ruchu. Przy innych założeniach wynik zmienia się o rząd wielkości.',
 '**Podział 33/33/33 jest celem, nie stanem.** Dziś w każdej klasie działa jeden wariant '
 'albo żaden. Doprowadzenie do trzech zaimplementowanych wariantów w kluczowych klasach '
 'to praca, którą korpus wycenia jako „potrojony nakład integracyjny” i której nikt '
 'jeszcze nie zaczął.',
 '**Deklaracja CE dostawcy nie została zweryfikowana u źródła.** Vitalera deklaruje '
 'oznakowanie CE wg MDR i tak jest to zapisane. Przed oparciem na tym funkcji warstwy C '
 'trzeba sprawdzić numer w EUDAMED i zakres przeznaczenia — deklaracja producenta '
 'to nie to samo co zgodność z naszym zastosowaniem.',
 '**Przypisanie funkcji do klasy komponentu jest regułowe, nie ręczne.** Reguły opierają '
 'się na module i na treści nazwy funkcji. Kontrola na czterech przypadkach wskazanych '
 'w korpusie (A3.5, A6.5, A6.8, D2.x) wypada zgodnie, ale pojedyncze przypisania '
 'w rejestrze mogą wymagać korekty przy pisaniu karty funkcji.',
 '**Pole „klasa MDR” ze źródłowego rejestru nie zostało użyte.** Jest artefaktem '
 'ekstrakcji: jako klasa IIb oznaczone są tam „Dashboard główny” i „Ręczne dodawanie '
 'danych”. Warstwa zgodności jest wyprowadzona z definicji, nie z tego pola.',
]:
    p = doc.add_paragraph()
    bogaty(p, t)

doc.add_paragraph()
p = doc.add_paragraph()
r = p.add_run('Dokument towarzyszy macierzy ETERNAL_MACIERZ_KOMPONENTOW.xlsx, w której '
              'każda z %d funkcji ma komplet przypisań. Podstawa: Specyfikacja Master 5.4 '
              '(sekcje 4.2, 7 i 11), Roadmapa v4 (tory alternatywne), analiza Vitalera '
              'z korpusu pytań i odpowiedzi. Stan na %s.' % (len(R), TODAY))
r.font.size = Pt(8.5)
r.italic = True
r.font.color.rgb = RGBColor.from_string('5D6B8A')

doc.save(OUT)
zn = (sum(len(p.text) for p in doc.paragraphs)
      + sum(len(c.text) for t in doc.tables for r in t.rows for c in r.cells))
print('%s -> %d B, %d tabel, ~%d stron' % (OUT, os.path.getsize(OUT), len(doc.tables),
                                           max(1, round(zn / 2600))))
