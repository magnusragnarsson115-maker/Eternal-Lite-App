# -*- coding: utf-8 -*-
"""Analiza wlasna: poprawnosc, legalnosc i oplacalnosc dokumentacji Eternal — DOCX.

Format wynika z reguly uzytkownika: HTML zostaje wylacznie dla roadmapy,
specyfikacja, biznesplan i ta analiza sa w DOCX, pitch decki w PPTX.
"""
import os
import sys
import datetime

sys.path.insert(0, os.path.dirname(os.path.abspath(__file__)))
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from mkdocx import setup, toc
from dane_analiza import BLEDY, ZRODLA, FINANSE, DOBRZE, WAGI, KOLOR, OPIS_WAGI

TODAY = datetime.date.today().strftime('%d.%m.%Y')
OUT = '/home/user/Eternal-Lite-App/out/ETERNAL_ANALIZA_POPRAWNOSCI.docx'


def bogaty(p, txt, size=10, color=None):
    """Wstawia tekst z wyroznieniem **pogrubionym** jako osobne runy."""
    for i, czesc in enumerate(txt.split('**')):
        if not czesc:
            continue
        r = p.add_run(czesc)
        r.font.size = Pt(size)
        r.bold = bool(i % 2)
        if color:
            r.font.color.rgb = RGBColor.from_string(color)
    return p


def cien(paragraph_or_cell, hexcol):
    el = paragraph_or_cell._tc if hasattr(paragraph_or_cell, '_tc') else paragraph_or_cell._p
    sh = OxmlElement('w:shd')
    sh.set(qn('w:val'), 'clear')
    sh.set(qn('w:fill'), hexcol)
    (el.get_or_add_tcPr() if hasattr(paragraph_or_cell, '_tc')
     else paragraph_or_cell._p.get_or_add_pPr()).append(sh)


def pasek(doc, txt, hexcol):
    """Etykieta wagi jako jednokomorkowa tabela z tlem."""
    t = doc.add_table(rows=1, cols=1)
    t.alignment = WD_TABLE_ALIGNMENT.LEFT
    c = t.rows[0].cells[0]
    c.width = Cm(4.2)
    cien(c, hexcol)
    c.text = ''
    r = c.paragraphs[0].add_run(txt)
    r.font.size = Pt(8)
    r.bold = True
    r.font.color.rgb = RGBColor.from_string('FFFFFF')
    return t


def tabela(doc, rows, szer=None, naglowek='1B3A6B'):
    n = max(len(r) for r in rows)
    t = doc.add_table(rows=0, cols=n)
    t.style = 'Table Grid'
    for ri, row in enumerate(rows):
        cells = t.add_row().cells
        for ci in range(n):
            cells[ci].text = ''
            p = cells[ci].paragraphs[0]
            bogaty(p, row[ci] if ci < len(row) else '', size=8.5,
                   color='FFFFFF' if ri == 0 else None)
            for r in p.runs:
                if ri == 0:
                    r.bold = True
            if ri == 0:
                cien(cells[ci], naglowek)
            if szer and ci < len(szer):
                cells[ci].width = Cm(szer[ci])
    return t


doc = Document()
setup(doc)

# --- strona tytulowa ---
for txt, sz, bold, col in [
        ('ETERNAL LIFE', 26, True, 'B8431F'),
        ('Analiza poprawności, legalności i opłacalności', 16, True, '1B3A6B'),
        ('Przegląd dokumentacji przeprowadzony niezależnie od jej autorów', 12, False, None)]:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(txt)
    r.font.size = Pt(sz)
    r.bold = bold
    if col:
        r.font.color.rgb = RGBColor.from_string(col)

kryt = sum(1 for b in BLEDY if b[0] == 'KRYT')
wys = sum(1 for b in BLEDY if b[0] == 'WYS')
sr = sum(1 for b in BLEDY if b[0] == 'SR')
nisk = sum(1 for b in BLEDY if b[0] == 'NISK')

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('%d ustaleń: %d krytyczne, %d wysokie, %d średnie, %d niskie\nStan na %s'
              % (len(BLEDY), kryt, wys, sr, nisk, TODAY))
r.font.size = Pt(10)
doc.add_page_break()

# --- wprowadzenie ---
doc.add_heading('Czego dotyczy ta analiza', 1)
for t in [
    'To nie jest ocena pomysłu. Pomysł jest spójny, a dokumentacja nietypowo szczegółowa jak na etap '
    'pre-seed. To lista miejsc, w których dokumenty przeczą sobie nawzajem, liczby się nie domykają '
    'albo deklaracja niesie zobowiązanie prawne, którego autorzy prawdopodobnie nie zamierzali podjąć.',
    'Każde ustalenie ma trzy pola: co jest w dokumentach, na czym polega błąd i co konkretnie z nim '
    'zrobić. Trzecie pole jest propozycją rozwiązania, nie tylko wskazaniem problemu.',
    'Podstawą jest oficjalny pitch deck (32 slajdy), Specyfikacja Master 5.4, Plan Korporacyjny 5.1, '
    'Biznesplan rozszerzony i Plan PWNŚ. Liczby finansowe policzone bezpośrednio z danych podanych '
    'w decku. Liczby rynkowe zweryfikowane w źródłach zewnętrznych wskazanych w aneksie — indeks '
    'plików archiwum jest osobnym dokumentem i nie służy tu jako podstawa.',
    'Jedno ustalenie (nr 10, ekonomika jednostkowa) jest konsekwencją mojej własnej rekomendacji '
    'darmowej aplikacji i zgłaszam je jako otwarte, a nie rozstrzygnięte.',
]:
    doc.add_paragraph(t)

doc.add_heading('Rozkład ustaleń', 2)
tabela(doc, [['Waga', 'Liczba', 'Co to znaczy']] +
       [[WAGI[k], str(n), OPIS_WAGI[k]]
        for k, n in [('KRYT', kryt), ('WYS', wys), ('SR', sr), ('NISK', nisk)]],
       szer=[3.0, 2.0, 11.0])

doc.add_page_break()
doc.add_heading('Spis treści', 1)
toc(doc)
doc.add_page_break()

# --- rachunek finansowy ---
doc.add_heading('Podsumowanie finansowe — na czym nie domyka się model', 1)
doc.add_paragraph(
    'Trzy ustalenia krytyczne dotyczą tej samej rzeczy: pieniędzy. Poniżej rachunek, który prowadzi '
    'do każdego z nich, policzony z liczb podanych w samym decku.')
tabela(doc, FINANSE, szer=[5.6, 2.0, 1.7, 1.7, 1.7, 1.7, 2.2])
doc.add_paragraph(
    'Skumulowana strata do osiągnięcia progu rentowności wynosi 8,11 mln PLN, a kapitał pozyskany '
    'do rundy A — od 6,11 do 6,81 mln PLN. Różnica 1,30–2,00 mln PLN nie ma w decku pokrycia, '
    'ponieważ runda A nie ma podanej daty.')

doc.add_page_break()

# --- ustalenia ---
doc.add_heading('Ustalenia według wagi', 1)
doc.add_paragraph(
    'Kolejność od najpoważniejszych. Pole „Co z tym zrobić” zawiera propozycję rozwiązania.')

for i, (w, tyt, jest, blad, fix) in enumerate(BLEDY, 1):
    pasek(doc, WAGI[w], KOLOR[w])
    doc.add_heading('%d. %s' % (i, tyt), 2)
    for etykieta, tresc in [('Co jest w dokumentach', jest),
                            ('Na czym polega błąd', blad),
                            ('Co z tym zrobić', fix)]:
        p = doc.add_paragraph()
        r = p.add_run(etykieta.upper())
        r.font.size = Pt(8)
        r.bold = True
        r.font.color.rgb = RGBColor.from_string(
            'B8431F' if etykieta == 'Co z tym zrobić' else '5D6B8A')
        p2 = doc.add_paragraph()
        p2.paragraph_format.left_indent = Cm(0.4)
        bogaty(p2, tresc)

doc.add_page_break()

# --- mocne strony ---
doc.add_heading('Co jest zrobione dobrze', 1)
doc.add_paragraph('Uczciwość wymaga wskazania także tego, co wytrzymuje kontrolę.')
tabela(doc, DOBRZE, szer=[5.0, 11.0])

doc.add_page_break()

# --- zrodla ---
doc.add_heading('ANEKS — źródła zewnętrzne użyte w tej analizie', 1)
doc.add_paragraph(
    'Weryfikacja liczb rynkowych i wymogów regulacyjnych. Indeks plików archiwum jest osobnym '
    'dokumentem — tutaj wyłącznie źródła zewnętrzne.')
for tyt, url, opis in ZRODLA:
    p = doc.add_paragraph()
    r = p.add_run(tyt)
    r.bold = True
    r.font.size = Pt(10)
    p2 = doc.add_paragraph()
    r2 = p2.add_run(url)
    r2.font.size = Pt(8.5)
    r2.font.color.rgb = RGBColor.from_string('B8431F')
    p3 = doc.add_paragraph()
    p3.paragraph_format.left_indent = Cm(0.4)
    r3 = p3.add_run(opis)
    r3.font.size = Pt(9)

doc.add_paragraph()
p = doc.add_paragraph()
r = p.add_run(
    'Analiza przygotowana na podstawie oficjalnego pitch decku (32 slajdy), Specyfikacji Master 5.4, '
    'Planu Korporacyjnego 5.1, Biznesplanu rozszerzonego i Planu PWNŚ. '
    'Stan na %s.' % TODAY)
r.font.size = Pt(8.5)
r.italic = True
r.font.color.rgb = RGBColor.from_string('5D6B8A')

os.makedirs(os.path.dirname(OUT), exist_ok=True)
doc.save(OUT)
znakow = (sum(len(p.text) for p in doc.paragraphs)
          + sum(len(c.text) for t in doc.tables for r in t.rows for c in r.cells))
print('%s -> %d B, %d akapitow, %d tabel, ~%d stron, ustalen %d (K%d W%d S%d N%d)'
      % (OUT, os.path.getsize(OUT), len(doc.paragraphs), len(doc.tables),
         max(1, round(znakow / 1800)), len(BLEDY), kryt, wys, sr, nisk))
