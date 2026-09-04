# -*- coding: utf-8 -*-
"""Biznesplan — wersja finalna, skrócona. Sekcje 1..21."""
import os, sys, collections
sys.path.insert(0, os.path.dirname(os.path.abspath(__file__)))
from mkdocx import *
from docx import Document
import dane_bp_final as F
import dane_odczyt_bp as B
import dane_odczyt_rm as R
import dane_produkty as PR
import dane_modele as MD
import rejestr as RJ
import dane_rynek as RY
import dane_analiza as AN
import dane_moduly as MO
import dane_komponenty as KP
import karty

OUT = '/home/user/Eternal-Lite-App/out/ETERNAL_BIZNESPLAN_FINALNY.docx'
doc = Document(); setup(doc)

def H(n, t): doc.add_heading("%s  %s" % (n, t), 1)
def h2(t): doc.add_heading(t, 2)
def P_(t): doc.add_paragraph(t)
def B_(t):
    p = doc.add_paragraph(); r = p.add_run(t); r.bold = True
def LIST(items):
    for it in items: doc.add_paragraph(str(it), style='List Bullet')
def TAB(rows, head=None):
    rows = [list(map(str, r)) for r in rows]
    if head: rows = [list(head)] + rows
    add_table(doc, rows)

for t, sz, b in [("ETERNAL ECOSYSTEM", 26, True), (F.TYTUL, 16, True), (F.PODTYTUL, 11, False)]:
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(t); r.font.size = Pt(sz); r.bold = b
    if sz >= 16: r.font.color.rgb = RGBColor.from_string('1F3864')
p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("Podstawa: pełny odczyt korpusu — 159 plików, 28 618 387 znaków\n"
              "Szkielet: Plan Korporacyjny 5.1 · Stan na %s" % TODAY)
r.font.size = Pt(10)
doc.add_page_break()
doc.add_heading("Nota metodyczna", 1)
for t in F.NOTA: P_(t)
h2("Hierarchia dokumentów biznesowych")
TAB(B.HIERARCHIA)
doc.add_page_break()
doc.add_heading("Spis treści", 1); toc(doc); doc.add_page_break()

# 1
H("1.", "Streszczenie zarządcze")
for t in F.S01: P_(t)
h2("Ustalenia w jednym zestawieniu")
TAB(B.STRESZCZENIE)
doc.add_page_break()

# 2
H("2.", "Problem")
for t in F.S02: P_(t)
h2("Skala problemu w liczbach")
TAB(B.PROBLEM_SKALA)
B_(B.PROBLEM_TEZA)
h2("Bilans z państwem — co zajęte, co zostawione")
TAB(B.BILANS_PANSTWO)
B_(B.BILANS_WZORZEC)
P_(B.LUKA)
doc.add_page_break()

# 3
H("3.", "Rozwiązanie")
for t in F.S03: P_(t)
B_(B.ETAP_ZEROWY)
h2("Zestaw podstawowy — co musi powstać, ile kosztuje, ile trwa")
TAB(B.ZESTAW_PODSTAWOWY)
h2("Elementy bez wariantu minimalnego")
TAB(B.BEZ_WARIANTU)
doc.add_page_break()

# 4
H("4.", "Dlaczego teraz")
for t in F.S04: P_(t)
h2("Kalendarz terminów zewnętrznych i bramek własnych")
TAB(R.DATY)
doc.add_page_break()

# 5
H("5.", "Rynek i segmenty klientów")
for t in F.S05: P_(t)
B_(B.ZASADA_RYNKU)
h2("Segmenty i dostępna część rynku")
TAB(B.SEGMENTY)
P_(B.SEGMENTY_B2B_STACJA)
h2("Luka weterynaryjna")
P_(B.LUKA_WETERYNARYJNA)
doc.add_page_break()

# 6
H("6.", "Produkt i portfel")
for t in F.S06: P_(t)
h2("Metoda doboru funkcji do produktu")
TAB(PR.METODA)
h2("Sześć produktów — skład, odbiorca, monetyzacja")
for p_ in PR.PRODUKTY:
    doc.add_heading("%s  %s" % (p_['kod'], p_['nazwa']), 3)
    B_(p_['claim'])
    TAB([[f, RJ.R[f]['nazwa'][:110], RJ.R[f]['etap'], RJ.R[f]['warstwa']]
         for f in p_['funkcje']], ["Funkcja", "Nazwa", "Etap", "Warstwa"])
    for etk, key in [("Niezastępowalność", 'niezast'), ("Odbiorca", 'odbiorca'),
                     ("Monetyzacja", 'monetyzacja'), ("Samodzielność", 'sam'),
                     ("Etap", 'etap'), ("Ryzyko", 'ryzyko')]:
        if key in p_:
            pp = doc.add_paragraph(); rr = pp.add_run(etk + ". "); rr.bold = True
            pp.add_run(str(p_[key]))
h2("Monetyzacja produktów")
TAB(PR.MONETYZACJA)
P_(PR.MONETYZACJA_ZASADA)
h2("Kiedy moduł staje się osobnym produktem")
TAB(PR.MODUL_NA_PRODUKT)
P_(PR.MODUL_NA_PRODUKT_WNIOSEK)
h2("Nisze i branże — te same funkcje, inna korelacja")
TAB(PR.BRANZE)
P_(PR.BRANZE_ZASADA)
h2("Portfel wg kryteriów własnych")
TAB(B.PORTFEL)
doc.add_page_break()

# 7
H("7.", "Model biznesowy")
for t in F.S07: P_(t)
h2("Kanały przychodu")
TAB(B.KANALY)
h2("Ranking typów przychodu")
TAB(B.RANKING_PRZYCHODU)
B_(B.MARZA_TEZA)
h2("Arytmetyka subskrypcji konsumenckiej")
TAB(B.ARYTMETYKA_ABO)
B_(B.ARYTMETYKA_WNIOSEK)
h2("Model odrzucony")
P_(B.MODEL_ODRZUCONY)
h2("Wczesne źródła przychodu")
TAB(B.WCZESNY_PRZYCHOD)
h2("Kanały przychodu w rejestrze funkcji")
kan = collections.Counter(str(v.get('kanal', '—')) for v in RJ.R.values())
zar = collections.Counter(str(v.get('zarabia', '—')) for v in RJ.R.values())
TAB([[k, str(n), ", ".join(sorted(v['kod'] for v in RJ.R.values()
      if str(v.get('kanal', '—')) == k))[:600]]
     for k, n in kan.most_common()],
    ["Kanał", "Pozycji rejestru", "Kody funkcji"])
P_("Rozkład odpowiedzi na pytanie, czy funkcja zarabia bezpośrednio: "
   + " · ".join("%s — %d" % (a, b) for a, b in zar.most_common()) +
   ". Funkcja, która nie zarabia bezpośrednio, nie jest funkcją zbędną — "
   "jest warunkiem działania kanału, który zarabia.")
doc.add_page_break()

# 8
H("8.", "Regulacje i strategia zgodności")
for t in F.S08: P_(t)
war = collections.Counter(v['warstwa'] for v in RJ.R.values())
TAB([[k] + list(v) + [str(war.get(k, 0))] for k, v in sorted(RJ.WARSTWA.items())],
    ["Warstwa", "Zakres", "Reżim", "Pozycji rejestru"])
h2("Cennik usług regulacyjnych jako kanał przychodu")
TAB(B.CENNIK_HUB)
B_(B.CENNIK_HUB_WNIOSEK)
doc.add_page_break()

# 9
H("9.", "Dane, dowód wartości i badania")
for t in F.S09: P_(t)
B_(B.GLEBIA)
h2("Zasoby potrzebne i moment, w którym są potrzebne")
TAB(B.ZASOBY)
P_(B.ZASOB_GLOWNY)
doc.add_page_break()

# 10
H("10.", "Go-to-market")
for t in F.S10: P_(t)
B_(B.FIZYKA_MARKETINGU)
h2("Fazy wejścia i mierniki")
TAB(B.WEJSCIE)
doc.add_page_break()

# 11
H("11.", "Konkurencja i system publiczny")
for t in F.S11: P_(t)
h2("Siedem kategorii konkurencji")
TAB(B.KONKURENCJA)
B_(B.KONKURENCJA_KALIBRACJA)
doc.add_page_break()

# 12
H("12.", "Fosa, własność intelektualna i kontrola")
for t in F.S12: P_(t)
h2("Sześć elementów fosy")
TAB(B.FOSA)
h2("Mechanizmy kontroli — konstrukcja, koszt, siła")
TAB(B.KONTROLA8)
h2("Plan utrzymania kontroli w czasie")
TAB(B.KONTROLA_PLAN)
B_(B.KONTROLA_WARUNEK)
h2("Zapis statutowy decydujący o trwałości")
P_(B.STATUT)
doc.add_page_break()

# 13
H("13.", "Technologia i operacje")
for t in F.S13: P_(t)
h2("Tanie warianty wobec budowy własnej")
TAB(B.WARIANTY_TANIE)
h2("Ekonomika stacji")
TAB(B.STACJA_EKONOMIKA)
P_(B.STACJA_WNIOSEK)
h2("Marża sprzętowa")
TAB(B.MARZA_SPRZET)
h2("Rynek dostawców — co da się kupić, co da się napisać samemu")
TAB([[x[0], x[1], " · ".join(x[2])[:220], " · ".join(x[3])[:220], str(x[4])[:260], str(x[5]), str(x[7])]
     for x in RY.POZYCJE],
    ["Obszar", "Producenci", "Agregatorzy", "OEM / ODM", "Czy da się samemu",
     "Publiczna specyfikacja", "Wchodzimy"])
P_(str(RY.TEST[0]) + "  " + str(RY.TEST[1]) + "  " + str(RY.TEST[2]))
h2("Agregatorzy danych — ocena dostawców")
TAB([[x[0], x[1], x[2], x[3], str(x[4]), str(x[5]), str(x[8])[:300]] for x in RY.AGREGATORY],
    ["Dostawca", "Grupa", "Reżim", "Cennik", "Pokrycie", "Ocena", "Uwaga"])
h2("Funkcje modułu agregacji a dostępność gotowego rozwiązania")
TAB([list(x) for x in RY.A1_FUNKCJE],
    ["Kod", "Funkcja", "Kategoria", "Dostawcy", "Czy kupić", "Decyzja"])
h2("Wspólny stos technologiczny")
TAB([list(x) for x in MO.STOS_WSPOLNY], ["Warstwa", "Rozstrzygnięcie"])
doc.add_page_break()

# 14
H("14.", "Zespół i struktura podmiotu")
for t in F.S14: P_(t)
h2("Zespół — role i stan obsadzenia")
TAB(B.ZESPOL)
h2("Struktura podmiotów")
TAB(B.STRUKTURA)
h2("Cel konstrukcji własnościowej")
P_(MD.STRUKTURA_CEL)
TAB([list(x) for x in MD.PODMIOTY])
h2("Modele wykonania ekosystemu")
TAB([list(x) for x in MD.MODELE])
P_(MD.MODELE_WNIOSEK)
h2("Fundusz badawczy")
P_(MD.FUNDUSZ)
h2("Źródła kontroli")
TAB([list(x) for x in MD.KONTROLA_ZRODLA])
doc.add_page_break()

# 15
H("15.", "Finanse i ekonomika")
for t in F.S15: P_(t)
h2("Struktura kosztów")
TAB(B.KOSZTY_STRUKTURA)
h2("Błędy wcześniejszych modeli kosztowych i ich skala")
TAB(B.BLEDY_KOSZTOWE)
h2("Budżet dziewięćdziesięciu dni")
TAB(B.BUDZET90)
h2("Koszt startowy komponentów — co realnie płacimy miesięcznie")
TAB([list(map(str, x)) for x in KP.EKONOMIA],
    ["Klasa", "Rozwiązanie startowe", "Model kosztu", "Koszt mies. (zł)",
     "Funkcji zależnych", "Podstawa"])
P_("Zestawienie pokazuje, dlaczego koszt technologiczny w pierwszym roku jest niski, a ryzyko "
   "wysokie: większość klas startuje na rozwiązaniach bezpłatnych albo tanich, ale liczba "
   "funkcji zależnych od pojedynczej klasy sięga kilkudziesięciu. Wzrost ceny u jednego "
   "dostawcy przekłada się wtedy na kilkadziesiąt funkcji naraz — stąd wymóg progu wyjścia "
   "zapisanego liczbowo dla każdej klasy.")
h2("Dźwignie niepieniężne")
TAB(B.DZWIGNIA)
B_(B.DZWIGNIA_WARUNEK)
h2("Prognoza wycofana — co zawierała i dlaczego nie obowiązuje")
TAB(AN.FINANSE)
P_("Powyższa prognoza pochodzi z materiału inwestorskiego sprzed audytu i jest w tym "
   "dokumencie pokazana wyłącznie po to, żeby wskazać, co zostało wycofane i dlaczego. "
   "Model przychodowy stojący za tymi liczbami zakładał subskrypcję konsumencką jako oś "
   "przychodu oraz strukturę kosztów bez wynagrodzeń. Oba założenia zostały zmienione, "
   "więc liczby nie mają już podstawy.")
doc.add_page_break()

# 16
H("16.", "Finansowanie")
for t in F.S16: P_(t)
h2("Kolejność źródeł finansowania")
TAB(B.ZRODLA_FINANSOWANIA)
h2("Budżet horyzontu zerowego")
TAB(R.BUDZET)
if hasattr(R, 'BUDZET_NOTA'): P_(R.BUDZET_NOTA)
h2("Arytmetyka projektów badawczych")
P_(B.MOONSHOT_ARYTMETYKA)
doc.add_page_break()

# 17
H("17.", "Ryzyka")
for t in F.S17: P_(t)
TAB(B.RYZYKA)
doc.add_page_break()

# 18
H("18.", "Kamienie milowe i wskaźniki")
for t in F.S18: P_(t)
h2("Bramki decyzyjne")
TAB(B.BRAMKI)
h2("Pięć torów równoległych")
TAB(R.TORY)
h2("Horyzont zerowy — zadania, właściciele, terminy")
TAB(R.H0)
h2("Czego nie robimy")
TAB(R.NIE_ROBIMY)
doc.add_page_break()

# 19
H("19.", "Czego plan nie obiecuje")
for t in F.S19: P_(t)
LIST(B.NIE_OBIECUJEMY)
h2("Pozycje usunięte z planu i z materiałów zewnętrznych")
P_(B.USUNIETE)
doc.add_page_break()

# 20
H("20.", "Indeks źródeł i proweniencja")
for t in F.S20: P_(t)
TAB(F.ZRODLA_ZEW)
doc.add_page_break()

# 21
H("21.", "Audyt źródeł i korekty")
for t in F.S21: P_(t)
h2("Korekty wobec treści źródłowej")
TAB(B.KOREKTY)
h2("Zmiany wobec wcześniejszych wersji roadmapy")
TAB(R.ZMIANY)
h2("Ustalenia audytu materiału inwestorskiego")
for x in AN.BLEDY:
    doc.add_heading("[%s]  %s" % (x[0], x[1]), 3)
    for etk, tr in [("Co deklarowano", x[2]), ("Na czym polega problem", x[3]),
                    ("Rozwiązanie", x[4])]:
        pp = doc.add_paragraph(); rr = pp.add_run(etk + ". "); rr.bold = True
        pp.add_run(str(tr))
h2("Co się broni bez zmian")
TAB([list(x) for x in AN.DOBRZE], ["Element", "Dlaczego się broni"])
doc.add_page_break()

# ---------------------------------------------------------- ZAŁĄCZNIK A
doc.add_heading("ZAŁĄCZNIK A — MODUŁY W UJĘCIU BIZNESOWYM", 1)
P_("Czterdzieści trzy moduły ekosystemu w ujęciu, które interesuje płatnika: po co moduł "
   "istnieje, kto z niego korzysta, na czym startuje, kiedy przechodzi na rozwiązanie własne "
   "i w którym kubełku regulacyjnym się znajduje.")
TAB([[k, str(v[13]), str(v[14]), str(v[15])] for k, v in sorted(MO.M.items())],
    ["Moduł", "Kubełek", "Priorytet", "Właściciel"])
for k, v in sorted(MO.M.items()):
    doc.add_heading("%s  %s" % (k, v[0]), 2)
    for etk, tr in [("Cel", v[1]), ("Problem", v[2]), ("Kto korzysta", v[3]),
                    ("Wejście", v[4]), ("Wyjście", v[5]),
                    ("Rozwiązanie startowe", v[10]), ("Próg przejścia na własne", v[11]),
                    ("Własne IP", v[9])]:
        pp = doc.add_paragraph(); rr = pp.add_run(etk + ". "); rr.bold = True
        pp.add_run(str(tr))
    TAB([[str(v[13]), str(v[14]), str(v[15]),
          (v[12][0] if isinstance(v[12], tuple) else str(v[12]))]],
        ["Kubełek", "Priorytet", "Właściciel", "Wymienialny"])
h2("Kubełki regulacyjne modułów")
TAB([[k] + list(v) for k, v in MO.KUBELKI.items()],
    ["Kubełek", "Nazwa", "Zakres", "Konsekwencja budżetowa"])
doc.add_page_break()

# ---------------------------------------------------------- ZAŁĄCZNIK B
doc.add_heading("ZAŁĄCZNIK B — REJESTR FUNKCJI W UJĘCIU MONETYZACYJNYM", 1)
P_("Pełny rejestr operacyjny w przekroju, który interesuje inwestora: która funkcja zarabia, "
   "w jakim kanale, na jakim etapie i w jakiej warstwie regulacyjnej. Kolumna „waga dla "
   "ekosystemu” mówi, czy funkcja jest warunkiem działania innych, nawet jeżeli sama nie "
   "generuje przychodu.")
TAB([[v['kod'], v['nazwa'][:80], v['produkt'].replace('Eternal ', ''),
      v['etap'], v['warstwa'], str(v.get('kanal', '—'))[:34],
      str(v.get('zarabia', '—'))[:18], str(v.get('waga_eko', '—'))[:12]]
     for v in sorted(RJ.R.values(), key=lambda x: (x['produkt'], x['kod']))],
    ["Kod", "Nazwa", "Produkt", "Etap", "Warstwa", "Kanał", "Zarabia", "Waga eko"])
doc.add_page_break()

# ---------------------------------------------------------- ZAŁĄCZNIK C
doc.add_heading("ZAŁĄCZNIK C — PYTANIA ROZSTRZYGNIĘTE", 1)
P_("Pytania, które wracają w każdej rozmowie inwestorskiej i technicznej, wraz z "
   "rozstrzygnięciem przyjętym w tym planie. Rozstrzygnięcie zapisane raz jest tańsze niż "
   "rozstrzygnięcie improwizowane przy stole.")
for x in MO.ODPOWIEDZI:
    doc.add_heading(str(x[0]), 3)
    B_(str(x[1]))
    for extra in x[2:]:
        P_(str(extra))

doc.save(OUT)
ch = sum(len(p.text) for p in doc.paragraphs) + sum(
    len(c.text) for t in doc.tables for r in t.rows for c in r.cells)
print('%s -> %d B, %d akapitow, %d tabel, ~%d stron' % (
    OUT, os.path.getsize(OUT), len(doc.paragraphs), len(doc.tables), round(ch / 1800)))
