# -*- coding: utf-8 -*-
"""Product Requirements Document — 43 moduly z kartami PRD i stosem technologicznym."""
import json
import os
import sys
import re
import collections
import datetime

sys.path.insert(0, os.path.dirname(os.path.abspath(__file__)))
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from mkdocx import setup, toc
from dane_komponenty import K, WARSTWA, SZCZEBEL, SKLADOWE
from dane_moduly import M, KUBELKI, STOS_WSPOLNY, STOS_ROZNICE, ODPOWIEDZI, WAGI

R = json.load(open('build/KOMPONENTY.json'))
FUN = collections.defaultdict(list)
for r in R:
    FUN[r['modul']].append(r)
TODAY = datetime.date.today().strftime('%d.%m.%Y')
OUT = '/home/user/Eternal-Lite-App/out/ETERNAL_PRD.docx'

KOL_KUB = {'W': 'D7F0DD', 'W>M': 'E3EDFA', 'M>W': 'F3E3C3', 'M': 'F8D7DA'}


def cien(cell, hexcol):
    sh = OxmlElement('w:shd')
    sh.set(qn('w:val'), 'clear')
    sh.set(qn('w:fill'), hexcol)
    cell._tc.get_or_add_tcPr().append(sh)


def bogaty(p, txt, size=9.5):
    for i, cz in enumerate(str(txt).split('**')):
        if cz:
            r = p.add_run(cz)
            r.font.size = Pt(size)
            r.bold = bool(i % 2)


def tab(doc, rows, szer, naglowek='1B3A6B', kol_fill=None):
    t = doc.add_table(rows=0, cols=len(rows[0]))
    t.style = 'Table Grid'
    for ri, row in enumerate(rows):
        cells = t.add_row().cells
        for ci, val in enumerate(row):
            cells[ci].text = ''
            bogaty(cells[ci].paragraphs[0], val, 8.5)
            if ri == 0:
                for r_ in cells[ci].paragraphs[0].runs:
                    r_.bold = True
                    r_.font.color.rgb = RGBColor.from_string('FFFFFF')
                cien(cells[ci], naglowek)
            elif kol_fill and ci == kol_fill[0] and str(val) in kol_fill[1]:
                cien(cells[ci], kol_fill[1][str(val)])
            cells[ci].width = Cm(szer[ci])
    return t


def pole(doc, etykieta, tresc, akcent=False):
    p = doc.add_paragraph()
    r = p.add_run(etykieta.upper())
    r.font.size = Pt(7.5)
    r.bold = True
    r.font.color.rgb = RGBColor.from_string('B8431F' if akcent else '5D6B8A')
    p2 = doc.add_paragraph()
    p2.paragraph_format.left_indent = Cm(0.4)
    p2.paragraph_format.space_after = Pt(6)
    bogaty(p2, tresc)


doc = Document()
setup(doc)
for txt, sz, bold, col in [
        ('ETERNAL LIFE', 26, True, 'B8431F'),
        ('Product Requirements Document', 16, True, '1B3A6B'),
        ('43 moduły · %d funkcji · karta PRD i stos technologiczny dla każdego' % len(R),
         12, False, None)]:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(txt)
    r.font.size = Pt(sz)
    r.bold = bold
    if col:
        r.font.color.rgb = RGBColor.from_string(col)
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('Stan na %s' % TODAY)
r.font.size = Pt(10)
doc.add_page_break()

doc.add_heading('Jak czytać ten dokument', 1)
for t in [
 'Każdy moduł ma dwie karty. **Karta PRD** odpowiada na pytanie, co moduł robi i dla kogo — '
 'cel, problem, użytkownik, wejście, wyjście, integracje, uprawnienia, regulacje, kryteria '
 'akceptacji. **Karta budowy** odpowiada na pytanie, czy to kupić czy zbudować — kandydat '
 'na cały moduł, jego pokrycie, kontrola, którą zostawia, i próg opłacalności własnego.',
 'Stos technologiczny jest wspólny dla całego produktu i podany raz, na początku. '
 'Przy modułach, które mają własne API zewnętrzne albo własne punkty końcowe, '
 'podane są tylko różnice — powtarzanie tego samego stosu przy czterdziestu trzech '
 'modułach nie niesie informacji.',
 '**Priorytet** P0 oznacza moduł, bez którego produkt nie działa; P1 — moduł, który '
 'decyduje o przychodzie albo o zgodności; P2 — moduł etapu docelowego.',
 '**Ewolucja wellness → medyczne** jest podana w czterech kubełkach, opisanych niżej. '
 'To jest odpowiedź na pytanie, które funkcje wolno wydać bez certyfikacji, a które '
 'tylko wyglądają, jakby wolno było.',
]:
    p = doc.add_paragraph()
    bogaty(p, t, 10)

doc.add_heading('Cztery kubełki: wellness czy medyczne', 2)
kub = collections.Counter(v[13] for v in M.values())
tab(doc, [['Kubełek', 'Co znaczy', 'Modułów', 'Co z tym robić']]
    + [[k, v[0] + ' — ' + v[1], str(kub.get(k, 0)), v[2]] for k, v in KUBELKI.items()],
    [2.2, 7.0, 1.8, 6.0], kol_fill=(0, KOL_KUB))

doc.add_heading('Model kontroli', 2)
p = doc.add_paragraph()
bogaty(p, 'Kontrola jest liczona jawnym wzorem, żeby dało się ją sprawdzić i żeby spór '
 'dotyczył wag, a nie wyniku: **kontrola = 0,40 × (szczebel/5) + 0,25 × dane + '
 '0,20 × wymienialność + 0,15 × wniosek**. Szczebel to pozycja 1–5 wg Master 5.4. '
 'Dane — czy mamy kopię działającą bez dostawcy. Wymienialność — czy adapter izoluje rdzeń. '
 'Wniosek — czy końcowa ocena należy do nas. Wagi są arbitralne, ale jawne: mówią, '
 'co uznajemy za kontrolę.', 10)

doc.add_page_break()
doc.add_heading('Spis treści', 1)
toc(doc)
doc.add_page_break()

doc.add_heading('Stos technologiczny — wspólny dla produktu', 1)
tab(doc, [['Warstwa', 'Rozwiązanie']] + [list(x) for x in STOS_WSPOLNY], [4.0, 13.0])

# --- karty modulow ---------------------------------------------------------
GRUPY = [('APLIKACJA ETERNAL', [k for k in M if k.startswith('A')]),
         ('ETERNAL STATION', [k for k in M if k.startswith('S')]),
         ('ETERNAL CAPSULE', [k for k in M if k.startswith('C')]),
         ('ETERNAL DIGITAL TWIN', [k for k in M if k.startswith('D')]),
         ('ETERNAL MATRIX', [k for k in M if k.startswith('X')])]


def klucz(k):
    return int(re.match(r'[A-Z]+(\d+)', k).group(1))


for tytul, kody in GRUPY:
    doc.add_page_break()
    doc.add_heading(tytul, 1)
    for kod in sorted(kody, key=klucz):
        v = M[kod]
        (nazwa, cel, problem, uzyt, inp, outp, kand, pokr, kontr, nasze,
         oss, wlasne, adapter, kubelek, prio, owner) = v
        fs = FUN.get(kod, [])
        warstwy = collections.Counter(f['warstwa'] for f in fs)
        klasy = sorted({f['klasa'] for f in fs} | {f['klasa'] for f in fs})
        wsp = sorted({x for f in fs for x in f['wspierajace'].split(' + ') if x})
        md = 'TAK' if warstwy.get('C') else ('GRANICA' if warstwy.get('B') else 'NIE')

        doc.add_heading('%s. %s' % (kod, nazwa), 2)
        tab(doc, [['Funkcji', 'Priorytet', 'Owner', 'Medical device', 'Kubełek',
                   'Warstwy zgodności'],
                  [str(len(fs)), prio, owner,
                   md + (' — %d funkcji w warstwie C' % warstwy['C'] if warstwy.get('C') else ''),
                   '%s — %s' % (kubelek, KUBELKI[kubelek][0]),
                   ' · '.join('%s: %d' % (w, n) for w, n in sorted(warstwy.items()))]],
            [2.0, 2.0, 2.2, 4.6, 3.6, 2.6], kol_fill=(4, {}))

        pole(doc, 'Cel', cel)
        pole(doc, 'Problem', problem)
        pole(doc, 'Użytkownik', uzyt)
        pole(doc, 'Opis funkcji', 'Moduł obejmuje %d funkcji: %s.'
             % (len(fs), ', '.join('%s %s' % (f['kod'], f['nazwa'][:44]) for f in fs[:12])
                + (' i dalsze' if len(fs) > 12 else '')))
        pole(doc, 'Input', inp)
        pole(doc, 'Output', outp)
        pole(doc, 'Przebieg użytkownika',
             'Wejście przez %s → normalizacja do Eternal Standard → zapis z proweniencją '
             '→ prezentacja w interfejsie modułu → eksport albo działanie następcze.'
             % (inp[:1].lower() + inp[1:]))
        pole(doc, 'Integracje',
             'Klasy komponentów: %s%s. Składowe: %s.'
             % (', '.join('%s (%s)' % (k_, K[k_][0]) for k_ in klasy),
                '; wspierające: ' + ', '.join(wsp) if wsp else '',
                fs[0]['skladowe'] if fs else '—'))
        pole(doc, 'API',
             '; '.join('%s: %s' % (a, b) for a, b in STOS_ROZNICE.get(kod, []))
             or 'Bez własnych integracji zewnętrznych — korzysta z punktów końcowych '
                'wspólnych i z danych modułów źródłowych.')
        pole(doc, 'Dane', 'Zasoby FHIR R4B mapowane na Eternal Standard, z proweniencją '
             '(źródło, czas, ścieżka). Retencja wg polityki modułu K3; usunięcie na żądanie '
             'obejmuje kopie w kolejkach i w kopii zapasowej.')
        pole(doc, 'Uprawnienia', 'Zakres zgody per funkcja, nie per aplikacja. Dostęp lekarza '
             'czasowy i odwoływalny (A24). Każde odczytanie zapisywane w dzienniku widocznym '
             'dla użytkownika (A18).')
        pole(doc, 'Bezpieczeństwo', 'Rdzeń nie woła dostawcy bezpośrednio. Klucze po naszej '
             'stronie. Umowa powierzenia przed pierwszym wywołaniem zewnętrznym. '
             'Tryb degradacji przy niedostępności dostawcy.')
        pole(doc, 'Regulacje',
             'RODO art. 9 (dane szczególnej kategorii) i art. 32. %s%s'
             % ('MDR — moduł zawiera %d funkcji w warstwie C, wymaga dossier klasy IIa '
                'albo modelu proxy. ' % warstwy['C'] if warstwy.get('C') else
                'Poza MDR przy obecnym przeznaczeniu. ',
                'AI Act — oznaczanie treści generowanej od 2.08.2026. '
                if kod in ('A6', 'A12', 'A17', 'A19', 'D2') else ''))
        p = doc.add_paragraph()
        r = p.add_run('CZY FUNKCJA JEST MEDICAL DEVICE:  ')
        r.font.size = Pt(7.5)
        r.bold = True
        r.font.color.rgb = RGBColor.from_string('5D6B8A')
        r2 = p.add_run(md)
        r2.font.size = Pt(9.5)
        r2.bold = True
        r2.font.color.rgb = RGBColor.from_string(
            'B8431F' if md == 'TAK' else ('B07419' if md == 'GRANICA' else '2E7D32'))
        pole(doc, 'Kryteria akceptacji',
             'Wszystkie funkcje modułu działają w trybie degradacji przy odciętym dostawcy. '
             'Adapter przechodzi test kontraktu niezależnie od dostawcy. '
             'Każde wywołanie zewnętrzne ma wpis w dzienniku. '
             'Zgoda na funkcję jest odrębna i odwoływalna.'
             + (' Deklaracja przeznaczenia zapisana i podpisana przed wydaniem.'
                if warstwy.get('C') or warstwy.get('B') else ''))
        pole(doc, 'Status', 'Etapy funkcji w module: '
             + ' · '.join('%s: %d' % (e, n) for e, n
                          in collections.Counter(f['etap'] for f in fs).most_common()))

        doc.add_heading('%s — karta budowy: kupić czy zbudować' % kod, 3)
        tab(doc, [['Pozycja', 'Ustalenie'],
                  ['Kandydat na cały moduł', kand],
                  ['Pokrycie funkcji modułu', '%d%%' % pokr],
                  ['Kontrola przy tym wariancie', '%d%% wg modelu kontroli' % kontr],
                  ['Co zostaje nasze bezwzględnie', nasze],
                  ['Alternatywa open source', oss],
                  ['Kiedy budujemy własne', wlasne],
                  ['Adapter wymagany', '%s — %s' % adapter]],
            [4.5, 12.5], naglowek='B8431F')
        if STOS_ROZNICE.get(kod):
            doc.add_paragraph()
            tab(doc, [['Stos — różnice wobec wspólnego', 'Rozwiązanie']]
                + [list(x) for x in STOS_ROZNICE[kod]], [4.5, 12.5])

doc.save(OUT)
zn = (sum(len(p.text) for p in doc.paragraphs)
      + sum(len(c.text) for t in doc.tables for r in t.rows for c in r.cells))
print('%s -> %d B, %d modulow, %d tabel, ~%d stron'
      % (OUT, os.path.getsize(OUT), len(M), len(doc.tables), max(1, round(zn / 2400))))
