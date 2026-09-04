# -*- coding: utf-8 -*-
"""Dokument CEO — 26 sekcji, produkty, modele, struktura podmiotu."""
import os, sys
sys.path.insert(0, os.path.dirname(os.path.abspath(__file__)))
from mkdocx import *
from docx import Document
import dane_ceo as C, dane_produkty as P, dane_modele as MD, rejestr as RJ

doc = Document(); setup(doc)
for t, sz, b in [("ETERNAL ECOSYSTEM", 26, True),
                 ("DOKUMENT ZARZĄDCZY", 16, True),
                 ("Dwadzieścia sześć sekcji · sześć produktów · modele wykonania · "
                  "struktura podmiotu", 11, False)]:
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(t); r.font.size = Pt(sz); r.bold = b
    if sz >= 16: r.font.color.rgb = RGBColor.from_string('1F3864')
p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("Podstawa: pełny odczyt korpusu — 159 plików, 28 618 387 znaków\n"
              "Rejestr: %d funkcji · Stan na %s" % (len(RJ.R), TODAY)); r.font.size = Pt(10)
doc.add_page_break()

doc.add_heading("Nota metodyczna", 1)
for t in [
 "Dokument jest wyciągiem zarządczym z korpusu, nie jego streszczeniem. Każda liczba ma "
 "źródło w rejestrze funkcji albo w ustaleniach pełnego odczytu; tam, gdzie źródła są "
 "sprzeczne, podana jest wartość obowiązująca wraz z informacją, co zastępuje.",
 "Zasada konstrukcyjna produktów jest inna niż we wcześniejszych wersjach. Produkt nie jest "
 "modułem ani zbiorem modułów. Produkt to pięć albo sześć funkcji z rejestru, dobranych tak, "
 "że razem robią jedną rzecz, której żadna z nich nie robi osobno. Żaden z sześciu produktów "
 "nie wprowadza funkcji, której nie ma w rejestrze.",
 "Warstwa wyłączona przez sekcję 38 specyfikacji — sterowanie zachowaniem ludzi, wpływ na "
 "decyzje wyborcze, oddziaływanie podprogowe, masowa implantacja — nie wchodzi do tego "
 "dokumentu i nie jest w nim rozwijana.",
]:
    doc.add_paragraph(t)
doc.add_page_break()
doc.add_heading("Spis treści", 1); toc(doc); doc.add_page_break()

def H(n, t): doc.add_heading("%s  %s" % (n, t), 1)
def h2(t): doc.add_heading(t, 2)
def P_(t): doc.add_paragraph(t)
def B_(t):
    p = doc.add_paragraph(); r = p.add_run(t); r.bold = True

H("00.", "Executive Summary")
for t in C.S00: P_(t)

H("01.", "Wizja i misja")
for t in C.S01: P_(t)

H("02.", "Cele strategiczne"); add_table(doc, C.S02)

H("03.", "Problem i potrzeba rynku"); add_table(doc, C.S03); B_(C.S03_NIE)

H("04.", "Grupy użytkowników"); add_table(doc, C.S04); B_(C.S04_ZASADA)

H("05.", "Model biznesowy"); add_table(doc, C.S05); B_(C.S05_ZASADA)

H("06.", "Ekosystem Eternal"); add_table(doc, C.S06)
P_("Modele wykonania ekosystemu — nie alternatywy do wyboru, tylko warstwy włączane "
   "w kolejności. Szczegóły w sekcji 19 i w załączniku o modelach.")

H("07.", "Produkty")
P_("Sześć produktów, każdy złożony z pięciu do sześciu funkcji rejestru. Metoda doboru:")
add_table(doc, P.METODA)
for pr in P.PRODUKTY:
    h2("%s  %s" % (pr['kod'], pr['nazwa']))
    B_(pr['claim'])
    rows = [["Funkcja", "Nazwa", "Etap", "Warstwa"]]
    for f in pr['funkcje']:
        v = RJ.R[f]
        rows.append([f, v['nazwa'], v['etap'], v['warstwa']])
    add_table(doc, rows)
    for et, tr in [("Niezastępowalność", pr['niezast']), ("Automatyzm", pr['automat']),
                   ("Rozwój", pr['rozwoj']), ("Personalizacja", pr['person']),
                   ("Odbiorca", pr['odbiorca']), ("Monetyzacja", pr['monetyzacja']),
                   ("Samodzielność", pr['sam']), ("Etap", pr['etap']),
                   ("Ryzyko", pr['ryzyko'])]:
        p = doc.add_paragraph(); r = p.add_run(et + ". "); r.bold = True; p.add_run(tr)
P_(P.PRODUKT_SIODMY)
h2("Kiedy moduł staje się osobnym produktem")
add_table(doc, P.MODUL_NA_PRODUKT); B_(P.MODUL_NA_PRODUKT_WNIOSEK)
h2("Produkty wielobranżowe z tego samego rdzenia")
P_(P.RDZEN); add_table(doc, P.BRANZE); B_(P.BRANZE_ZASADA)
h2("Monetyzacja produktów"); add_table(doc, P.MONETYZACJA); B_(P.MONETYZACJA_ZASADA)
h2("Dobór funkcji pod klienta"); add_table(doc, P.DOBOR); B_(P.DOBOR_ZASADA)
h2("Kiedy budujemy własne, a kiedy kupujemy")
add_table(doc, P.BUILD_BUY); add_table(doc, P.ALTERNATYWY)

H("08.", "Moduły")
P_("Moduł jest porządkiem katalogowym, nie jednostką sprzedaży. Rozkład rejestru na moduły:")
import collections
mods = collections.Counter((RJ.R[k]['produkt'], RJ.R[k]['modul']) for k in RJ.R)
rows = [["Produkt", "Moduł", "Funkcji"]]
for (prod, mod), n in sorted(mods.items()):
    rows.append([prod, mod, str(n)])
add_table(doc, rows)

H("09.", "Funkcje"); add_table(doc, C.S09_LICZBY)
P_("Pełne karty wszystkich %d funkcji w szablonie osiemnastopolowym: "
   "ETERNAL_KARTY_FUNKCJI.docx." % len(RJ.R))

H("10.", "Priorytety i roadmapa")
st = collections.Counter(RJ.R[k]['etap'] for k in RJ.R)
add_table(doc, [["Etap", "Funkcji"]] + [[k, str(v)] for k, v in st.most_common()])
add_table(doc, C.S24)

H("11.", "Regulacje i compliance"); add_table(doc, C.S11); B_(C.S11_REGULA)

H("12.", "IP i własność technologiczna"); add_table(doc, C.S12); B_(C.S12_ZASADA)

H("13.", "Dane")
for t in C.S13: P_(t)

H("14.", "AI"); add_table(doc, C.S14)

H("15.", "Cyberbezpieczeństwo"); add_table(doc, C.S15)

H("16.", "Integracje"); add_table(doc, C.S16); B_(C.S16_LUKA)

H("17.", "Hardware i software")
for t in C.S17: P_(t)

H("18.", "Partnerstwa, OEM, API, SDK"); add_table(doc, C.S18)

H("19.", "Model operacyjny"); add_table(doc, C.S19); B_(C.S19_ZASADA)
h2("Struktura podmiotu — model badawczo-biznesowy")
P_(MD.STRUKTURA_CEL); add_table(doc, MD.PODMIOTY); B_(MD.FUNDUSZ)
add_table(doc, MD.KONTROLA_ZRODLA)
h2("Modele wykonania ekosystemu")
for m in MD.MODELE:
    doc.add_heading("%s  %s" % (m['kod'], m['nazwa']), 3)
    add_table(doc, [["Wymiar", "Ustalenie"],
                    ["Istota", m['istota']], ["Koszt", m['koszt']],
                    ["Kontrola", m['kontrola']], ["Czas do przychodu", m['czas']],
                    ["Ryzyko regulacyjne", m['ryzyko_reg']],
                    ["Kiedy stosujemy", m['kiedy']], ["Warunek wyjścia", m['wyjscie']]])
B_(MD.MODELE_WNIOSEK); P_(MD.PORTFEL_ZASADA)

H("20.", "Finanse"); add_table(doc, C.S20); B_(C.S20_KOSZTY)

H("21.", "KPI"); add_table(doc, C.S21)

H("22.", "Ryzyka"); add_table(doc, C.S22)

H("23.", "Decyzje strategiczne"); add_table(doc, C.S23)

H("24.", "Roadmapa 1–3–5–10 lat"); add_table(doc, C.S24)

H("25.", "Załączniki"); add_table(doc, C.S25)

out = '/home/user/Eternal-Lite-App/out/ETERNAL_CEO.docx'
os.makedirs(os.path.dirname(out), exist_ok=True); doc.save(out)
ch = sum(len(p.text) for p in doc.paragraphs) + sum(
    len(c.text) for t in doc.tables for r in t.rows for c in r.cells)
print('%s -> %d B, %d akapitow, %d tabel, ~%d stron' % (
    out, os.path.getsize(out), len(doc.paragraphs), len(doc.tables), round(ch / 1800)))
