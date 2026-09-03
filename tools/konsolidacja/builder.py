# -*- coding: utf-8 -*-
import json, os, sys
sys.path.insert(0, os.path.dirname(os.path.abspath(__file__)))
from mkdocx import *
from docx import Document
from dane_ustalenia import U, KAT
from dane_pliki import P as UPL
import dane_odczyt as OD
import wyklucz
import dane_odczyt_bp as OB


def _ustalenia(doc, sec):
    """Czesc 0 — ustalenia z plikow bez kodow funkcji, wlasciwe dla tej sekcji."""
    GR = {'S': ('C', 'K', 'T', 'D', 'M', 'S', 'L'), 'B': ('P', 'E', 'S', 'K', 'M', 'L'),
          'R': ('S', 'T', 'D', 'L'), 'P': ('P', 'E', 'C', 'L')}
    poz = [u for u in U if u[0][0] in GR.get(sec, ())]
    if not poz:
        return
    doc.add_heading('CZĘŚĆ 0 — USTALENIA Z PLIKÓW BEZ KODÓW FUNKCJI', 1)
    doc.add_paragraph(
        'Siedemdziesiąt cztery pliki korpusu nie zawierają kodów funkcji i przez to nie '
        'występowały w rejestrze, z którego budowano dokumenty analityczne. Ich treść była '
        'w częściach poniżej — nie była w żadnym wniosku. Ta część zbiera %d ustaleń '
        'właściwych dla tej sekcji, z numerem pliku źródłowego przy każdym. '
        'Pełny rejestr %d ustaleń: ETERNAL_USTALENIA_KORPUSU.docx.' % (len(poz), len(U)))
    rows = [['Kod', 'Kategoria', 'Ustalenie', 'Co zmienia', 'Pliki']]
    for kod, kat, tyt, ust, zm, pl in poz:
        rows.append([kod, KAT[kat][0], tyt + ' — ' + ust, zm, pl])
    add_table(doc, rows)
    doc.add_page_break()

    W = {'KOR': 'KOREKTA', 'ROZ': 'ROZSTRZYGNIĘCIE', 'NOW': 'NOWE', 'RYZ': 'RYZYKO',
         'POT': 'POTWIERDZENIE'}
    mine = [i for i, (sc, st, ro) in sorted(M.items()) if sec in sc.split(',') and i in UPL]
    if mine:
        doc.add_heading('CZĘŚĆ 0B — USTALENIE Z KAŻDEGO PLIKU TEJ SEKCJI', 1)
        doc.add_paragraph(
            'Przejście po kolei przez wszystkie %d plików przypisanych do tej sekcji, '
            'w paczkach po dziesięć. Każdy plik ma jeden wpis: co z niego wynika i jaką '
            'ma wagę. Pełny rejestr wszystkich 159 plików: '
            'ETERNAL_USTALENIA_PER_PLIK.docx.' % len(mine))
        r2 = [['#', 'Plik', 'Waga', 'Ustalenie z tego pliku']]
        for i in mine:
            u, wg = UPL[i]
            r2.append([str(i), INV[i]['name'].replace('.txt', '')[:52], W[wg], u])
        add_table(doc, r2)
        doc.add_page_break()


def _odczyt(doc, sec):
    """Czesc 0C — ustalenia z pelnego odczytu korpusu. Tylko sekcja SPECYFIKACJA."""
    if sec != 'S':
        return
    H1 = doc.add_heading
    H1('CZĘŚĆ 0C — USTALENIA Z PEŁNEGO ODCZYTU KORPUSU', 1)
    doc.add_paragraph(
        'Ta część powstała z odczytu całej treści 159 plików korpusu — 28 618 387 znaków '
        'surowo, 13 020 154 po usunięciu duplikatów i treści powtarzającej się jeden do '
        'jednego. Nie powiela bloków źródłowych z części I–X: zawiera wyłącznie ustalenia, '
        'które powstały z zestawienia wielu plików albo prostują treść źródłową. Dziennik '
        'odczytu: tools/konsolidacja/odczyt/USTALENIA_ODCZYT.md.')

    H1('0C.1  Hierarchia dokumentów obowiązujących', 2)
    doc.add_paragraph(
        'Korpus zawiera kilka pokoleń tych samych dokumentów. Poniższa tabela rozstrzyga, '
        'która wersja obowiązuje w każdym zakresie i co zostało przez nią zastąpione.')
    add_table(doc, OD.HIERARCHIA)

    H1('0C.2  Dziesięć różnych liczb funkcji — rozstrzygnięcie', 2)
    doc.add_paragraph(
        'W korpusie występuje dziesięć różnych liczb funkcji. Nie są sprzeczne — opisują '
        'różne zakresy albo różne etapy tego samego rejestru. Liczba użyta bez wskazania '
        'ujęcia jest nieinformatywna.')
    add_table(doc, OD.LICZBY)

    H1('0C.3  Cztery statusy regulacyjne i granica wyrobu', 2)
    doc.add_paragraph(
        'Podział na „certyfikowane i niecertyfikowane” jest źródłem błędów, bo sugeruje, '
        'że brak MDR oznacza brak regulacji. Obowiązują cztery statusy.')
    add_table(doc, OD.STATUSY)
    p = doc.add_paragraph(); r = p.add_run('Reguła robocza granicy: ' + OD.REGULA_GRANICY)
    r.bold = True

    H1('0C.4  Model ewolucji wellness → wyrób medyczny', 2)
    doc.add_paragraph(
        'Ewolucja nie polega na dodaniu kodu. Polega na zmianie jednego zdania '
        'w przeznaczeniu i poniesieniu kosztu dossier. Ta sama funkcja techniczna może '
        'mieć dwa reżimy.')
    add_table(doc, OD.EWOLUCJA)
    doc.add_paragraph(
        'Alternatywa dla etapów 4–6: proxy do cudzego wyrobu z CE. Działa dla funkcji '
        'jednorodnych — jedno wejście, jeden wynik. Nie działa dla funkcji łączących dane '
        'z wielu źródeł. ' + OD.REGULA_PROXY)

    H1('0C.5  Czterdzieści pięć reguł: kiedy funkcja staje się wyrobem', 2)
    doc.add_paragraph(
        'Lista deduplikowana ze 183 kart funkcji. To jest lista kontrolna do każdego '
        'przeglądu przeznaczenia — jeżeli którekolwiek zdanie opisuje projektowaną '
        'funkcję, funkcja jest wyrobem medycznym.')
    for t in OD.KIEDY_MDR:
        doc.add_paragraph(t, style='List Bullet')

    H1('0C.6  Bezpieczne sformułowania interfejsu', 2)
    doc.add_paragraph(
        'Gotowe brzmienia komunikatów po bezpiecznej stronie granicy, zestawione z tym, '
        'co ją przekracza. Warstwa nieobecna w Master 3.0 — do przeniesienia do '
        'projektu interfejsu bez zmian.')
    add_table(doc, OD.BEZPIECZNE)

    H1('0C.7  Warstwa orkiestracji — osiem modułów zarządzających', 2)
    doc.add_paragraph(
        'W Master 3.0 warstwa kontrolna występuje jako cecha funkcji, nie jako byt. '
        'W pełnym tekście oryginału słowa „rejestr”, „silnik reguł”, „proweniencja” '
        'i „terminologia” nie występują ani razu. Rzeczy bez kodu nie mają właściciela, '
        'budżetu ani terminu.')
    add_table(doc, OD.ORKIESTRACJA)

    H1('0C.8  Klasy komponentów K01–K28 — próg wyjścia i decyzja na dzień 1', 2)
    doc.add_paragraph(
        'Funkcje nie mają osobnych dostawców. Sprowadzają się do 28 klas komponentów; '
        'trzy alternatywy wybiera się raz na klasę, nie raz na funkcję. Zamiast 588 '
        'kombinacji pozostaje 28 decyzji. ' + OD.BRAKI_KLAS)
    add_table(doc, OD.KLASY)

    H1('0C.9  Blokady licencyjne', 2)
    doc.add_paragraph(
        'Cztery pozycje blokują zamknięty model komercyjny wprost, cztery kolejne zmieniły '
        'licencje na restrykcyjne po 2021 roku. Każda wymaga decyzji przed rozpoczęciem '
        'prac, nie po.')
    add_table(doc, OD.LICENCJE)

    H1('0C.10  Terminy regulacyjne wiążące dla projektu', 2)
    add_table(doc, OD.TERMINY)
    doc.add_paragraph(OD.TERMIN_RYNKOWY)

    H1('0C.11  Korekty do treści źródłowej', 2)
    doc.add_paragraph(
        'Pozycje, w których treść źródłowa jest błędna. W dokumentach wyjściowych '
        'obowiązuje kolumna trzecia; brzmienie źródłowe pozostaje w częściach I–X, '
        'więc korekta musi być czytana razem z nim.')
    add_table(doc, OD.KOREKTY)

    H1('0C.12  Trzynaście funkcji obowiązkowych w MVP', 2)
    doc.add_paragraph(
        'Funkcje, których nie da się odłożyć. Sześć z nich nie występowało w żadnej '
        'wcześniejszej wersji planu: A18.3, A18.8, A18.10, A18.11, A18.12 i A19.1. '
        'Wszystkie sześć to warstwa A — nie wymagają certyfikacji, a bez nich MVP jest '
        'nielegalne. Trzy z nich — log dostępu, granularne zgody i usunięcie danych — '
        'są jednocześnie wymogami formalnymi i najtańszymi wyróżnikami handlowymi '
        'w całym rejestrze.')
    add_table(doc, OD.MVP_OBOWIAZKOWE)

    H1('0C.13  Szesnaście modułów technicznych: co bierzemy, co zostaje nasze', 2)
    doc.add_paragraph(
        'Najczystsze w korpusie zastosowanie zasady „kupuj funkcję, kontroluj interfejs”. '
        'Dziewięć z szesnastu modułów ma koszt licencji zero.')
    add_table(doc, OD.MODULY16)
    doc.add_paragraph(OD.MODULY16_MVP)

    H1('0C.14  Graf zależności między funkcjami', 2)
    doc.add_paragraph(OD.GRAF)

    H1('0C.15  Do rozstrzygnięcia przed budową', 2)
    for t in OD.PRZED_BUDOWA:
        doc.add_paragraph(t, style='List Bullet')

    H1('0C.16  Pozycje wymagające weryfikacji przed budżetowaniem', 2)
    doc.add_paragraph(
        'Nie zostały potwierdzone źródłowo. Nie należy na nich opierać decyzji '
        'finansowych bez sprawdzenia.')
    for t in OD.DO_WERYFIKACJI:
        doc.add_paragraph(t, style='List Bullet')

    doc.add_page_break()



def _wykluczenie(doc, stat):
    """Nota o warstwie wylaczonej — w kazdym dokumencie, ktory jej dotyka."""
    if not stat['pliki'] and not stat['bloki']:
        return
    doc.add_heading('Warstwa wyłączona z dokumentacji', 2)
    doc.add_paragraph(
        'Specyfikacja Master 5.4 w sekcji 38 wyłącza z dokumentacji warstwę sterowania '
        'zachowaniem ludzi, wpływu na decyzje wyborcze, oddziaływania podprogowego, masowej '
        'implantacji i niejawnego podawania nanotechnologii. Ta sama granica jest zapisana '
        'niezależnie w trzech innych miejscach korpusu: etapy 7–11 roadmapy są oznaczone '
        '[FIKCJA]; dodatek z epikami strategicznymi opisuje je jako „motywy dystopijne, '
        'konflikty fabularne i ostrzeżenia — to nie są realne instrukcje wdrożeniowe”; '
        'a Plan PWNŚ świadomie ich nie zoperacjonalizował, uzasadniając to zdaniem: '
        '„nie da się zbudować dla nich budżetu, listy partnerów i harmonogramu, bo to nie '
        'jest plan firmy”.')
    doc.add_paragraph(
        'Ten dokument tej warstwy nie rozwija. Usunięto %d bloków treści%s. Zachowano '
        'natomiast każdy zapis, który tę warstwę NAZYWA i wyklucza — rejestry epików, '
        'listy skreśleń, argumenty odrzucające — ponieważ to jest zapis granicy i musi '
        'pozostać widoczny. Zachowano też całą treść weterynaryjną dotyczącą transponderów '
        'w standardzie ISO 11784/11785, która z tą warstwą nie ma związku.'
        % (stat['bloki'],
           ' oraz %d plików w całości poświęconych tej warstwie' % len(stat['pliki'])
           if stat['pliki'] else ''))
    if stat['pliki']:
        rows = [['#', 'Plik pominięty w całości', 'Czym jest']]
        for i in sorted(stat['pliki']):
            rows.append([str(i), INV[i]['name'].replace('.txt', '')[:56], wyklucz.PLIKI[i]])
        add_table(doc, rows)
    doc.add_paragraph('Epiki wyłączone — nazwane, żeby nie wróciły przez pomyłkę:')
    add_table(doc, [['Kod epiku', 'Czego dotyczy']] + [[k, o] for k, o in wyklucz.EPIKI])
    doc.add_paragraph(
        'Korpus wskazuje dla nich legalne odpowiedniki i to one są przyjęte '
        'w dokumentacji:')
    add_table(doc, wyklucz.ODPOWIEDNIKI)



def _odczyt_bp(doc, sec):
    """Czesc 0C dla sekcji BIZNESPLAN — ustalenia z pelnego odczytu korpusu."""
    if sec != 'B':
        return
    H = doc.add_heading
    H('CZĘŚĆ 0C — USTALENIA Z PEŁNEGO ODCZYTU KORPUSU', 1)
    doc.add_paragraph(
        'Ta część powstała z odczytu całej treści 159 plików korpusu — 28 618 387 znaków '
        'surowo, 13 020 154 po usunięciu duplikatów i treści powtarzającej się jeden do '
        'jednego. Nie powiela bloków źródłowych z części I–XIII: zawiera wyłącznie '
        'ustalenia, które powstały z zestawienia wielu plików albo prostują treść '
        'źródłową. Dziennik odczytu: tools/konsolidacja/odczyt/USTALENIA_ODCZYT.md.')

    H('0C.1  Hierarchia dokumentów biznesowych', 2)
    add_table(doc, OB.HIERARCHIA)

    H('0C.2  Streszczenie zarządcze w siedmiu wymiarach', 2)
    add_table(doc, OB.STRESZCZENIE)

    H('0C.3  Skala problemu — liczby systemowe', 2)
    add_table(doc, OB.PROBLEM_SKALA)
    doc.add_paragraph(OB.PROBLEM_TEZA)
    doc.add_paragraph(OB.LUKA)

    H('0C.4  Bilans wobec systemu publicznego', 2)
    add_table(doc, OB.BILANS_PANSTWO)
    doc.add_paragraph(OB.BILANS_WZORZEC)

    H('0C.5  Segmenty rynku i zasada prezentacji rynku', 2)
    add_table(doc, OB.SEGMENTY)
    p = doc.add_paragraph(); r = p.add_run(OB.ZASADA_RYNKU); r.bold = True

    H('0C.6  Portfel pierwszej fali i etap zerowy', 2)
    add_table(doc, OB.PORTFEL)
    doc.add_paragraph(OB.ETAP_ZEROWY)
    doc.add_paragraph(OB.USUNIETE)

    H('0C.7  Kanały przychodu i ich ranking', 2)
    add_table(doc, OB.KANALY)
    add_table(doc, OB.RANKING_PRZYCHODU)
    doc.add_paragraph(OB.MARZA_TEZA)
    p = doc.add_paragraph(); r = p.add_run(OB.MODEL_ODRZUCONY); r.bold = True

    H('0C.8  Dlaczego marża sprzętowa nie może być źródłem głównym', 2)
    add_table(doc, OB.MARZA_SPRZET)

    H('0C.9  Arytmetyka abonamentu konsumenckiego', 2)
    add_table(doc, OB.ARYTMETYKA_ABO)
    doc.add_paragraph(OB.ARYTMETYKA_WNIOSEK)

    H('0C.10  Gdzie naprawdę leży zasób', 2)
    add_table(doc, OB.ZASOBY)
    doc.add_paragraph(OB.ZASOB_GLOWNY)
    doc.add_paragraph(OB.GLEBIA)

    H('0C.11  Finansowanie: kolejność źródeł i wczesny przychód', 2)
    add_table(doc, OB.ZRODLA_FINANSOWANIA)
    add_table(doc, OB.WCZESNY_PRZYCHOD)
    add_table(doc, OB.CENNIK_HUB)
    doc.add_paragraph(OB.CENNIK_HUB_WNIOSEK)

    H('0C.12  Dźwignia niepieniężna', 2)
    add_table(doc, OB.DZWIGNIA)
    doc.add_paragraph(OB.DZWIGNIA_WARUNEK)

    H('0C.13  Kontrola technologii, której nie budujemy', 2)
    add_table(doc, OB.KONTROLA8)
    add_table(doc, OB.KONTROLA_PLAN)
    p = doc.add_paragraph(); r = p.add_run(OB.KONTROLA_WARUNEK); r.bold = True
    doc.add_paragraph(OB.MOONSHOT_ARYTMETYKA)

    H('0C.14  Warianty tanie i zestaw podstawowy', 2)
    add_table(doc, OB.WARIANTY_TANIE)
    add_table(doc, OB.ZESTAW_PODSTAWOWY)
    add_table(doc, OB.BEZ_WARIANTU)

    H('0C.15  Budżet okna dziewięćdziesięciu dni', 2)
    add_table(doc, OB.BUDZET90)

    H('0C.16  Struktura kosztów i błędy poprzednich modeli', 2)
    add_table(doc, OB.KOSZTY_STRUKTURA)
    add_table(doc, OB.BLEDY_KOSZTOWE)

    H('0C.17  Korekty liczb biznesowych', 2)
    add_table(doc, OB.KOREKTY)

    H('0C.18  Fosa i konkurencja', 2)
    add_table(doc, OB.FOSA)
    add_table(doc, OB.KONKURENCJA)
    doc.add_paragraph(OB.KONKURENCJA_KALIBRACJA)
    doc.add_paragraph(OB.LUKA_WETERYNARYJNA)

    H('0C.19  Wejście na rynek i bramki decyzyjne', 2)
    doc.add_paragraph(OB.FIZYKA_MARKETINGU)
    add_table(doc, OB.WEJSCIE)
    add_table(doc, OB.BRAMKI)

    H('0C.20  Zespół i struktura podmiotów', 2)
    add_table(doc, OB.ZESPOL)
    add_table(doc, OB.STRUKTURA)
    p = doc.add_paragraph(); r = p.add_run(OB.STATUT); r.bold = True

    H('0C.21  Ekonomika warstwy sprzętowej', 2)
    add_table(doc, OB.STACJA_EKONOMIKA)
    doc.add_paragraph(OB.STACJA_WNIOSEK)
    doc.add_paragraph(OB.SEGMENTY_B2B_STACJA)

    H('0C.22  Czego ten plan nie obiecuje', 2)
    for t in OB.NIE_OBIECUJEMY:
        doc.add_paragraph(t, style='List Bullet')

    doc.add_page_break()


def build(sec, tytul, podtytul, podstawa, nota, wersje, kanon, klastry, reszta_tytul, out):
    PARTS={x[0]:(x[1],x[2],x[3]) for x in json.load(open('build/PARTS_%s.json'%sec))}
    PARTS, WSTAT = wyklucz.filtruj(PARTS)
    doc=Document(); setup(doc)
    for t,sz,b in [("ETERNAL ECOSYSTEM",26,True),(tytul,16,True),(podtytul,12,False)]:
        p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER
        r=p.add_run(t); r.font.size=Pt(sz); r.bold=b
        if sz>=16: r.font.color.rgb=RGBColor.from_string('1F3864')
    p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER
    r=p.add_run("%s\nKonsolidacja %d plików źródłowych\nStan na %s"%(podstawa,len(PARTS),TODAY)); r.font.size=Pt(10)
    doc.add_page_break()
    doc.add_heading("Nota metodyczna — jak powstał ten dokument",1)
    for t in nota: doc.add_paragraph(t)
    if wersje:
        doc.add_heading("Łańcuch wersji i pliki zastąpione",2); add_table(doc,wersje)
    _wykluczenie(doc, WSTAT)
    doc.add_page_break(); doc.add_heading("Spis treści",1); toc(doc); doc.add_page_break()
    _ustalenia(doc, sec)
    _odczyt(doc, sec)
    _odczyt_bp(doc, sec)

    uzyte=set()
    for i,(ktyt,kopis) in kanon:
        if i not in PARTS or not PARTS[i][2]: continue
        doc.add_heading(ktyt,1); doc.add_paragraph(kopis)
        emit(doc, PARTS[i][2], base_level=2,
             src_tag="źródło: #%d %s"%(i, INV[i]['name'].replace('.txt','')))
        uzyte.add(i); doc.add_page_break()

    for tyt,opis,idxs in klastry:
        have=[i for i in idxs if i in PARTS and PARTS[i][2] and i not in uzyte]
        if not have: continue
        doc.add_heading(tyt,1); doc.add_paragraph(opis)
        for i in have:
            st,rola,blocks=PARTS[i]
            doc.add_heading("%s  [#%d]"%(INV[i]['name'].replace('.txt',''),i),2)
            p=doc.add_paragraph(); r=p.add_run("Wkład do tej sekcji: %s"%rola)
            r.font.size=Pt(8.5); r.italic=True
            emit(doc, blocks, base_level=3); uzyte.add(i)
        doc.add_page_break()

    reszta=[i for i in sorted(PARTS) if i not in uzyte and PARTS[i][2]]
    if reszta:
        doc.add_heading(reszta_tytul,1)
        for i in reszta:
            st,rola,blocks=PARTS[i]
            doc.add_heading("%s  [#%d]"%(INV[i]['name'].replace('.txt',''),i),2)
            p=doc.add_paragraph(); r=p.add_run("Wkład: %s"%rola); r.font.size=Pt(8.5); r.italic=True
            emit(doc, blocks, base_level=3)
        doc.add_page_break()

    doc.add_heading("ANEKS A — INDEKS ŹRÓDEŁ TEJ SEKCJI",1)
    doc.add_paragraph("Wszystkie pliki korpusu przypisane do tej sekcji wraz z informacją, co dokładnie "
      "z każdego zostało wzięte i czy plik jest wersją obowiązującą. Pliki ze statusem „zastąpiony” "
      "lub „duplikat” nie wnoszą treści — ich zawartość zawiera się w wersji nowszej, wskazanej w statusie.")
    rows=[["#","Plik","Status","Bloków przyjętych","Co wnosi do tej sekcji"]]
    for i,(s,st,rola) in sorted(M.items()):
        if sec not in s.split(','): continue
        n=len(PARTS[i][2]) if i in PARTS else 0
        rows.append([str(i), INV[i]['name'].replace('.txt','')[:60], st, str(n), rola])
    add_table(doc,rows)
    os.makedirs(os.path.dirname(out),exist_ok=True); doc.save(out)
    ch=sum(len(p.text) for p in doc.paragraphs)+sum(len(c.text) for t in doc.tables for r in t.rows for c in r.cells)
    print('%s -> %d B, %d akapitow, %d tabel, ~%d stron'%(out,os.path.getsize(out),
          len(doc.paragraphs),len(doc.tables),round(ch/1800)))
