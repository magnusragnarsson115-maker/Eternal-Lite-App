# -*- coding: utf-8 -*-
import json, os, sys, datetime
sys.path.insert(0, os.path.dirname(os.path.abspath(__file__)))
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.enum.section import WD_SECTION
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from mapa import M
INV={r['idx']:r for r in json.load(open('INVENTORY.json'))}
TODAY=datetime.date.today().strftime('%d.%m.%Y')

def setup(doc):
    st=doc.styles['Normal']; st.font.name='Calibri'; st.font.size=Pt(10)
    st._element.rPr.rFonts.set(qn('w:eastAsia'),'Calibri')
    st.paragraph_format.space_after=Pt(4); st.paragraph_format.line_spacing=1.08
    for i,(sz,bold,col) in enumerate([(20,True,'1F3864'),(15,True,'2E5496'),(12,True,'2E5496'),(11,True,'404040')],1):
        s=doc.styles['Heading %d'%i]; s.font.name='Calibri'; s.font.size=Pt(sz)
        s.font.bold=bold; s.font.color.rgb=RGBColor.from_string(col)
        s.paragraph_format.space_before=Pt(12 if i<3 else 8); s.paragraph_format.space_after=Pt(4)
        s.paragraph_format.keep_with_next=True

def toc(doc):
    p=doc.add_paragraph(); r=p.add_run()
    f=OxmlElement('w:fldChar'); f.set(qn('w:fldCharType'),'begin'); r._r.append(f)
    i=OxmlElement('w:instrText'); i.set(qn('xml:space'),'preserve')
    i.text=r'TOC \o "1-3" \h \z \u'; r._r.append(i)
    f2=OxmlElement('w:fldChar'); f2.set(qn('w:fldCharType'),'separate'); r._r.append(f2)
    t=OxmlElement('w:t'); t.text='Spis tresci — kliknij prawym > Aktualizuj pole'; r._r.append(t)
    f3=OxmlElement('w:fldChar'); f3.set(qn('w:fldCharType'),'end'); r._r.append(f3)

def add_table(doc, rows):
    rows=[r for r in rows if any(c.strip() for c in r)]
    if not rows: return
    n=max(len(r) for r in rows)
    if n>12: n=12
    t=doc.add_table(rows=0, cols=n); t.style='Table Grid'; t.autofit=True
    for ri,r in enumerate(rows[:400]):
        cells=t.add_row().cells
        for ci in range(n):
            v=r[ci] if ci<len(r) else ''
            cells[ci].text=v[:900]
            for p in cells[ci].paragraphs:
                for run in p.runs:
                    run.font.size=Pt(8)
                    if ri==0: run.font.bold=True

def emit(doc, blocks, base_level=1, src_tag=None):
    if src_tag:
        p=doc.add_paragraph(); r=p.add_run(src_tag); r.font.size=Pt(7.5); r.font.italic=True
        r.font.color.rgb=RGBColor.from_string('808080')
    for kind,lvl,payload in blocks:
        if kind=='h':
            L=min(max((lvl if lvl is not None else 2)+base_level-1, base_level),4)
            doc.add_heading(payload[:300], level=L)
        elif kind=='t':
            add_table(doc,payload)
        else:
            txt=payload
            b=txt.startswith('**') and txt.endswith('**')
            txt=txt.strip('*')
            if txt.startswith(('- ','* ','• ','– ')):
                p=doc.add_paragraph(txt[2:].strip(), style='List Bullet')
            else:
                p=doc.add_paragraph()
                r=p.add_run(txt); r.bold=b
