# -*- coding: utf-8 -*-
"""Architektura modularna Eternal: A1 jako wzorzec, moduly architektury,
modularnosc, strategia integracji zamiast budowy, hierarchia ekosystemu."""
import json
import os
import sys
import collections
import datetime

sys.path.insert(0, os.path.dirname(os.path.abspath(__file__)))
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from mkdocx import setup, toc
from dane_architektura import (MODULY, KONTROLNE, MODULARNOSC, INTEGRACJA, HIERARCHIA,
                               PRZEBIEG, FILTR_TWARDY, PUNKTACJA, TRYBY, LOKALNE_GLOBALNE,
                               BRAMA_BEZPIECZENSTWO, BRAMA_ZASTRZEZENIA, POZIOMY_SPRZEDAZY,
                               REGULA_KATALOGU, DOWOD, PRZEPISANIE, WERDYKTY,
                               REGULA_KRZYWEJ, PULAPKA_E5, ZRODLA_SEKCJI)
from dane_rynek import TEST, POZYCJE, AGREGATORY, A1_FUNKCJE
from dane_moduly import M as MOD

R = json.load(open('build/KOMPONENTY.json'))
A1 = [r for r in R if r['modul'] == 'A1']
TODAY = datetime.date.today().strftime('%d.%m.%Y')
OUT = '/home/user/Eternal-Lite-App/out/ETERNAL_ARCHITEKTURA_MODULARNA.docx'


def cien(cell, hexcol):
    sh = OxmlElement('w:shd')
    sh.set(qn('w:val'), 'clear')
    sh.set(qn('w:fill'), hexcol)
    cell._tc.get_or_add_tcPr().append(sh)


def bogaty(p, txt, size=9.5):
    for i, cz in enumerate(str(txt).split('**')):
        if cz:
            r = p.add_run(cz)
            r.font.size = Pt(size)
            r.bold = bool(i % 2)


def tab(doc, rows, szer, naglowek='1B3A6B'):
    t = doc.add_table(rows=0, cols=len(rows[0]))
    t.style = 'Table Grid'
    for ri, row in enumerate(rows):
        cells = t.add_row().cells
        for ci, val in enumerate(row):
            cells[ci].text = ''
            bogaty(cells[ci].paragraphs[0], val, 8.5)
            if ri == 0:
                for r_ in cells[ci].paragraphs[0].runs:
                    r_.bold = True
                    r_.font.color.rgb = RGBColor.from_string('FFFFFF')
                cien(cells[ci], naglowek)
            cells[ci].width = Cm(szer[ci])
    return t


def pole(doc, etykieta, tresc, akcent=False):
    p = doc.add_paragraph()
    r = p.add_run(etykieta.upper())
    r.font.size = Pt(7.5)
    r.bold = True
    r.font.color.rgb = RGBColor.from_string('B8431F' if akcent else '5D6B8A')
    p2 = doc.add_paragraph()
    p2.paragraph_format.left_indent = Cm(0.4)
    p2.paragraph_format.space_after = Pt(6)
    bogaty(p2, tresc)


def akapit(doc, t, size=10):
    p = doc.add_paragraph()
    bogaty(p, t, size)


doc = Document()
setup(doc)
s = doc.sections[0]
s.orientation = WD_ORIENT.LANDSCAPE
s.page_width, s.page_height = s.page_height, s.page_width
s.left_margin = s.right_margin = Cm(1.8)

for txt, sz, bold, col in [
        ('ETERNAL LIFE', 26, True, 'B8431F'),
        ('Architektura modularna i strategia integracji', 16, True, '1B3A6B'),
        ('Moduł A1 jako wzorzec · adapter, brama, mapper · '
         'jak kontrolować technologię, której się nie buduje', 12, False, None)]:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(txt)
    r.font.size = Pt(sz)
    r.bold = bold
    if col:
        r.font.color.rgb = RGBColor.from_string(col)
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('Stan na %s' % TODAY)
r.font.size = Pt(10)
doc.add_page_break()

doc.add_heading('Nota — druga wersja po audycie pokrycia źródeł', 1)
akapit(doc, 'Pierwsza wersja tego dokumentu powstała bez trzech plików korpusu, które '
 'opisują dokładnie to samo, i opisują lepiej: **#117 ETL-034 „Eternal API Gateway", '
 '#119 ETL-031 „Model orkiestratora" i #118 ETL-032 „Wykonalność naukowa i kontrola '
 'technologii".** Wypadły z mojego pola widzenia, bo nie zawierają kodów funkcji '
 'i przez to nie ma ich w rejestrze funkcji, z którego pracowałem. Ich treść jest '
 'w specyfikacji scalonej — nie było jej w moim rozumowaniu.')
akapit(doc, 'Ta wersja jest przepisana wokół nich. **Jedno twierdzenie z poprzedniej '
 'wersji było błędne i zostało poprawione**: napisałem, że dostawca z własnym oznakowaniem '
 'CE obsłuży naszą funkcję warstwy C w modelu proxy. ETL-031 rozstrzyga inaczej '
 'i ma rację — oznakowanie obejmuje konkretną funkcję w aplikacji producenta i w jego '
 'przeznaczeniu, dane surowe z interfejsu nie są nim objęte, a cudzego oznakowania '
 'nie da się odziedziczyć przez adapter.')
tab(doc, [['Sekcja', 'Pliki korpusu, na których stoi']] + [list(x) for x in ZRODLA_SEKCJI],
    [6.0, 20.0], naglowek='B8431F')
doc.add_page_break()

doc.add_heading('Spis treści', 1)
toc(doc)
doc.add_page_break()

# ---------------------------------------------------------------- 0
doc.add_heading('0. Trzy zdania, od których zależy reszta', 1)
akapit(doc, '**Adapter adaptuje, brama wywołuje, mapper mapuje** — tak, ale ta skrótowa '
 'wersja gubi to, co w każdym z nich jest istotne.')
tab(doc, [['Moduł', 'Skrót', 'Czego skrót nie oddaje']] + [
 ['Adapter', 'adaptuje', 'Jest **jeden na klasę komponentu, nie jeden na dostawcę**. '
  'Dopiero to czyni dostawcę wymienialnym — adapter per dostawca to tylko przepisany '
  'klient API'],
 ['Brama', 'wywołuje', 'Wywoływanie jest jej drugą funkcją. Pierwszą jest '
  '**przypisanie użytkownika do dostawcy** — i to jest funkcja, której nikt nie ma. '
  'Bez niej reguły 33% nie da się wyegzekwować, a marży jednostkowej policzyć'],
 ['Mapper', 'mapuje', 'Mapuje **znaczenia, nie pola**. „HGB", „Hb" i „hemoglobina" '
  'to jedno pojęcie z kodem LOINC. Konwerter formatu przepisuje pola i nie wie, '
  'co przepisał'],
], [3.0, 3.0, 20.0])
akapit(doc, '**„Mapujemy rozwiązania i agregujemy?"** — prawie. Mapujemy **zdolności**, '
 'nie rozwiązania. Reguła z korpusu brzmi: *Eternal nie wiąże się z technologią, tylko '
 'ze zdolnością. Nie „Terra API", lecz „zdolność: dane z urządzeń".* Rozwiązanie jest '
 'implementacją zdolności i wymienia się je bez ruszania mapy. Gdyby mapa była mapą '
 'rozwiązań, każda wymiana dostawcy byłaby przebudową mapy.')

# ---------------------------------------------------------------- 1
doc.add_page_break()
doc.add_heading('1. Moduł A1 rozpisany w całości — wzorzec dla pozostałych', 1)
akapit(doc, 'A1 jest wzorcem, bo ma wszystkie cztery przypadki naraz: funkcję, którą '
 'da się kupić w całości (A1.1), funkcję, której kupić nie wolno z powodu licencji (A1.2), '
 'funkcje, których nikt nie sprzedaje (A1.3, A1.4, A1.6), i funkcję na granicy — '
 'silnik kupny, sens własny (A1.5).')

doc.add_heading('1.1 Dziesięć funkcji i ich odpowiedniki rynkowe', 2)
tab(doc, [['Kod', 'Funkcja', 'Odpowiednik rynkowy', 'Kto to robi', 'Da się kupić?',
           'Nasza decyzja']]
    + [list(x) for x in A1_FUNKCJE], [1.6, 5.0, 3.6, 4.4, 4.6, 7.0])

doc.add_heading('1.2 Co pokrywa CAŁY moduł A1', 2)
akapit(doc, 'Pytanie brzmi: czy jedna firma daje synchronizację, ręczne wprowadzanie, '
 'FHIR, deduplikację i przechowywanie razem. Odpowiedź: **żadna nie daje wszystkiego, '
 'ale Vitalera daje najwięcej.**')
tab(doc, [['Kandydat', 'A1.1', 'A1.2', 'A1.3', 'A1.4', 'A1.5', 'A1.6', 'A1.7', 'A1.8',
           'A1.9', 'A1.10', 'Pokrycie']] + [
 ['Vitalera', '●', '○', '—', '—', '●', '○', '○', '●', '—', '○', '85%'],
 ['Terra / Rook / Junction', '●', '—', '—', '—', '○', '—', '○', '●', '—', '—', '70%'],
 ['Thryve (mio)', '●', '—', '—', '—', '●', '—', '○', '●', '—', '—', '70%'],
 ['HealthKit + Health Connect', '○', '—', '—', '—', '—', '—', '—', '●', '—', '○', '70%'],
 ['Medplum (serwer FHIR)', '—', '—', '○', '○', '●', '●', '—', '—', '—', '●', '55%'],
 ['Własne adaptery GATT', '●', '●', '—', '—', '—', '—', '—', '—', '●', '—', '60%'],
], [5.0, 1.3, 1.3, 1.3, 1.3, 1.3, 1.3, 1.3, 1.3, 1.3, 1.4, 2.4])
akapit(doc, '● pełne pokrycie · ○ częściowe · — brak', 8)
akapit(doc, '**Wniosek dla A1: nie kupujemy jednego dostawcy, tylko składamy trzy warstwy.** '
 'HealthKit i Health Connect od dnia pierwszego (zero kosztu, ~70% przypadków), '
 'Medplum jako serwer FHIR z własnym mapperem, agregator dopiero gdy klient B2B zażąda '
 'Garmina albo Oury. Własne adaptery GATT wchodzą przy progu 3 000 zł/mies albo '
 '5 000 aktywnych userów.')

doc.add_heading('1.3 Adapter A1 — co konkretnie budujemy', 2)
tab(doc, [['Element adaptera', 'Co robi', 'Skąd']] + [
 ['Kontrakt danych', 'Jeden interfejs: pobierz obserwacje dla użytkownika w zakresie dat, '
  'zwróć zasoby FHIR R4B w Eternal Standard', 'Własne — to jest definicja zdolności'],
 ['Implementacja per dostawca', 'HealthKit, Health Connect, Terra, Rook, Junction, '
  'Vitalera, Thryve, GATT — osiem implementacji jednego kontraktu',
  'Własne, po ~5–10 osobodni na dostawcę'],
 ['Normalizacja jednostek', 'mmol/l ↔ mg/dl, bpm, mmHg, kg ↔ lb — z zachowaniem '
  'precyzji źródła', 'Własne, wspólne dla całej klasy'],
 ['Proweniencja', 'Które źródło, kiedy, jaką ścieżką, z jaką wersją mappera',
  'Własne — bez tego rozbieżność jest nierozstrzygalna'],
 ['Rozstrzyganie konfliktów', 'Dwa źródła, to samo tętno o 8:00, różne wartości: '
  'reguła pierwszeństwa plus zapis rozbieżności', 'Własne (Universal Sync)'],
 ['Tryb degradacji', 'Dostawca niedostępny → funkcja działa na danych lokalnych '
  'i mówi użytkownikowi, co przestało działać', 'Własne'],
 ['Test kontraktu', 'Przechodzi na atrapie, bez żadnego dostawcy. To jest dowód, '
  'że adapter naprawdę izoluje', 'Własne'],
], [4.6, 11.0, 10.4])
akapit(doc, 'Koszt własnych adapterów GATT z macierzy dostawców: **45 osobodni ≈ 36 tys. zł**. '
 'To jest cała cena wyjścia z klasy K01 — i jest niższa niż roczny rachunek za Terra '
 'przy 5 000 użytkowników.')

doc.add_heading('1.4 Moduły zabezpieczające wewnątrz A1 — niefunkcjonalne', 2)
akapit(doc, 'Nie są funkcjami, nie ma ich w rejestrze 337 pozycji, a bez nich A1 nie może '
 'zostać wydane. To jest ta część, której korpus nie opisuje, bo klasa V (funkcjonalna) '
 'jest jedyną opisaną z pięciu.')
tab(doc, [['Moduł zabezpieczający', 'Co robi w A1', 'Kiedy', 'Bez tego']] + [
 ['K5 Zgody', 'Zgoda osobna na każde źródło danych — Apple, Garmin, laboratorium. '
  'Wycofanie działa wstecz', 'MVP', 'Zgoda pakietowa jest nieswobodna i upada '
  'przy pierwszej kontroli'],
 ['K6 Dziennik audytowy', 'Zapis każdego pobrania: kto, co, kiedy, na jakiej podstawie',
  'MVP', 'Brak dowodu wobec organu; brak wyróżnika produktowego'],
 ['K11 Bezpieczeństwo', 'Klucze API dostawców w magazynie sekretów, rotacja, '
  'model zagrożeń dla ścieżki synchronizacji', 'MVP', 'Klucz w kodzie źródłowym '
  'to incydent, nie ryzyko'],
 ['K12 Obserwowalność', 'Licznik udziału dostawcy w klasie — egzekucja reguły 33%',
  'MVP', 'Reguła 33% jest zadeklarowana, nie egzekwowana'],
 ['K10 Zgodność i jakość', 'Rejestr SOUP: każdy komponent zewnętrzny z wersją '
  'i licencją', 'Przed dossier', 'Nie da się złożyć dossier — SOUP jest wymogiem'],
 ['K13 Tożsamość', 'Powiązanie konta z urządzeniem; izolacja danych między profilami '
  'w koncie rodzinnym', 'MVP', 'Dane dziecka i rodzica w jednym worku'],
], [4.2, 10.6, 2.6, 8.6], naglowek='B8431F')

# ---------------------------------------------------------------- 2
doc.add_page_break()
doc.add_heading('2. Moduły architektury — działanie, budowa, monetyzacja', 1)
for (kod, nazwa, jedno, dzialanie, przezn, zczego, kto, kiedy, wlasne, kontrola,
     widzi, produkt, monet, stan) in MODULY:
    doc.add_heading('%s — %s' % (kod, nazwa), 2)
    p = doc.add_paragraph()
    r = p.add_run(jedno)
    r.font.size = Pt(11)
    r.bold = True
    r.font.color.rgb = RGBColor.from_string('B8431F')
    pole(doc, 'Działanie', dzialanie)
    pole(doc, 'Przeznaczenie', przezn)
    pole(doc, 'Z czego budować', zczego)
    tab(doc, [['Kto', 'Kiedy', 'Kiedy własne', 'Kontrola', 'Stan wg Master 5.4'],
              [kto, kiedy, wlasne, kontrola, stan]], [4.0, 4.6, 6.0, 3.4, 7.0])
    doc.add_paragraph()
    tab(doc, [['Co widzi użytkownik', 'Osobny produkt?', 'Monetyzacja'],
              [widzi, produkt, monet]], [8.0, 6.0, 11.0], naglowek='5D6B8A')
    doc.add_paragraph()

doc.add_heading('2.1 Brama — przebieg jednego zapytania (ETL-034, plik #117)', 2)
akapit(doc, 'Brama nie pyta o dostawcę — pyta o **zdolność**. Nie „pobierz dane z konta '
 'użytkownika u dostawcy X", tylko „podaj tętno z ostatnich siedmiu dni". '
 'Jedenaście kroków, z których każdy ma warunek zatrzymujący:')
tab(doc, [['Krok', 'Co się dzieje', 'Co blokuje']] + [list(x) for x in PRZEBIEG],
    [4.0, 13.0, 9.0])

doc.add_heading('2.2 Jak brama wybiera — filtr twardy przed punktacją', 2)
akapit(doc, '**Najpierw odrzucenie binarne, dopiero potem ocena.** Odrzucony nie wraca '
 'do punktacji — to jest różnica między filtrem a wagą. Moja poprzednia wersja opisywała '
 'cztery filtry w kolejności; poprawna konstrukcja to siedem kryteriów binarnych, '
 'a następnie ranking ważony.')
tab(doc, [['Kryterium odrzucenia', 'Dlaczego']] + [list(x) for x in FILTR_TWARDY],
    [10.0, 16.0])
doc.add_paragraph()
tab(doc, [['Kryterium punktacji', 'Waga', 'Co mierzy']] + [list(x) for x in PUNKTACJA],
    [7.0, 2.0, 17.0])
doc.add_paragraph()
tab(doc, [['Tryb rozstrzygnięcia', 'Kiedy', 'Jak działa']] + [list(x) for x in TRYBY],
    [4.0, 7.0, 15.0])

doc.add_heading('2.3 Lokalne czy globalne — rezydencja per zdolność', 2)
akapit(doc, '**Podróż użytkownika nie zmienia regionu przetwarzania.** Gdy użytkownik jest '
 'za granicą, jego dane nadal są przetwarzane w regionie macierzystym. Zmienia się tylko '
 'to, co brama może dodatkowo udostępnić na miejscu. Zmiana regionu wymagałaby osobnej '
 'zgody i osobnej podstawy prawnej.')
tab(doc, [['Rodzaj zdolności', 'Reguła', 'Bez wyjątku?']]
    + [list(x) for x in LOKALNE_GLOBALNE], [7.0, 15.0, 4.0])

doc.add_heading('2.4 Dziesięć mechanizmów zabezpieczających bramy', 2)
akapit(doc, 'Pytanie „jak to zabezpiecza" ma dziesięć odpowiedzi, które działają razem. '
 'Żadna z osobna nie wystarcza.')
tab(doc, [['Mechanizm', 'Przed czym chroni']]
    + [list(x) for x in BRAMA_BEZPIECZENSTWO], [8.0, 18.0])

doc.add_heading('2.5 Cztery zastrzeżenia, których nie wolno pominąć', 2)
tab(doc, [['Zastrzeżenie', 'Na czym polega', 'Źródło']]
    + [list(x) for x in BRAMA_ZASTRZEZENIA], [7.0, 15.0, 4.0], naglowek='B8431F')

doc.add_heading('2.6 Trzy poziomy dostępności zdolności — to jest model sprzedaży', 2)
akapit(doc, 'Odpowiedź na pytanie, co jest zintegrowane, co wyszukane, a co składane '
 'na zamówienie. To nie jest kwestia techniczna — to jest cennik.')
tab(doc, [['Poziom', 'Stan', 'Czas', 'Kto płaci', 'Cena']]
    + [list(x) for x in POZIOMY_SPRZEDAZY], [5.4, 8.0, 4.0, 3.6, 5.0])
doc.add_paragraph()
p = doc.add_paragraph()
r = p.add_run(REGULA_KATALOGU[0])
r.font.size = Pt(11); r.bold = True
r.font.color.rgb = RGBColor.from_string('B8431F')
akapit(doc, REGULA_KATALOGU[1])
tab(doc, [['Wariant dla zamawiającego', 'Co dostaje', 'Cena']]
    + [list(x) for x in REGULA_KATALOGU[2]], [6.0, 12.0, 8.0])
akapit(doc, '**To jest odpowiedź na pytanie o mapper integracji i o to, kto za niego '
 'zapłaci.** Katalog rośnie na cudzy koszt, a każdy kolejny klient dostaje więcej '
 'za tę samą cenę. Wyłączności trwałej nie oferujemy, bo trwale blokuje wartość, '
 'którą sprzedajemy wszystkim pozostałym.')

doc.add_heading('2.7 Czternaście modułów kontrolnych', 2)
akapit(doc, 'Lista pochodzi z Master 5.4 wraz z oceną stanu. **Pięć z czternastu nie '
 'istnieje wcale**, jeden jest ukryty w orkiestratorze zamiast być wydzielony. '
 'To nie są moduły opcjonalne — to jest warstwa, bez której nie ma dossier ani sprzedaży B2B.')
tab(doc, [['Kod', 'Moduł', 'Zakres', 'Stan', 'Kiedy budować', 'Kto widzi', 'Monetyzacja']]
    + [list(x) for x in KONTROLNE], [1.4, 4.6, 5.0, 2.6, 4.6, 4.6, 4.2])
brak = sum(1 for k in KONTROLNE if 'NIE ISTNIEJE' in k[3])
akapit(doc, '**%d z 14 modułów kontrolnych nie istnieje.** Wszystkie pięć — zgody, dziennik '
 'audytowy, rejestr, zgodność i jakość, bezpieczeństwo — są warunkiem sprzedaży B2B albo '
 'certyfikacji. To jest największa luka wykonawcza w całej architekturze i nie widać jej '
 'w rejestrze funkcji, bo to nie są funkcje.' % brak)

# ---------------------------------------------------------------- 3
doc.add_page_break()
doc.add_heading('3. Sześć agregatorów — czym się różnią i jak je pogrupować', 1)
akapit(doc, 'Pytanie „czym różni się Terra od Vitalery" ma odpowiedź krótką: '
 '**Terra jest produktem wellness, Vitalera deklaruje własne CE wg MDR.** To nie jest '
 'różnica stopnia, tylko kategorii — decyduje o tym, czy dostawca może obsłużyć funkcję '
 'warstwy C.')
tab(doc, [['Dostawca', 'Grupa', 'Klasa', 'Model rozliczenia', 'Pokrycie A1',
           'Zgodność z Eternal', 'Rozwój', 'Adaptowalność', 'Czym się różni']]
    + [[n, g, k, m, '%d%%' % p, '%d%%' % z, r, a, d]
       for n, g, k, m, p, z, r, a, d in AGREGATORY],
    [4.4, 4.0, 3.0, 4.0, 1.8, 2.0, 4.6, 3.6, 8.6])
akapit(doc, '**Grupowanie**: G1 (Terra, Rook, Junction) to jedna kategoria — ten sam model, '
 'te same SDK źródłowe producentów opasek, klasa wellness, brak własnego CE. '
 'Wybór między nimi jest wyborem cennika, nie zdolności. **G2 (Vitalera, Validic, Thryve) '
 'to inna kategoria** — orientacja na zdalny monitoring pacjenta i profil regulacyjny. '
 'G3 (HealthKit, Health Connect) nie jest dostawcą — to system operacyjny użytkownika, '
 'bez umowy i bez ryzyka odcięcia. G4 (własne GATT) nie ma dostawcy wcale.')
akapit(doc, '**Pułapka**, którą korpus nazywa wprost: trzej dostawcy z G1 sięgający po ten '
 'sam SDK Garmina to **formalnie trzech dostawców i faktycznie jeden punkt awarii**. '
 'Reguła 33% mierzy dostawcę; potrzebny jest drugi licznik na poziomie technologii źródłowej. '
 'Prawdziwa redundancja w klasie K01 to: jeden z G1 + jeden z G2 + G3 + G4 — cztery różne '
 'technologie, nie czterech dostawców tej samej.')

# ---------------------------------------------------------------- 4
doc.add_page_break()
doc.add_heading('4. Test otwartego standardu — jak kontrolować to, czego się nie buduje', 1)
akapit(doc, '**%s**' % TEST[0])
tab(doc, [['Odpowiedź', 'Co z tego wynika'], ['TAK', TEST[1]], ['NIE', TEST[2]]],
    [3.0, 23.0], naglowek='B8431F')
std = sum(1 for p_ in POZYCJE if p_[7] == 'TAK')
akapit(doc, 'Na 22 pozycjach z macierzy dostawców **%d pozwala budować rdzeń, %d nie**. '
 'Te siedem — CGM, baza leków, P1, płatności, smart clothes, CDMO, nanotech — '
 'jest wprost oznaczonych: mogą być funkcjami, nigdy fundamentem. To jest jedyny '
 'mechanizm, który w tej architekturze pełni rolę ochrony przed uzależnieniem, '
 'i jest wcześniejszy niż reguła 33%%: **reguła mówi, ilu dostawców; test mówi, '
 'czy wolno na nich w ogóle stanąć.**' % (std, len(POZYCJE) - std))
tab(doc, [['Pozycja', 'Konkurencja', 'Trzy opcje rynkowe', 'White label / OEM',
           'Wyjście, gdy cała trójka odpadnie', 'Standard', 'Rdzeń?']]
    + [[p_[0], p_[1], ' · '.join(p_[2]), ' · '.join(p_[3]), p_[4], p_[5], p_[7]]
       for p_ in POZYCJE], [4.0, 4.0, 5.4, 4.4, 6.6, 3.0, 1.6])

# ---------------------------------------------------------------- 5
doc.add_page_break()
doc.add_heading('5. Integracja zamiast budowy — mechanizm', 1)
for tyt, krotka, dluga in INTEGRACJA:
    doc.add_heading(tyt, 2)
    p = doc.add_paragraph()
    r = p.add_run(krotka)
    r.font.size = Pt(11)
    r.bold = True
    r.font.color.rgb = RGBColor.from_string('B8431F')
    akapit(doc, dluga)

# ---------------------------------------------------------------- 6
doc.add_page_break()
doc.add_heading('6. Modularność — czy jest potrzebna i komu', 1)
for tyt, krotka, dluga, koszt in MODULARNOSC:
    doc.add_heading(tyt, 2)
    p = doc.add_paragraph()
    r = p.add_run(krotka)
    r.font.size = Pt(11)
    r.bold = True
    r.font.color.rgb = RGBColor.from_string('B8431F')
    akapit(doc, dluga)
    akapit(doc, koszt)

# ---------------------------------------------------------------- 7a
doc.add_page_break()
doc.add_heading('7. Ocena technologii — czego nie budować i dlaczego (ETL-032, plik #118)', 1)
akapit(doc, '**Dwie zasady rozstrzygają większość pozycji, zanim zacznie się je oceniać.** '
 'Pierwsza: brak zaprzeczenia nie jest dowodem — że czegoś nie obalono, znaczy tylko tyle, '
 'że nikt tego nie sprawdził. Druga: cel nie jest technologią. „Chcemy nanoboty" nie jest '
 'celem; celem jest „działać precyzyjnie w tkance bez otwierania ciała" — a to otwiera '
 'listę alternatyw dostępnych dziś.')
tab(doc, [['Stopień dowodu', 'Co znaczy', 'Co wolno na tej podstawie']]
    + [list(x) for x in DOWOD], [6.0, 9.0, 11.0], naglowek='B8431F')

doc.add_heading('7.1 Przepisanie celów z technologii na funkcje', 2)
akapit(doc, 'To ćwiczenie **zdejmuje z planu więcej pozycji niż jakakolwiek analiza '
 'kosztowa** — i jest bezpośrednią odpowiedzią na pytanie, jak nie budować rzeczy '
 'wartych miliardy.')
tab(doc, [['Zapis w macierzy', 'Cel po przepisaniu', 'Co jest dostępne dziś']]
    + [list(x) for x in PRZEPISANIE], [6.0, 8.0, 12.0])
akapit(doc, '**„Nanomedycyna JUŻ JEST na rynku — tylko nie wygląda jak roje robotów. '
 'Nasza rola nie leży w nośniku, tylko w danych o tym, komu i kiedy."** To zdanie '
 'z korpusu jest całą strategią wobec moonshotów w jednym wierszu.')

doc.add_heading('7.2 Werdykty dla pięciu projektów', 2)
tab(doc, [['Projekt', 'Dowód', 'TRL', 'Werdykt', 'Warunek']]
    + [list(x) for x in WERDYKTY], [7.4, 4.0, 2.0, 5.0, 7.6])
doc.add_paragraph()
p = doc.add_paragraph()
r = p.add_run('Pułapka wzorcowa: ' + PULAPKA_E5[0])
r.font.size = Pt(11); r.bold = True
r.font.color.rgb = RGBColor.from_string('B8431F')
akapit(doc, PULAPKA_E5[1])
doc.add_paragraph()
p = doc.add_paragraph()
r = p.add_run(REGULA_KRZYWEJ[0])
r.font.size = Pt(11); r.bold = True
r.font.color.rgb = RGBColor.from_string('B8431F')
akapit(doc, REGULA_KRZYWEJ[1])

# ---------------------------------------------------------------- 7
doc.add_page_break()
doc.add_heading('8. Hierarchia ekosystemu', 1)
akapit(doc, 'Korpus zostawia tę sprawę nierozstrzygniętą: w dotychczasowych pracach '
 'przyjęto „cel → projekt → produkt → funkcja" (projekt nad produktem), a w Macierzy 40 '
 'jest odwrotnie. **Rekomendacja korpusu, którą przyjmuję: produkt → projekt → funkcja**, '
 'bo tylko ta konwencja pozwala przypisać funkcje do budżetów i etapów. Czterdzieści '
 'pozycji z Macierzy to wtedy **inicjatywy**, nie projekty.')
tab(doc, [['Poziom', 'Ile', 'Co to jest', 'Rola']]
    + [list(x) for x in HIERARCHIA], [5.0, 2.6, 9.0, 9.4])
akapit(doc, '**Luka planistyczna wskazana w korpusie i nadal otwarta**: filar Digital Twin '
 'ma w Macierzy 40 tylko dwa projekty, a zawiera oba komponenty kluczowe — bazę pacjenta '
 'i symulację ryzyka. Cały ekosystem zależy od filaru, który nie ma budżetu ani terminu '
 'przed 2028 rokiem. Filar Hub nie ma ani jednego projektu.')

doc.add_heading('8.1 Moonshoty — czym są w tej strukturze', 2)
akapit(doc, 'Moonshot nie jest projektem produktowym. Jest **najlepszym dostępnym '
 'rozwiązaniem komponentu, środowiska albo architektury**, wybranym z horyzontem, '
 'w którym jeszcze nie jest dostępne. Dwa rodzaje:')
tab(doc, [['Rodzaj', 'Przykład', 'Co z nim robimy']] + [
 ['Moonshot komponentowy', 'OpenBCI jako droga do interfejsu mózg-komputer bez setek '
  'mln USD; lab-on-chip; nanotech jako protokół, nie urządzenie',
  'Obserwacja plus tor walidacyjny na czymś tańszym. Korpus wskazuje CGM i Pet Bio-Tag '
  'jako substytut funkcjonalny dla nanotechu — walidujemy protokół, nie technologię'],
 ['Moonshot rynkowy', 'Roboty humanoidalne (Japonia, Korea) — rynek rośnie, '
  'wejście produkcyjne jest bardzo drogie',
  'NIE produkować. Wejść jako warstwa danych zdrowotnych i dystrybucja: robot '
  'towarzyszący jest kolejnym źródłem w klasie K01 i kolejnym kanałem w K21. '
  'Kontrola przez standard, przychód przez integrację'],
], [4.6, 9.0, 12.4])
akapit(doc, '**Zasada wspólna dla obu**: moonshot wchodzi do ekosystemu jako **zdolność '
 'z pustą implementacją** — kontrakt danych i test akceptacyjny istnieją, zanim istnieje '
 'dostawca. Gdy dostawca się pojawia, integracja jest przejściem testu, a nie projektem. '
 'To jest cała treść pomysłu na mapper integracji.')

doc.add_paragraph()
p = doc.add_paragraph()
r = p.add_run('Podstawa: Specyfikacja Master 5.4 (§4.2, §4.3, §7, §11), '
              'ETERNAL_Macierz_Dostawcow.xlsx (22 pozycje, test otwartego standardu), '
              'Eternal_Projekty_P1-P5_definicja.md, Roadmapa v4. '
              'Oceny pokrycia i zgodności są autorskie i wymagają potwierdzenia '
              'u dostawców. Stan na %s.' % TODAY)
r.font.size = Pt(8.5)
r.italic = True
r.font.color.rgb = RGBColor.from_string('5D6B8A')

doc.save(OUT)
zn = (sum(len(p.text) for p in doc.paragraphs)
      + sum(len(c.text) for t in doc.tables for r in t.rows for c in r.cells))
print('%s -> %d B, %d tabel, ~%d stron' % (OUT, os.path.getsize(OUT), len(doc.tables),
                                           max(1, round(zn / 2600))))
