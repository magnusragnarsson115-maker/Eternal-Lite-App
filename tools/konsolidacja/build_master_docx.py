# -*- coding: utf-8 -*-
"""Dokument nadrzedny Eternal — 26 sekcji plus odpowiedzi na pytania strategiczne."""
import json
import os
import sys
import collections
import datetime

sys.path.insert(0, os.path.dirname(os.path.abspath(__file__)))
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from mkdocx import setup, toc
from dane_master import S
from dane_moduly import M, KUBELKI, ODPOWIEDZI

R = json.load(open('build/KOMPONENTY.json'))
CW = collections.Counter(r['warstwa'] for r in R)
TODAY = datetime.date.today().strftime('%d.%m.%Y')
OUT = '/home/user/Eternal-Lite-App/out/ETERNAL_DOKUMENT_NADRZEDNY.docx'


def cien(cell, hexcol):
    sh = OxmlElement('w:shd')
    sh.set(qn('w:val'), 'clear')
    sh.set(qn('w:fill'), hexcol)
    cell._tc.get_or_add_tcPr().append(sh)


def bogaty(p, txt, size=10):
    for i, cz in enumerate(str(txt).split('**')):
        if cz:
            r = p.add_run(cz)
            r.font.size = Pt(size)
            r.bold = bool(i % 2)


def tab(doc, naglowki, wiersze, szer, kolor='1B3A6B'):
    t = doc.add_table(rows=0, cols=len(naglowki))
    t.style = 'Table Grid'
    for ri, row in enumerate([naglowki] + wiersze):
        cells = t.add_row().cells
        for ci, val in enumerate(row):
            cells[ci].text = ''
            bogaty(cells[ci].paragraphs[0], val, 8.5)
            if ri == 0:
                for r_ in cells[ci].paragraphs[0].runs:
                    r_.bold = True
                    r_.font.color.rgb = RGBColor.from_string('FFFFFF')
                cien(cells[ci], kolor)
            cells[ci].width = Cm(szer[ci])
    return t


doc = Document()
setup(doc)
for txt, sz, bold, col in [
        ('ETERNAL LIFE', 26, True, 'B8431F'),
        ('Dokument nadrzędny', 16, True, '1B3A6B'),
        ('Dwadzieścia sześć sekcji — od wizji po załączniki', 12, False, None)]:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(txt)
    r.font.size = Pt(sz)
    r.bold = bold
    if col:
        r.font.color.rgb = RGBColor.from_string(col)
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('%d funkcji · 43 moduły · 30 klas komponentów\nStan na %s' % (len(R), TODAY))
r.font.size = Pt(10)
doc.add_page_break()

# --- odpowiedzi na pytania strategiczne, na poczatku ----------------------
doc.add_heading('Sześć pytań i sześć odpowiedzi', 1)
doc.add_paragraph(
 'Ta część odpowiada na pytania postawione wprost. Reszta dokumentu jest kontekstem '
 'dla tych odpowiedzi.')
for pyt, krotka, dlug, dane in ODPOWIEDZI:
    doc.add_heading(pyt, 2)
    p = doc.add_paragraph()
    r = p.add_run(krotka)
    r.font.size = Pt(11)
    r.bold = True
    r.font.color.rgb = RGBColor.from_string('B8431F')
    for t in (dlug, dane):
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Cm(0.4)
        bogaty(p, t)

doc.add_page_break()
doc.add_heading('Spis treści', 1)
toc(doc)
doc.add_page_break()

for nr, tyt, tresc in S:
    doc.add_heading('%s. %s' % (nr, tyt), 1)
    for el in tresc:
        if isinstance(el, tuple) and el and el[0] == 'T':
            _, nag, wier, szer = el
            tab(doc, nag, wier, szer)
            doc.add_paragraph()
        else:
            p = doc.add_paragraph()
            bogaty(p, el % (len(R), CW['A'], CW['B'], CW['C'])
                   if el.count('%d') == 4 else
                   (el % len(R) if el.count('%d') == 1 else
                    (el % CW['C'] if '%d funkcji w warstwie C' in el else el)))
    if nr not in ('25',):
        doc.add_page_break()

p = doc.add_paragraph()
r = p.add_run('Dokument nadrzędny scala ustalenia ze specyfikacji Master 5.4, Planu '
              'Korporacyjnego 5.1, Biznesplanu rozszerzonego, Planu PWNŚ, Roadmapy v4 '
              'i statutu Fundacji, uzupełnione o macierz komponentów i analizę '
              'poprawności. Stan na %s.' % TODAY)
r.font.size = Pt(8.5)
r.italic = True
r.font.color.rgb = RGBColor.from_string('5D6B8A')

doc.save(OUT)
zn = (sum(len(p.text) for p in doc.paragraphs)
      + sum(len(c.text) for t in doc.tables for r in t.rows for c in r.cells))
print('%s -> %d B, %d sekcji, %d tabel, ~%d stron'
      % (OUT, os.path.getsize(OUT), len(S), len(doc.tables), max(1, round(zn / 2400))))
