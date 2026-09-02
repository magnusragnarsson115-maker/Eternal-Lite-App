# -*- coding: utf-8 -*-
"""Ustalenie z kazdego pliku — 159 pozycji, po kolei, z indeksem i waga."""
import os, sys, collections, datetime
sys.path.insert(0, os.path.dirname(os.path.abspath(__file__)))
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from mkdocx import setup, toc, INV
from mapa import M
from dane_pliki import P

TODAY = datetime.date.today().strftime('%d.%m.%Y')
OUT = '/home/user/Eternal-Lite-App/out/ETERNAL_USTALENIA_PER_PLIK.docx'
W = {'KOR': ('KOREKTA', 'B8431F'), 'ROZ': ('ROZSTRZYGNIĘCIE', '1B3A6B'),
     'NOW': ('NOWE', '2E7D32'), 'RYZ': ('RYZYKO', 'B07419'),
     'POT': ('POTWIERDZENIE', '5D6B8A')}
OPIS = {'KOR': 'poprawia albo obala twierdzenie z innego dokumentu',
        'ROZ': 'zamyka sprawę otwartą', 'NOW': 'wnosi treść nieobecną gdzie indziej',
        'RYZ': 'wskazuje zagrożenie', 'POT': 'potwierdza treść już ujętą, bez nowych twierdzeń'}


def cien(cell, hexcol):
    sh = OxmlElement('w:shd'); sh.set(qn('w:val'), 'clear'); sh.set(qn('w:fill'), hexcol)
    cell._tc.get_or_add_tcPr().append(sh)


doc = Document(); setup(doc)
s = doc.sections[0]; s.orientation = WD_ORIENT.LANDSCAPE
s.page_width, s.page_height = s.page_height, s.page_width
s.left_margin = s.right_margin = Cm(1.6)

for txt, sz, bold, col in [('ETERNAL LIFE', 26, True, 'B8431F'),
        ('Ustalenie z każdego pliku', 16, True, '1B3A6B'),
        ('159 plików korpusu, po kolei, z wagą i indeksem', 12, False, None)]:
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(txt); r.font.size = Pt(sz); r.bold = bold
    if col: r.font.color.rgb = RGBColor.from_string(col)
c = collections.Counter(v[1] for v in P.values())
p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('%d korekt · %d rozstrzygnięć · %d treści nowych · %d ryzyk · %d potwierdzeń\n'
              'Stan na %s' % (c['KOR'], c['ROZ'], c['NOW'], c['RYZ'], c['POT'], TODAY))
r.font.size = Pt(10)
doc.add_page_break()

doc.add_heading('Jak czytać', 1)
for t in ['Każdy z 159 plików korpusu ma tu jeden wpis: co z niego wynika i jaką ma wagę. '
          'Kolejność jest numeryczna, zgodna z indeksem źródeł. Pliki przetwarzane były '
          'w szesnastu paczkach po dziesięć.',
          'Waga POTWIERDZENIE nie znaczy, że plik jest zbędny — znaczy, że jego treść pokrywa '
          'się z tym, co jest już ujęte w dokumentach obowiązujących. Trzydzieści plików ma tę '
          'wagę i to jest zdrowy wynik dla korpusu, w którym większość dokumentów to kolejne '
          'wersje tych samych ustaleń.']:
    doc.add_paragraph(t)
t = doc.add_table(rows=0, cols=3); t.style = 'Table Grid'
for ri, row in enumerate([['Waga', 'Ile plików', 'Co znaczy']]
                         + [[W[k][0], str(c[k]), OPIS[k]] for k in ('KOR','ROZ','NOW','RYZ','POT')]):
    cells = t.add_row().cells
    for ci, v in enumerate(row):
        cells[ci].text = ''
        r = cells[ci].paragraphs[0].add_run(v); r.font.size = Pt(9)
        if ri == 0:
            r.bold = True; r.font.color.rgb = RGBColor.from_string('FFFFFF')
            cien(cells[ci], '1B3A6B')
        cells[ci].width = Cm([4.5, 3.0, 18.5][ci])
doc.add_page_break()
doc.add_heading('Spis treści', 1); toc(doc); doc.add_page_break()

for lo in range(1, 160, 10):
    hi = min(lo + 9, 159)
    doc.add_heading('Paczka %d — pliki #%d–#%d' % ((lo // 10) + 1, lo, hi), 1)
    for i in range(lo, hi + 1):
        if i not in P: continue
        ust, waga = P[i]
        et, kol = W[waga]
        p = doc.add_paragraph()
        r = p.add_run(et); r.font.size = Pt(7.5); r.bold = True
        r.font.color.rgb = RGBColor.from_string(kol)
        r2 = p.add_run('   sekcja %s · %s · %s znaków'
                       % (M[i][0], M[i][1], format(INV[i]['chars'], ',').replace(',', ' ')))
        r2.font.size = Pt(7.5); r2.font.color.rgb = RGBColor.from_string('5D6B8A')
        doc.add_heading('#%d  %s' % (i, INV[i]['name'].replace('.txt', '')[:70]), 2)
        p = doc.add_paragraph(); p.paragraph_format.left_indent = Cm(0.3)
        r = p.add_run(ust); r.font.size = Pt(9.5)
    doc.add_page_break()

doc.save(OUT)
print('%s -> %d B, %d plikow' % (OUT, os.path.getsize(OUT), len(P)))
