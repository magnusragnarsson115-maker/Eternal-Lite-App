# -*- coding: utf-8 -*-
"""Ustalenia z plikow bez kodow funkcji — tresc unikatowa dolaczona do wersji finalnej."""
import json
import os
import sys
import collections
import datetime

sys.path.insert(0, os.path.dirname(os.path.abspath(__file__)))
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from mkdocx import setup, toc, INV
from mapa import M
from dane_ustalenia import U, KAT

TODAY = datetime.date.today().strftime('%d.%m.%Y')
OUT = '/home/user/Eternal-Lite-App/out/ETERNAL_USTALENIA_KORPUSU.docx'


def cien(cell, hexcol):
    sh = OxmlElement('w:shd')
    sh.set(qn('w:val'), 'clear')
    sh.set(qn('w:fill'), hexcol)
    cell._tc.get_or_add_tcPr().append(sh)


def bogaty(p, txt, size=10):
    for i, cz in enumerate(str(txt).split('**')):
        if cz:
            r = p.add_run(cz)
            r.font.size = Pt(size)
            r.bold = bool(i % 2)


def tab(doc, rows, szer, naglowek='1B3A6B'):
    t = doc.add_table(rows=0, cols=len(rows[0]))
    t.style = 'Table Grid'
    for ri, row in enumerate(rows):
        cells = t.add_row().cells
        for ci, val in enumerate(row):
            cells[ci].text = ''
            bogaty(cells[ci].paragraphs[0], val, 8.5)
            if ri == 0:
                for r_ in cells[ci].paragraphs[0].runs:
                    r_.bold = True
                    r_.font.color.rgb = RGBColor.from_string('FFFFFF')
                cien(cells[ci], naglowek)
            cells[ci].width = Cm(szer[ci])
    return t


doc = Document()
setup(doc)
for txt, sz, bold, col in [
        ('ETERNAL LIFE', 26, True, 'B8431F'),
        ('Ustalenia z plików bez kodów funkcji', 16, True, '1B3A6B'),
        ('%d ustaleń z 32 plików, które wypadły z rejestru — treść unikatowa '
         'dołączona do wersji obowiązującej' % len(U), 12, False, None)]:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(txt)
    r.font.size = Pt(sz)
    r.bold = bold
    if col:
        r.font.color.rgb = RGBColor.from_string(col)
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('Stan na %s' % TODAY)
r.font.size = Pt(10)
doc.add_page_break()

doc.add_heading('Dlaczego ten dokument istnieje', 1)
for t in [
 'Rejestr funkcji powstał z plików, które zawierają kody funkcji. **Siedemdziesiąt cztery '
 'pliki ze stu pięćdziesięciu dziewięciu kodów nie zawierają** — i przez to nie było ich '
 'w rejestrze, a więc i w rozumowaniu przy budowie dokumentów analitycznych. '
 'Ich treść była w specyfikacji scalonej. Nie była w żadnym wniosku.',
 'Trzydzieści trzy z tych plików należą do sekcji SPECYFIKACJA. Przeczytane w całości, '
 'dają %d ustaleń, z których **%d obala albo poprawia twierdzenie z wcześniejszych '
 'dokumentów**, %d zamyka sprawę otwartą, %d wnosi treść nieobecną gdzie indziej, '
 'a %d wskazuje ryzyko spoza rejestru ryzyk.'
 % (len(U), sum(1 for u in U if u[1] == 'KOREKTA'),
    sum(1 for u in U if u[1] == 'ROZSTRZ'), sum(1 for u in U if u[1] == 'NOWE'),
    sum(1 for u in U if u[1] == 'RYZYKO')),
 'Zarzut nie jest nowy. **Korpus zawiera własny audyt pokrycia źródeł** (plik #114), '
 'który stawia dokładnie to samo pytanie i kończy się zdaniem: „dopóki nie zostanie '
 'przeczytane i zestawione, każde zdanie o ujednoliceniu wszystkich źródeł jest '
 'nieprawdziwe". Ten dokument jest wykonaniem tamtego zalecenia.',
 'Każde ustalenie ma numer pliku źródłowego. Kolumna „co zmienia" mówi, do którego '
 'dokumentu wynikowego trafia.',
]:
    p = doc.add_paragraph()
    bogaty(p, t)

doc.add_heading('Rozkład ustaleń', 2)
c = collections.Counter(u[1] for u in U)
tab(doc, [['Kategoria', 'Ile', 'Co znaczy']]
    + [[KAT[k][0], str(c[k]), KAT[k][2]] for k in ('KOREKTA', 'ROZSTRZ', 'NOWE', 'RYZYKO')],
    [4.0, 1.6, 11.0])

doc.add_page_break()
doc.add_heading('Spis treści', 1)
toc(doc)
doc.add_page_break()

GRUPY = [
 ('C', 'Certyfikacja, agregacja i granica wyrobu'),
 ('P', 'Państwo, IKP i elektroniczna dokumentacja'),
 ('S', 'Struktura ekosystemu i projekty'),
 ('E', 'Ekonomia i model biznesowy'),
 ('K', 'Kontrola, ład korporacyjny i standard'),
 ('T', 'Technologia i sprzęt'),
 ('D', 'Dane, użytkownik i horyzont'),
 ('M', 'Metodyka i luki w materiale'),
]

for pref, tytul in GRUPY:
    poz = [u for u in U if u[0].startswith(pref)]
    if not poz:
        continue
    doc.add_page_break()
    doc.add_heading('%s (%d)' % (tytul, len(poz)), 1)
    for kod, kat, tyt, ust, zmienia, pliki in poz:
        et, kol, _ = KAT[kat]
        p = doc.add_paragraph()
        r = p.add_run(et)
        r.font.size = Pt(7.5)
        r.bold = True
        r.font.color.rgb = RGBColor.from_string(kol)
        r2 = p.add_run('   ' + pliki)
        r2.font.size = Pt(7.5)
        r2.font.color.rgb = RGBColor.from_string('5D6B8A')
        doc.add_heading('%s. %s' % (kod, tyt), 2)
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Cm(0.3)
        bogaty(p, ust, 10)
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Cm(0.3)
        r = p.add_run('CO ZMIENIA:  ')
        r.font.size = Pt(7.5)
        r.bold = True
        r.font.color.rgb = RGBColor.from_string('B8431F')
        bogaty(p, zmienia, 9.5)

# --- indeks plikow --------------------------------------------------------
doc.add_page_break()
doc.add_heading('Indeks plików przeanalizowanych w tej rundzie', 1)
doc.add_paragraph(
 'Pliki, które nie zawierają kodów funkcji i przez to nie występowały w rejestrze. '
 'Kolumna „ustalenia" podaje kody pozycji z tego dokumentu, które z danego pliku wynikają.')
pliki = collections.defaultdict(list)
for u in U:
    for p_ in u[5].replace('#', '').split(', '):
        pliki[int(p_)].append(u[0])
rows = [['#', 'Plik', 'Sekcja', 'Znaków', 'Ustalenia z tego pliku']]
for i in sorted(pliki):
    rows.append([str(i), INV[i]['name'].replace('.txt', '')[:56], M[i][0],
                 format(INV[i]['chars'], ',').replace(',', ' '), ', '.join(pliki[i])])
tab(doc, rows, [1.2, 7.4, 1.6, 2.0, 4.8])

doc.add_paragraph()
p = doc.add_paragraph()
r = p.add_run('Pozostałe pliki bez kodów funkcji zostały przeczytane i nie wniosły ustaleń '
              'wykraczających poza treść już ujętą — ich zawartość pokrywa się z plikami '
              'wymienionymi wyżej albo z dokumentami obowiązującymi. Stan na %s.' % TODAY)
r.font.size = Pt(8.5)
r.italic = True
r.font.color.rgb = RGBColor.from_string('5D6B8A')

doc.save(OUT)
print('%s -> %d B, %d ustalen, %d plikow, %d tabel'
      % (OUT, os.path.getsize(OUT), len(U), len(pliki), len(doc.tables)))
