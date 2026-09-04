# -*- coding: utf-8 -*-
"""Karty funkcji — 337 kart w szablonie CEO."""
import os, sys, collections
sys.path.insert(0, os.path.dirname(os.path.abspath(__file__)))
from mkdocx import *
from docx import Document
import karty, rejestr as RJ, dane_produkty as PR

K = karty.wszystkie()
doc = Document(); setup(doc)
for t, sz, b in [("ETERNAL ECOSYSTEM", 26, True), ("KARTY FUNKCJI", 16, True),
                 ("%d funkcji w szablonie osiemnastopolowym z warstwą rozszerzoną"
                  % len(K), 11, False)]:
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(t); r.font.size = Pt(sz); r.bold = b
    if sz >= 16: r.font.color.rgb = RGBColor.from_string('1F3864')
p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("Stan na %s" % TODAY); r.font.size = Pt(10)
doc.add_page_break()

doc.add_heading("Jak czytać kartę", 1)
doc.add_paragraph(
 "Osiemnaście pól szablonu zarządczego, a pod nimi warstwa rozszerzona: monetyzacja, "
 "efekt końcowy, rozdzielenie tego, co widzi użytkownik, od tego, co widzimy my, komponenty, "
 "architektura, stack z decyzją build/buy, próg przejścia na własne, kontrola, granica "
 "wellness–wyrób oraz informacja, czy funkcja działa samodzielnie, czy tylko w ekosystemie.")
doc.add_paragraph(
 "Trzydzieści funkcji oznaczonych jako rdzeniowe wchodzi w skład sześciu produktów "
 "i ma treść pisaną osobno. Pozostałe karty są wyprowadzone z rejestru — pola regulacyjne, "
 "bezpieczeństwa i kryteriów akceptacji wynikają z warstwy A, B albo C oraz z klasy "
 "komponentu, więc są spójne w całym zbiorze.")
doc.add_paragraph(
 "Priorytet: P0 to funkcje pierwszej wersji rdzeniowe dla produktu albo o wysokiej wadze "
 "dla ekosystemu; P1 pozostałe pierwszej wersji i ważne z drugiej fali; P2 reszta.")

st = collections.Counter(k['priorytet'] for k in K)
mw = collections.Counter(k['medical'] for k in K)
add_table(doc, [["Miara", "Wartość"],
                ["Kart", str(len(K))],
                ["Rdzeniowych (w produktach)", str(len(karty.RDZEN))],
                ["P0 / P1 / P2", "%d / %d / %d" % (st['P0'], st['P1'], st['P2'])],
                ["Wyrób medyczny TAK / NIE", "%d / %d" % (mw['TAK'], mw['NIE'])]])
doc.add_page_break()

doc.add_heading("Indeks funkcji rdzeniowych produktów", 1)
rows = [["Produkt", "Funkcje"]]
for p in PR.PRODUKTY:
    rows.append(["%s %s" % (p['kod'], p['nazwa']), " · ".join(p['funkcje'])])
add_table(doc, rows)
doc.add_page_break()

POLA = [
 ('Cel', 'cel'), ('Problem', 'problem'), ('Użytkownik', 'uzytkownik'),
 ('Opis funkcji', 'opis'), ('Input', 'input'), ('Output', 'output'),
 ('Przebieg użytkownika', 'przebieg'), ('Integracje', 'integracje'), ('API', 'api'),
 ('Dane', 'dane'), ('Uprawnienia', 'uprawnienia'),
]
ROZSZ = [
 ('Monetyzacja', 'monetyzacja'), ('Efekt końcowy', 'efekt'),
 ('Co widzi użytkownik', 'widzi_user'), ('Co widzimy my', 'widzimy_my'),
 ('Komponenty', 'komponenty'), ('Architektura i infrastruktura', 'architektura'),
 ('Stack i decyzja build/buy', 'stack'), ('Etap i certyfikacja', 'etapy'),
 ('Kontrola technologii', 'kontrola'), ('Wellness czy wyrób', 'med_wellness'),
 ('Moduł', 'modul'), ('Samodzielność', 'samodzielnosc'),
 ('Czas do niezależności', 'czas_wyjscia'),
]

prod_akt = None
for k in K:
    v = RJ.R[k['kod']]
    if v['produkt'] != prod_akt:
        prod_akt = v['produkt']
        doc.add_page_break(); doc.add_heading(prod_akt, 1)
    doc.add_heading("%s  %s" % (k['kod'], k['nazwa']), 2)
    if k['produkt_rdzenny']:
        p = doc.add_paragraph()
        r = p.add_run("Funkcja rdzeniowa: " + ", ".join(k['produkt_rdzenny']))
        r.bold = True; r.font.size = Pt(9)
    for et, key in POLA:
        p = doc.add_paragraph(); r = p.add_run(et + ": "); r.bold = True
        p.add_run(str(k[key]))
    for et, key in [('Bezpieczeństwo', 'bezpieczenstwo'), ('Regulacje', 'regulacje'),
                    ('Kryteria akceptacji', 'kryteria')]:
        p = doc.add_paragraph(); r = p.add_run(et + ":"); r.bold = True
        for it in k[key]:
            doc.add_paragraph(it, style='List Bullet')
    p = doc.add_paragraph()
    r = p.add_run("Czy funkcja jest medical device: "); r.bold = True
    r2 = p.add_run(k['medical']); r2.bold = True
    r2.font.color.rgb = RGBColor.from_string('B8431F' if k['medical'] == 'TAK' else '14602C')
    p.add_run(" — " + k['medical_uzas'])
    add_table(doc, [["Priorytet", "Status", "Owner"],
                    [k['priorytet'], k['status'], k['owner']]])
    doc.add_heading("Warstwa rozszerzona", 3)
    add_table(doc, [["Wymiar", "Ustalenie"]] + [[et, str(k[key])] for et, key in ROZSZ])

out = '/home/user/Eternal-Lite-App/out/ETERNAL_KARTY_FUNKCJI.docx'
doc.save(out)
ch = sum(len(p.text) for p in doc.paragraphs) + sum(
    len(c.text) for t in doc.tables for r in t.rows for c in r.cells)
print('%s -> %d B, %d kart, %d akapitow, %d tabel, ~%d stron' % (
    out, os.path.getsize(out), len(K), len(doc.paragraphs), len(doc.tables), round(ch / 1800)))
