# -*- coding: utf-8 -*-
"""Specyfikacja techniczna — wersja finalna, skrócona.

SPEC-00..SPEC-21 plus Aneks A. Buduje z modułów danych, nie z bloków korpusu:
dokument dowodowy (ponad 9 000 stron) pozostaje w archiwum, ten ma być czytany.
"""
import os, sys, collections
sys.path.insert(0, os.path.dirname(os.path.abspath(__file__)))
from mkdocx import *
from docx import Document
import dane_spec_final as F
import dane_odczyt as O
import dane_odczyt_rm as R
import dane_architektura as A
import dane_komponenty as KP
import dane_produkty as PR
import rejestr as RJ
import karty
import wyklucz

OUT = '/home/user/Eternal-Lite-App/out/ETERNAL_SPECYFIKACJA_FINALNA.docx'
doc = Document(); setup(doc)


def H(n, t): doc.add_heading("%s  %s" % (n, t), 1)
def h2(t): doc.add_heading(t, 2)
def h3(t): doc.add_heading(t, 3)
def P_(t): doc.add_paragraph(t)
def B_(t):
    p = doc.add_paragraph(); r = p.add_run(t); r.bold = True
def LIST(items):
    for it in items: doc.add_paragraph(str(it), style='List Bullet')
def TAB(rows, head=None):
    rows = [list(map(str, r)) for r in rows]
    if head: rows = [list(head)] + rows
    add_table(doc, rows)


# ------------------------------------------------------------------ okładka
for t, sz, b in [("ETERNAL ECOSYSTEM", 26, True), (F.TYTUL, 16, True), (F.PODTYTUL, 11, False)]:
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(t); r.font.size = Pt(sz); r.bold = b
    if sz >= 16: r.font.color.rgb = RGBColor.from_string('1F3864')
p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("Podstawa: pełny odczyt korpusu — 159 plików, 28 618 387 znaków\n"
              "Rejestr operacyjny: %d pozycji · Stan na %s" % (len(RJ.R), TODAY))
r.font.size = Pt(10)
doc.add_page_break()

doc.add_heading("Nota metodyczna", 1)
for t in F.NOTA: P_(t)
doc.add_page_break()
doc.add_heading("Spis treści", 1); toc(doc); doc.add_page_break()

# ------------------------------------------------------------------ SPEC-00
H("SPEC-00.", "Metodyka konsolidacji i deduplikacji")
for t in F.S00: P_(t)
h2("Hierarchia wersji — co obowiązuje i co zostało zastąpione")
TAB(O.HIERARCHIA)
h2("Materiał wyłączony z produkcji dokumentów")
P_("Wyłączenie obejmuje warstwę sterowania zachowaniem ludzi i powiązane epiki. Zapisy, "
   "które tę warstwę nazywają i wykluczają, pozostają w dokumentacji — wykluczenie jest "
   "ustaleniem, nie luką.")
TAB([[str(i), n] for i, n in sorted(wyklucz.PLIKI.items())], ["#", "Plik wyłączony w całości"])
TAB([[k, w] for k, w in wyklucz.EPIKI], ["Epik", "Czego dotyczy"])
if hasattr(wyklucz, 'ODPOWIEDNIKI'):
    h3("Legalne odpowiedniki z korpusu")
    TAB([list(x) for x in wyklucz.ODPOWIEDNIKI])
doc.add_page_break()

# ------------------------------------------------------------------ SPEC-01
H("SPEC-01.", "Zakres produktu i taksonomia kanoniczna")
for t in F.S01: P_(t)
h2("Liczby funkcji — wszystkie ujęcia obok siebie")
TAB(O.LICZBY)
h2("Rejestr operacyjny w przekrojach")
prod = collections.Counter(v['produkt'] for v in RJ.R.values())
war = collections.Counter(v['warstwa'] for v in RJ.R.values())
et = collections.Counter(v['etap'] for v in RJ.R.values())
K = karty.wszystkie()
pri = collections.Counter(k['priorytet'] for k in K)
TAB([["Pozycji w rejestrze", str(len(RJ.R))],
     ["Produkty", " · ".join("%s %d" % (a, b) for a, b in prod.most_common())],
     ["Warstwy regulacyjne", " · ".join("%s %d" % (a, b) for a, b in sorted(war.items()))],
     ["Etapy", " · ".join("%s %d" % (a, b) for a, b in et.most_common())],
     ["Priorytety", "P0 %d · P1 %d · P2 %d" % (pri['P0'], pri['P1'], pri['P2'])]],
    ["Przekrój", "Wartość"])
h2("Warstwy regulacyjne — definicje")
TAB([[k] + list(v) for k, v in sorted(RJ.WARSTWA.items())],
    ["Warstwa", "Zakres", "Reżim"])
h2("Graf zależności funkcji")
P_(O.GRAF)
doc.add_page_break()

# ------------------------------------------------------------------ SPEC-02
H("SPEC-02.", "Przeznaczenie, wellness a wyrób i granica regulacyjna")
for t in F.S02: P_(t)
B_("Reguła granicy: " + O.REGULA_GRANICY)
h2("Cztery statusy regulacyjne")
TAB(O.STATUSY)
h2("Ewolucja funkcji od wellness do wyrobu — koszt każdego etapu")
TAB(O.EWOLUCJA)
h2("Katalog: kiedy funkcja staje się wyrobem medycznym")
P_("Czterdzieści pięć reguł kwalifikacji wyprowadzonych z kart funkcji specyfikacji "
   "kanonicznej. Każda z nich opisuje zmianę, która sama w sobie przesuwa funkcję do "
   "warstwy C.")
LIST(O.KIEDY_MDR)
h2("Bezpieczne sformułowania interfejsu")
TAB(O.BEZPIECZNE)
doc.add_page_break()

# ------------------------------------------------------------------ SPEC-03
H("SPEC-03.", "Stan prawny i terminy zewnętrzne")
for t in F.S03: P_(t)
h2("Terminy ustawowe")
TAB(O.TERMINY)
h2("Kalendarz pełny — terminy ustawowe, zewnętrzne i bramki wewnętrzne")
TAB(R.DATY)
B_(O.TERMIN_RYNKOWY)
doc.add_page_break()

# ------------------------------------------------------------------ SPEC-04
H("SPEC-04.", "Architektura referencyjna ekosystemu")
for t in F.S04: P_(t)
h2("Hierarchia warstw")
TAB([list(x) for x in A.HIERARCHIA], ["Poziom", "#", "Nazwa", "Zakres"])
h2("Przebieg wywołania przez warstwę orkiestracji")
TAB([list(x) for x in A.PRZEBIEG], ["Krok", "Co się dzieje", "Reguła"])
h2("Modularność — dlaczego architektura jest podzielona tak, a nie inaczej")
TAB([list(x) for x in A.MODULARNOSC], ["Pytanie", "Rozstrzygnięcie", "Uzasadnienie"])
doc.add_page_break()

# ------------------------------------------------------------------ SPEC-05
H("SPEC-05.", "Model danych, jakość i proweniencja")
for t in F.S05: P_(t)
h2("Tryby rozstrzygania przy wielu dostawcach tej samej zdolności")
TAB([list(x) for x in A.TRYBY], ["Tryb", "Warunek", "Zachowanie"])
doc.add_page_break()

# ------------------------------------------------------------------ SPEC-06
H("SPEC-06.", "Integracje i brama API")
for t in F.S06: P_(t)
h2("Brama API — co daje każda z zasad")
TAB([list(x) for x in A.BRAMA_BEZPIECZENSTWO], ["Zasada", "Co daje"])
h2("Zastrzeżenia do konstrukcji bramy")
TAB([list(x) for x in A.BRAMA_ZASTRZEZENIA], ["Zastrzeżenie", "Konsekwencja", "Ślad"])
h2("Zasada integracji zamiast budowy")
TAB([list(x) for x in A.INTEGRACJA], ["Pozycja", "Ustalenie", "Uzasadnienie"])
doc.add_page_break()

# ------------------------------------------------------------------ SPEC-07
H("SPEC-07.", "AI, wyszukiwanie z kontekstem, orkiestracja i nadzór modeli")
for t in F.S07: P_(t)
h2("Warstwa orkiestracji K1–K8 — koszt i uzasadnienie")
TAB(O.ORKIESTRACJA)
doc.add_page_break()

# ------------------------------------------------------------------ SPEC-08
H("SPEC-08.", "Bezpieczeństwo, prywatność, zgody i audyt")
for t in F.S08: P_(t)
h2("Zasady bezpieczeństwa architektury komponentowej")
TAB([list(x) for x in KP.BEZPIECZENSTWO], ["Zasada", "Wykonanie", "Podstawa"])
doc.add_page_break()

# ------------------------------------------------------------------ SPEC-09
H("SPEC-09.", "Architektura zgodności: MDR, EHDS, AI Act, KSC, EUDAMED")
for t in F.S09: P_(t)
B_("Reguła proxy: " + O.REGULA_PROXY)
h2("Licencje komponentów obcych — skutek prawny i zamiennik")
TAB(O.LICENCJE)
h2("Szczeble zależności od dostawcy a rola regulacyjna")
TAB([[str(k)] + list(v) for k, v in sorted(KP.SZCZEBEL.items())],
    ["Szczebel", "Pozycja", "Ryzyko", "Rola wg MDR"])
doc.add_page_break()

# ------------------------------------------------------------------ SPEC-10..14
PRODUKTY_SEK = [
    ("SPEC-10.", "Eternal App — moduły i funkcje", "App", F.S10),
    ("SPEC-11.", "Eternal Station — moduły i funkcje", "Station", F.S11),
    ("SPEC-12.", "Eternal Capsule — moduły i funkcje", "Capsule", F.S12),
    ("SPEC-13.", "Eternal Digital Twin — moduły i funkcje", "Twin", F.S13),
    ("SPEC-14.", "Eternal Matrix — moduły i funkcje", "Matrix", F.S14),
]
for nr, tyt, prod_kod, proza in PRODUKTY_SEK:
    H(nr, tyt)
    for t in proza: P_(t)
    poz = [v for v in RJ.R.values() if v['produkt'].startswith(prod_kod) or v['produkt'] == prod_kod]
    if not poz:
        poz = [v for v in RJ.R.values() if prod_kod.lower() in str(v['produkt']).lower()]
    mods = collections.OrderedDict()
    for v in sorted(poz, key=lambda x: x['kod']):
        mods.setdefault(v['modul'], []).append(v)
    h2("Moduły produktu — %d pozycji rejestru w %d modułach" % (len(poz), len(mods)))
    TAB([[m, str(len(fs)),
          " · ".join(sorted({f['warstwa'] for f in fs})),
          ", ".join(f['kod'] for f in fs)[:850]] for m, fs in mods.items()],
        ["Moduł", "Pozycji", "Warstwy", "Kody funkcji"])
    h2("Funkcje produktu")
    TAB([[v['kod'], v['nazwa'][:110], v['etap'], v['warstwa'],
          next((k['priorytet'] for k in K if k['kod'] == v['kod']), '—')]
         for v in sorted(poz, key=lambda x: x['kod'])],
        ["Kod", "Nazwa", "Etap", "Warstwa", "Priorytet"])
    if prod_kod == 'App':
        h2("Funkcje obowiązkowe od pierwszego dnia")
        TAB(O.MVP_OBOWIAZKOWE)
    doc.add_page_break()

# ------------------------------------------------------------------ SPEC-15
H("SPEC-15.", "Kontrola technologii: budować, kupić, zintegrować")
for t in F.S15: P_(t)
h2("Klasy komponentów K01–K28 — próg wyjścia i rekomendacja startowa")
TAB(O.KLASY)
h2("Ekonomia komponentów — koszt startowy i liczba funkcji zależnych")
TAB([list(map(str, x)) for x in KP.EKONOMIA],
    ["Klasa", "Rozwiązanie startowe", "Model kosztu", "Koszt mies.", "Funkcji", "Podstawa"])
h2("Szesnaście modułów technicznych — rozwiązanie gotowe czy własne IP")
TAB(O.MODULY16)
P_(O.MODULY16_MVP)
h2("Moduły kontrolne K1–K14")
TAB([list(x) for x in A.KONTROLNE],
    ["Kod", "Moduł", "Zakres", "Gotowe?", "Etap", "Co widzi użytkownik", "Monetyzacja"])
h2("Decyzja budować czy kupić — reguła i przykłady z rejestru")
TAB(PR.BUILD_BUY)
h2("Alternatywy wobec budowy własnej")
TAB(PR.ALTERNATYWY)
doc.add_page_break()

# ------------------------------------------------------------------ SPEC-16
H("SPEC-16.", "Roadmapa wykonawcza i fazy")
for t in F.S16: P_(t)
B_(R.ZASADA if isinstance(R.ZASADA, str) else "")
h2("Pięć torów równoległych")
TAB(R.TORY)
h2("Horyzont 0 — do końca 2026")
TAB(R.H0)
h2("Horyzont 1")
LIST(R.H1)
h2("Horyzont 2")
LIST(R.H2)
h2("Horyzont 3")
LIST(R.H3)
h2("Horyzont 4 — pozycje odłożone i warunki reaktywacji")
TAB(R.H4)
h2("Czego nie robimy")
TAB(R.NIE_ROBIMY)
h2("Budżet horyzontu zerowego")
TAB(R.BUDZET)
if hasattr(R, 'BUDZET_NOTA'): P_(R.BUDZET_NOTA)
h2("Zmiany wobec wcześniejszych wersji roadmapy")
TAB(R.ZMIANY)
doc.add_page_break()

# ------------------------------------------------------------------ SPEC-17
H("SPEC-17.", "Jakość, testy i cykl życia oprogramowania")
for t in F.S17: P_(t)
h2("Co musi powstać przed pierwszą linią kodu")
LIST(O.PRZED_BUDOWA)
doc.add_page_break()

# ------------------------------------------------------------------ SPEC-18
H("SPEC-18.", "Operacje, obserwowalność i odporność")
for t in F.S18: P_(t)
doc.add_page_break()

# ------------------------------------------------------------------ SPEC-19
H("SPEC-19.", "Konflikty, otwarte decyzje i kryteria akceptacji")
for t in F.S19: P_(t)
h2("Korekty do treści źródłowej — co było błędne i co obowiązuje")
TAB(O.KOREKTY)
h2("Do weryfikacji przed przyjęciem założeń")
LIST(O.DO_WERYFIKACJI)
h2("Werdykty wykonalności obszarów")
TAB([list(x) for x in A.WERDYKTY],
    ["Obszar", "Ocena", "Punkty", "Werdykt", "Czego wymaga"])
doc.add_page_break()

# ------------------------------------------------------------------ SPEC-20
H("SPEC-20.", "Produkty z korelacji funkcji i podział na moduły")
for t in F.S20: P_(t)
h2("Metoda doboru funkcji do produktu")
TAB(PR.METODA)
for p_ in PR.PRODUKTY:
    h2("%s  %s" % (p_['kod'], p_['nazwa']))
    B_(p_['claim'])
    TAB([[f, RJ.R[f]['nazwa'][:110], RJ.R[f]['etap'], RJ.R[f]['warstwa']]
         for f in p_['funkcje']], ["Funkcja", "Nazwa", "Etap", "Warstwa"])
    for etk, tr in [("Niezastępowalność", 'niezast'), ("Automatyzm", 'automat'),
                    ("Rozwój", 'rozwoj'), ("Personalizacja", 'person'),
                    ("Odbiorca", 'odbiorca'), ("Monetyzacja", 'monetyzacja'),
                    ("Samodzielność", 'sam'), ("Warstwa", 'warstwa'),
                    ("Etap", 'etap'), ("Ryzyko", 'ryzyko')]:
        if tr in p_:
            pp = doc.add_paragraph(); rr = pp.add_run(etk + ". "); rr.bold = True
            pp.add_run(str(p_[tr]))
P_(PR.PRODUKT_SIODMY)
h2("Kiedy moduł zostaje modułem, a kiedy staje się produktem")
TAB(PR.MODUL_NA_PRODUKT)
P_(PR.MODUL_NA_PRODUKT_WNIOSEK)
h2("Monetyzacja produktów")
TAB(PR.MONETYZACJA)
P_(PR.MONETYZACJA_ZASADA)
h2("Nisze i branże — te same funkcje, inna korelacja")
TAB(PR.BRANZE)
P_(PR.BRANZE_ZASADA)
h2("Dobór zestawu pod klienta")
TAB(PR.DOBOR)
P_(PR.DOBOR_ZASADA)
doc.add_page_break()

# ------------------------------------------------------------------ SPEC-21
H("SPEC-21.", "Indeks źródeł i proweniencja")
for t in F.S21: P_(t)
h2("Źródła zewnętrzne użyte do aktualizacji stanu")
TAB(F.ZRODLA_ZEW)
doc.add_page_break()

# ------------------------------------------------------------------ ANEKS A
doc.add_heading("ANEKS A — REJESTR FUNKCJI I KARTY PRIORYTETU P0", 1)
for t in F.ANEKS_A_NOTA: P_(t)
h2("Rejestr operacyjny — %d pozycji" % len(RJ.R))
PRIO = {k['kod']: k['priorytet'] for k in K}
TAB([[v['kod'], v['nazwa'][:100], v['produkt'], v['modul'], v['etap'], v['warstwa'],
      PRIO.get(v['kod'], '—'), str(v.get('kanal', '—'))[:40]]
     for v in sorted(RJ.R.values(), key=lambda x: (x['produkt'], x['kod']))],
    ["Kod", "Nazwa", "Produkt", "Moduł", "Etap", "Warstwa", "Priorytet", "Kanał"])
doc.add_page_break()

POLA = [('Cel', 'cel'), ('Problem', 'problem'), ('Użytkownik', 'uzytkownik'),
        ('Opis funkcji', 'opis'), ('Input', 'input'), ('Output', 'output'),
        ('Przebieg użytkownika', 'przebieg'), ('Integracje', 'integracje'), ('API', 'api'),
        ('Dane', 'dane'), ('Uprawnienia', 'uprawnienia')]

P0 = [k for k in K if k['priorytet'] == 'P0']
doc.add_heading("Karty funkcji priorytetu P0 — %d pozycji" % len(P0), 1)
prod_akt = None
for k in P0:
    v = RJ.R[k['kod']]
    if v['produkt'] != prod_akt:
        prod_akt = v['produkt']; h2(prod_akt)
    h3("%s  %s" % (k['kod'], k['nazwa']))
    if k['produkt_rdzenny']:
        pp = doc.add_paragraph()
        rr = pp.add_run("Funkcja rdzeniowa: " + ", ".join(k['produkt_rdzenny']))
        rr.bold = True; rr.font.size = Pt(9)
    for etk, key in POLA:
        pp = doc.add_paragraph(); rr = pp.add_run(etk + ": "); rr.bold = True
        pp.add_run(str(k[key]))
    for etk, key in [('Bezpieczeństwo', 'bezpieczenstwo'), ('Regulacje', 'regulacje'),
                     ('Kryteria akceptacji', 'kryteria')]:
        pp = doc.add_paragraph(); rr = pp.add_run(etk + ":"); rr.bold = True
        LIST(k[key])
    pp = doc.add_paragraph()
    rr = pp.add_run("Czy funkcja jest medical device: "); rr.bold = True
    r2 = pp.add_run(k['medical']); r2.bold = True
    r2.font.color.rgb = RGBColor.from_string('B8431F' if k['medical'] == 'TAK' else '14602C')
    pp.add_run(" — " + k['medical_uzas'])
    TAB([[k['priorytet'], k['status'], k['owner']]], ["Priorytet", "Status", "Owner"])

P1 = [k for k in K if k['priorytet'] == 'P1']
doc.add_page_break()
doc.add_heading("Karty skrócone priorytetu P1 — %d pozycji" % len(P1), 1)
P_("Wersja skrócona zawiera sześć pól rozstrzygających: cel, opis funkcji, wejście, wyjście, "
   "status wyrobu medycznego i kryteria akceptacji. Pełne karty w szablonie osiemnastopolowym "
   "znajdują się w dokumencie kart funkcji.")
prod_akt = None
for k in P1:
    v = RJ.R[k['kod']]
    if v['produkt'] != prod_akt:
        prod_akt = v['produkt']; h2(prod_akt)
    h3("%s  %s" % (k['kod'], k['nazwa']))
    for etk, key in [('Cel', 'cel'), ('Opis funkcji', 'opis'),
                     ('Input', 'input'), ('Output', 'output')]:
        pp = doc.add_paragraph(); rr = pp.add_run(etk + ": "); rr.bold = True
        pp.add_run(str(k[key]))
    pp = doc.add_paragraph()
    rr = pp.add_run("Medical device: "); rr.bold = True
    r2 = pp.add_run(k['medical']); r2.bold = True
    r2.font.color.rgb = RGBColor.from_string('B8431F' if k['medical'] == 'TAK' else '14602C')
    pp.add_run(" — " + k['medical_uzas'])
    pp = doc.add_paragraph(); rr = pp.add_run("Kryteria akceptacji:"); rr.bold = True
    LIST(k['kryteria'])

doc.save(OUT)
ch = sum(len(p.text) for p in doc.paragraphs) + sum(
    len(c.text) for t in doc.tables for r in t.rows for c in r.cells)
print('%s -> %d B, %d akapitow, %d tabel, ~%d stron' % (
    OUT, os.path.getsize(OUT), len(doc.paragraphs), len(doc.tables), round(ch / 1800)))
