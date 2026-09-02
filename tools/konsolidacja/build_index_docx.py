# -*- coding: utf-8 -*-
"""Indeks 159 plikow korpusu w podziale na 4 sekcje — DOCX.

HTML zostaje wylacznie dla roadmapy; indeks jest dokumentem do czytania i druku,
wiec idzie w DOCX razem ze specyfikacja, biznesplanem i analiza.
"""
import json
import os
import sys
import datetime

sys.path.insert(0, os.path.dirname(os.path.abspath(__file__)))
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from mkdocx import setup, toc
from mapa import M
from dane_ustalenia import U

INV = {r['idx']: r for r in json.load(open('INVENTORY.json'))}
TODAY = datetime.date.today().strftime('%d.%m.%Y')
OUT = '/home/user/Eternal-Lite-App/out/ETERNAL_INDEKS_ZRODEL.docx'

SEK = [('B', 'BIZNESPLAN', 'ETERNAL_BIZNESPLAN_SCALONY.docx'),
       ('R', 'ROADMAPA', 'ETERNAL_ROADMAPA_SCALONA.html + ETERNAL_ROADMAPA_APLIKACJA.html'),
       ('S', 'SPECYFIKACJA', 'ETERNAL_SPECYFIKACJA_SCALONA.docx + ETERNAL_SPECYFIKACJA_TEMATYCZNA.docx'),
       ('P', 'PITCH DECK', 'ETERNAL_PITCH_APLIKACJA.pptx + ETERNAL_PITCH_EKOSYSTEM.pptx')]

PARTS = {}
for k, _, _ in SEK:
    try:
        PARTS[k] = {x[0]: x[3] for x in json.load(open('build/PARTS_%s.json' % k))}
    except Exception:
        PARTS[k] = {}

ROWS = [(i, INV[i]['name'].replace('.txt', ''), INV[i]['chars'], st, s.split(','), rola)
        for i, (s, st, rola) in sorted(M.items())]

# ustalenia per plik — z rundy analizy plikow bez kodow funkcji
UST = {}
for u in U:
    for p_ in u[5].replace('#', '').split(', '):
        UST.setdefault(int(p_), []).append(u[0])


def cien(cell, hexcol):
    sh = OxmlElement('w:shd')
    sh.set(qn('w:val'), 'clear')
    sh.set(qn('w:fill'), hexcol)
    cell._tc.get_or_add_tcPr().append(sh)


TLA = {'FINAL': 'D7F0DD', 'UNIKAT': 'DFE8FB', 'ZASTAPIONY': 'F3E3C3',
       'DUPLIKAT': 'F8D7DA', 'SUROWIEC': 'EFEFEF'}


def tabela(doc, rows, szer, statuscol=None):
    t = doc.add_table(rows=0, cols=len(rows[0]))
    t.style = 'Table Grid'
    for ri, row in enumerate(rows):
        cells = t.add_row().cells
        for ci, val in enumerate(row):
            cells[ci].text = ''
            p = cells[ci].paragraphs[0]
            r = p.add_run(str(val))
            r.font.size = Pt(8 if ri else 8.5)
            if ri == 0:
                r.bold = True
                r.font.color.rgb = RGBColor.from_string('FFFFFF')
                cien(cells[ci], '1B3A6B')
            elif statuscol is not None and ci == statuscol:
                cien(cells[ci], TLA.get(str(val).split(':')[0], 'FFFFFF'))
            cells[ci].width = Cm(szer[ci])
    return t


doc = Document()
setup(doc)
s = doc.sections[0]
s.orientation = WD_ORIENT.LANDSCAPE
s.page_width, s.page_height = s.page_height, s.page_width
s.left_margin = s.right_margin = Cm(1.6)

for txt, sz, bold, col in [('ETERNAL LIFE', 26, True, 'B8431F'),
                           ('Indeks źródeł — 159 plików korpusu', 16, True, '1B3A6B'),
                           ('Co z którego pliku weszło do którego dokumentu', 12, False, None)]:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(txt)
    r.font.size = Pt(sz)
    r.bold = bold
    if col:
        r.font.color.rgb = RGBColor.from_string(col)
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('28,6 mln znaków · 4 sekcje docelowe · stan na %s' % TODAY)
r.font.size = Pt(10)
doc.add_page_break()

doc.add_heading('Metoda porządkowania', 1)
zast = sum(1 for x in ROWS if x[3].split(':')[0] in ('ZASTAPIONY', 'DUPLIKAT'))
tabela(doc, [['Liczba', 'Co oznacza'],
             ['159', 'unikalnych plików po deduplikacji MD5'],
             ['28,6 mln', 'znaków korpusu'],
             ['4', 'sekcje docelowe: biznesplan, roadmapa, specyfikacja, pitch'],
             [str(zast), 'plików zastąpionych lub duplikatów']], [3.5, 20.0])

doc.add_heading('Znaczenie statusów', 2)
tabela(doc, [['Status', 'Znaczenie'],
             ['FINAL', 'wersja obowiązująca, kanon dla swojego obszaru'],
             ['UNIKAT', 'jedyna wersja tego materiału — treść wchodzi w całości'],
             ['ZASTAPIONY:n', 'istnieje wersja nowsza (plik nr n); powielana jest wyłącznie treść, '
                              'której nowsza wersja nie zawiera'],
             ['DUPLIKAT:n', 'kopia pliku nr n'],
             ['SUROWIEC', 'materiał źródłowy (konwersacje, listy pytań) o niższym statusie '
                          'niż dokumenty scalone']], [4.0, 19.5], statuscol=0)

for t in [
    'Zasada rozstrzygania: przy kilku wersjach tego samego dokumentu obowiązuje najnowsza, ale '
    'wersje wcześniejsze są wczytywane i sprawdzane pod kątem treści, której nowsza nie zawiera — '
    'i tylko taka treść jest dobierana. Etykieta statusu decyduje o pierwszeństwie, nigdy '
    'o pominięciu pliku.',
    'Weryfikacja jest mechaniczna: każdy blok tekstu sprowadzany jest do postaci znormalizowanej '
    'i porównywany z blokami już przyjętymi. Dzięki temu twierdzenie „App 5.4 zawiera się '
    'w Master 5.4” nie jest oceną, tylko wynikiem pomiaru — 89% jej bloków powtarza się dosłownie.',
    'Bloki krótsze niż 40 znaków znormalizowanych (etykiety pól kart funkcji, nagłówki tabel) są '
    'zachowywane mimo powtarzalności, ponieważ ich powtórzenia są strukturalne — bez nich '
    'rozpadłyby się karty 185 funkcji w biznesplanie rozszerzonym.',
    'Drugi przebieg (filtr artefaktów) łapie duplikaty międzyformatowe: ten sam dokument w PDF '
    'i w DOCX tnie się na inne bloki, więc hasze się nie zgadzają, choć treść jest identyczna.',
    'RUNDA ANALIZY PLIKÓW BEZ KODÓW FUNKCJI. Siedemdziesiąt cztery pliki nie zawierają kodów '
    'funkcji i przez to nie występowały w rejestrze, z którego budowano dokumenty analityczne. '
    'Przeczytane osobno, dały %d ustaleń z 32 plików — kolumna „Ustalenia" wskazuje, które '
    'pozycje z ETERNAL_USTALENIA_KORPUSU.docx z danego pliku wynikają. Pliki bez wpisu '
    'w tej kolumnie nie wniosły treści wykraczającej poza już ujętą.' % len(U),
]:
    doc.add_paragraph(t)

doc.add_page_break()
doc.add_heading('Spis treści', 1)
toc(doc)
doc.add_page_break()

for k, nazwa, wynik in SEK:
    r = [x for x in ROWS if k in x[4]]
    dup = [x for x in r if x[3].split(':')[0] in ('ZASTAPIONY', 'DUPLIKAT')]
    doc.add_heading('%s — %d plików' % (nazwa, len(r)), 1)
    doc.add_paragraph('Dokument wynikowy: %s. Plików wnoszących treść: %d. Plików zastąpionych '
                      'lub duplikatów: %d — ich treść zawiera się w wersji nowszej wskazanej '
                      'w statusie.' % (wynik, len(r) - len(dup), len(dup)))
    tabela(doc, [['#', 'Plik źródłowy', 'Znaków', 'Status', 'Bloków', 'Ustalenia',
                  'Co z niego wchodzi do tej sekcji']]
           + [[str(i), n[:70], format(c, ',').replace(',', ' '), st,
               str(len(PARTS[k].get(i, []))), ', '.join(UST.get(i, [])) or '—', rola]
              for i, n, c, st, ss, rola in r],
           [1.1, 6.4, 1.7, 2.4, 1.4, 2.8, 8.2], statuscol=3)
    doc.add_page_break()

multi = [x for x in ROWS if len(x[4]) > 1]
NAZWY = {k: n for k, n, _ in SEK}
doc.add_heading('Pliki występujące w wielu sekcjach — %d pozycji' % len(multi), 1)
doc.add_paragraph('Ten sam plik może zasilać kilka sekcji, ale każda bierze z niego co innego. '
                  'Kolumna po prawej mówi, co dokładnie jest brane pod uwagę.')
tabela(doc, [['#', 'Plik', 'Sekcje', 'Co jest z niego brane']]
       + [[str(i), n[:70], ' + '.join(NAZWY[y] for y in ss), rola]
          for i, n, c, st, ss, rola in multi],
       [1.1, 7.5, 6.0, 8.9])

doc.save(OUT)
print('%s -> %d B, %d tabel, plikow w wielu sekcjach: %d'
      % (OUT, os.path.getsize(OUT), len(doc.tables), len(multi)))
